#!/usr/bin/env python3
"""
Cloud-safe Streamlit dashboard for the Interpreter Audio Generator.

Deploy with Streamlit Community Cloud using this file as the entrypoint, or run locally:
    streamlit run audio_generator_streamlit_cloud_V2.py

Generated files are session/runtime files. Download ZIP/CSV/JSON outputs if you want to keep them.
"""
from __future__ import annotations

import asyncio
import contextlib
import importlib.util
import io
import json
import os
import shutil
import tempfile
import csv
import subprocess
import zipfile
import time
import base64
import hashlib
import re
import uuid
import difflib
import html
import unicodedata
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

import streamlit as st
import streamlit.components.v1 as components
import pandas as pd

APP_DIR = Path(__file__).resolve().parent
CLOUD_MODE = True
RUNTIME_DIR = Path(tempfile.gettempdir()) / "interpreter_audio_generator_web"
RUNTIME_OUTPUT_DIR = RUNTIME_DIR / "audio_output"
RUNTIME_CACHE_DIR = RUNTIME_DIR / "cache"
DEFAULT_GENERATOR_PATH = APP_DIR / "Current_Audio_Generator_Script_JSON_to_MP3_V7.py"
DEFAULT_STUDY_PROGRESS_DIR = RUNTIME_DIR / "study_progress"
DEFAULT_PRACTICE_LOG_PATH = DEFAULT_STUDY_PROGRESS_DIR / "practice_log.csv"
DEFAULT_TERM_REVIEW_LOG_PATH = DEFAULT_STUDY_PROGRESS_DIR / "term_review_log.csv"
DEFAULT_TERM_SRS_STATE_PATH = DEFAULT_STUDY_PROGRESS_DIR / "term_srs_state.csv"
DEFAULT_SIMULTANEOUS_LOG_PATH = DEFAULT_STUDY_PROGRESS_DIR / "simultaneous_practice_log.csv"
DEFAULT_DRILL_STUDIO_LOG_PATH = DEFAULT_STUDY_PROGRESS_DIR / "drill_studio_log.csv"
DEFAULT_EXAM_OUTPUT_DIR = RUNTIME_DIR / "exam_output"
DEFAULT_EXAM_TRANSCRIPT_DIR = DEFAULT_EXAM_OUTPUT_DIR / "transcripts"
OPENAI_TRANSCRIBE_MODELS = ["whisper-1", "gpt-4o-mini-transcribe", "gpt-4o-transcribe"]
OPENAI_AI_FEEDBACK_MODELS = ["gpt-5.5", "gpt-5", "gpt-4o-mini", "gpt-4o"]
OPENAI_TRANSCRIBE_LIMIT_BYTES = 25 * 1024 * 1024

for _runtime_path in [RUNTIME_DIR, RUNTIME_OUTPUT_DIR, RUNTIME_CACHE_DIR, DEFAULT_STUDY_PROGRESS_DIR, DEFAULT_EXAM_OUTPUT_DIR, DEFAULT_EXAM_TRANSCRIPT_DIR]:
    try:
        _runtime_path.mkdir(parents=True, exist_ok=True)
    except Exception:
        pass

APPROVED_EDGE_VOICES = {
    "en-US-AvaNeural", "en-US-AndrewNeural", "en-US-EmmaNeural", "en-US-BrianNeural",
    "en-US-AnaNeural", "en-US-AndrewMultilingualNeural", "en-US-AriaNeural",
    "en-US-AvaMultilingualNeural", "en-US-BrianMultilingualNeural", "en-US-ChristopherNeural",
    "en-US-EmmaMultilingualNeural", "en-US-EricNeural", "en-US-GuyNeural", "en-US-JennyNeural",
    "en-US-MichelleNeural", "en-US-RogerNeural", "en-US-SteffanNeural",
    "es-ES-XimenaNeural", "es-MX-DaliaNeural", "es-MX-JorgeNeural", "es-ES-AlvaroNeural",
    "es-ES-ElviraNeural", "es-US-AlonsoNeural", "es-US-PalomaNeural",
}

REQUIRED_TOP_LEVEL_FIELDS = [
    "script_id", "category", "subcategory", "title", "format", "terms_used",
    "english_script", "spanish_script", "paired_segments",
]


def load_generator_module(generator_path: Path):
    if not generator_path.exists():
        raise FileNotFoundError(f"Generator script not found: {generator_path}")
    spec = importlib.util.spec_from_file_location("current_audio_generator_v7", generator_path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load generator module from {generator_path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def safe_script_id(data: dict[str, Any], fallback: str) -> str:
    return str(data.get("script_id") or fallback).strip() or fallback


def load_json_bytes(name: str, raw: bytes) -> tuple[dict[str, Any] | None, str | None]:
    try:
        payload = json.loads(raw.decode("utf-8-sig"))
        if not isinstance(payload, dict):
            return None, "Top-level JSON value is not an object."
        return payload, None
    except Exception as exc:
        return None, f"Could not parse {name}: {exc}"


def collect_voice_names(audio_profile: dict[str, Any]) -> list[str]:
    voices: list[str] = []
    defaults = audio_profile.get("default_voices", {}) or {}
    if isinstance(defaults, dict):
        voices.extend(str(v).strip() for v in defaults.values() if str(v).strip())

    assignments = audio_profile.get("voice_assignments", {}) or {}
    if isinstance(assignments, dict):
        for value in assignments.values():
            if isinstance(value, dict):
                voices.extend(str(v).strip() for key, v in value.items() if key in {"english", "spanish"} and str(v).strip())
    return sorted(set(voices))


def validate_json_payload(data: dict[str, Any]) -> tuple[list[str], list[str]]:
    warnings: list[str] = []
    errors: list[str] = []

    for field in REQUIRED_TOP_LEVEL_FIELDS:
        if field not in data:
            warnings.append(f"Missing recommended top-level field: {field}")

    terms = data.get("terms_used", []) or []
    if not isinstance(terms, list):
        warnings.append("terms_used should be a list.")

    paired = data.get("paired_segments", []) or []
    if not isinstance(paired, list) or not paired:
        warnings.append("No paired_segments found. The generator may fall back to paragraph splitting.")
        return warnings, errors

    for i, seg in enumerate(paired, start=1):
        if not isinstance(seg, dict):
            errors.append(f"Segment {i}: segment is not a JSON object.")
            continue
        if not str(seg.get("english", "")).strip():
            errors.append(f"Segment {i}: missing English text.")
        if not str(seg.get("spanish", "")).strip():
            errors.append(f"Segment {i}: missing Spanish text.")

        order = seg.get("playback_order")
        if order not in (["english", "spanish"], ["spanish", "english"]):
            warnings.append(f"Segment {i}: playback_order should be ['english','spanish'] or ['spanish','english'].")

        source_language = str(seg.get("source_language") or seg.get("original_language") or "").strip().lower()
        if source_language in {"english", "spanish"} and isinstance(order, list) and order:
            if source_language != str(order[0]).lower():
                warnings.append(
                    f"Segment {i}: source_language is {source_language!r}, but playback_order starts with {order[0]!r}."
                )
        elif not source_language:
            warnings.append(f"Segment {i}: missing source_language/original_language.")

        if not str(seg.get("speaker", "")).strip():
            warnings.append(f"Segment {i}: missing speaker label; speaker-specific voices may not apply.")

    audio_profile = data.get("audio_profile", {}) or {}
    if isinstance(audio_profile, dict) and audio_profile:
        voices = collect_voice_names(audio_profile)
        unapproved = [v for v in voices if v not in APPROVED_EDGE_VOICES]
        if unapproved:
            warnings.append("Unapproved or unconfirmed Edge TTS voice(s): " + ", ".join(unapproved))

        assignments = audio_profile.get("voice_assignments", {}) or {}
        if isinstance(assignments, dict):
            assigned = {str(k).strip() for k in assignments.keys()}
            speakers = {str(seg.get("speaker", "")).strip() for seg in paired if isinstance(seg, dict) and str(seg.get("speaker", "")).strip()}
            missing = sorted(speakers - assigned)
            if missing:
                warnings.append("Speakers without explicit audio_profile assignment: " + ", ".join(missing[:12]))

    return warnings, errors


def summarize_payload(data: dict[str, Any], filename: str) -> dict[str, Any]:
    paired = data.get("paired_segments", []) or []
    if not isinstance(paired, list):
        paired = []
    terms = data.get("terms_used", []) or []
    if not isinstance(terms, list):
        terms = []
    english_source = 0
    spanish_source = 0
    speakers: set[str] = set()
    segment_types: set[str] = set()
    for seg in paired:
        if not isinstance(seg, dict):
            continue
        src = str(seg.get("source_language") or seg.get("original_language") or "").strip().lower()
        if src == "english":
            english_source += 1
        elif src == "spanish":
            spanish_source += 1
        if str(seg.get("speaker", "")).strip():
            speakers.add(str(seg.get("speaker", "")).strip())
        if str(seg.get("segment_type", "")).strip():
            segment_types.add(str(seg.get("segment_type", "")).strip())
    return {
        "file": filename,
        "script_id": safe_script_id(data, Path(filename).stem),
        "title": data.get("title", ""),
        "format": data.get("format", ""),
        "segments": len(paired),
        "terms": len(terms),
        "english_source": english_source,
        "spanish_source": spanish_source,
        "speakers": ", ".join(sorted(speakers))[:180],
        "segment_types": ", ".join(sorted(segment_types))[:180],
    }


async def run_generator(generator_path: Path, json_files: list[Path], output_dir: Path, settings: dict[str, Any]) -> list[Path]:
    generator = load_generator_module(generator_path)

    generator.OUTPUT_DIR = output_dir
    generator.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    for key, value in settings.items():
        if hasattr(generator, key):
            setattr(generator, key, value)

    if not shutil.which("ffmpeg"):
        raise FileNotFoundError("ffmpeg was not found on PATH. Install ffmpeg before generating audio.")

    await generator.initialize_voice_validation()

    outputs: list[Path] = []
    for path in json_files:
        results = await generator.build_audio_for_script(path)
        outputs.extend(results or [])

    if getattr(generator, "GENERATE_COMBINED_AUDIO_ALL_SCRIPTS", False):
        all_script_mp3s = (
            generator.get_finished_script_rendition_mp3s(output_dir)
            if getattr(generator, "SPLIT_FLASHCARDS_AND_SCRIPT", False)
            else generator.get_all_finished_script_mp3s(output_dir)
        )
        if all_script_mp3s:
            combined_out = output_dir / "combined_audio_all_scripts.mp3"
            generator.concat_mp3s_with_gap(
                all_script_mp3s,
                combined_out,
                gap_seconds=getattr(generator, "COMBINED_SCRIPT_GAP", 3.0),
            )
            outputs.append(combined_out)
    return outputs





# -----------------------------
# Interactive practice audio helpers
# -----------------------------
DEFAULT_PRACTICE_VOICES = {
    "english": "en-US-JennyNeural",
    "spanish": "es-MX-DaliaNeural",
}


def sanitize_filename_part(value: str, max_len: int = 80) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(value).strip())
    cleaned = cleaned.strip("._") or "item"
    return cleaned[:max_len]


def short_hash(*parts: str) -> str:
    h = hashlib.sha1()
    for part in parts:
        h.update(str(part).encode("utf-8", errors="ignore"))
        h.update(b"\0")
    return h.hexdigest()[:12]


def get_rate_for_language(speed_mode: str, lang: str) -> str:
    presets = {
        "learning": {"english": "-17%", "spanish": "-17%"},
        "normal": {"english": "-4%", "spanish": "-4%"},
        "fast": {"english": "+6%", "spanish": "+6%"},
    }
    return presets.get(speed_mode, presets["learning"]).get(lang, "-17%")


def get_practice_voice(data: dict[str, Any], seg: dict[str, Any], lang: str, use_audio_profile: bool = True) -> str:
    """Resolve a voice for the practice player using the same intended fallback order as the generator."""
    lang = str(lang).strip().lower()
    fallback = DEFAULT_PRACTICE_VOICES.get(lang, "en-US-JennyNeural")
    if not use_audio_profile:
        return fallback

    def voice_from_entry(entry: Any) -> str:
        if isinstance(entry, str):
            return entry.strip()
        if isinstance(entry, dict):
            return str(entry.get(lang) or entry.get("voice") or entry.get("default") or "").strip()
        return ""

    # Segment-level override.
    if seg.get("voice"):
        voice = voice_from_entry(seg.get("voice"))
        if voice:
            return voice
    for key in ("voices", "voice_overrides"):
        entry = seg.get(key)
        if isinstance(entry, dict):
            voice = str(entry.get(lang) or "").strip()
            if voice:
                return voice

    profile = data.get("audio_profile") or {}
    if not isinstance(profile, dict):
        return fallback

    assignments = profile.get("voice_assignments") or {}
    speaker = str(seg.get("speaker") or "").strip().upper()
    if isinstance(assignments, dict) and speaker:
        # Match exact key first, then upper-case normalized key.
        entry = assignments.get(seg.get("speaker")) or assignments.get(speaker)
        if entry is None:
            for k, v in assignments.items():
                if str(k).strip().upper() == speaker:
                    entry = v
                    break
        voice = voice_from_entry(entry)
        if voice:
            return voice

    defaults = profile.get("default_voices") or {}
    if isinstance(defaults, dict):
        voice = str(defaults.get(lang) or "").strip()
        if voice:
            return voice
    return fallback


def get_practice_voice_for_chunk(data: dict[str, Any], chunk: dict[str, Any], lang: str, use_audio_profile: bool = True) -> str:
    """Resolve a practice voice for a simultaneous chunk.

    Single-chunk mode uses the segment's speaker voice. Full-passage mode may wrap
    several paired_segments; when they share a speaker, keep that assigned speaker
    voice instead of reverting to the global default. If there are multiple
    speakers, use the first segment's assigned voice as the practical fallback,
    because one continuous generated MP3 can only use one voice.
    """
    fallback = DEFAULT_PRACTICE_VOICES.get(str(lang).strip().lower(), "en-US-JennyNeural")
    chunk_segments = chunk.get("segments") or []
    if isinstance(chunk_segments, list) and chunk_segments:
        usable_segments = [seg for seg in chunk_segments if isinstance(seg, dict)]
        if usable_segments:
            return get_practice_voice(data, usable_segments[0], lang, use_audio_profile)

    # Paragraph/thought chunks keep a direct speaker field on the chunk.
    pseudo_segment = {"speaker": chunk.get("speaker", "")}
    return get_practice_voice(data, pseudo_segment, lang, use_audio_profile) if use_audio_profile else fallback


async def generate_practice_audio_file(text: str, voice: str, rate: str, out_path: Path, generator_path: Path | None = None) -> None:
    """Generate a small MP3 for the interactive practice player."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if out_path.exists() and out_path.stat().st_size > 0:
        return

    # Prefer the generator's hardened tts_to_mp3 when available, because it includes
    # voice validation and fallback behavior.
    if generator_path and generator_path.exists():
        try:
            generator = load_generator_module(generator_path)
            if hasattr(generator, "VERIFY_TTS_VOICES"):
                generator.VERIFY_TTS_VOICES = True
            if hasattr(generator, "initialize_voice_validation"):
                await generator.initialize_voice_validation()
            if hasattr(generator, "tts_to_mp3"):
                await generator.tts_to_mp3(text, voice, rate, out_path)
                return
        except Exception:
            # Fall back to direct edge_tts below so one helper mismatch does not break practice mode.
            pass

    import edge_tts
    communicate = edge_tts.Communicate(text=str(text), voice=voice, rate=rate)
    await communicate.save(str(out_path))


def get_practice_audio_path(
    cache_dir: Path,
    script_id: str,
    seg_index: Any,
    role: str,
    lang: str,
    text: str,
    voice: str,
    rate: str,
) -> Path:
    key = short_hash(script_id, str(seg_index), role, lang, text, voice, rate)
    filename = f"{sanitize_filename_part(script_id)}_seg_{int(seg_index):04d}_{role}_{lang}_{key}.mp3" if str(seg_index).isdigit() else f"{sanitize_filename_part(script_id)}_{role}_{lang}_{key}.mp3"
    return cache_dir / filename


def render_audio_player(
    path: Path,
    label: str,
    autoplay: bool = False,
    show_native_player: bool = True,
    nonce: str = "",
    playback_speed: float = 1.0,
) -> None:
    """Render an MP3 with a forced-refresh HTML audio element.

    Uses components.html instead of st.markdown so long base64 data URLs are
    rendered as an audio element rather than occasionally appearing as raw HTML
    text in Streamlit. This renderer is used by the Consecutive and Term Review
    tabs; the simultaneous tab has a speed-aware variant below.
    """
    if not path.exists():
        st.warning(f"Audio file not found: {path.name}")
        return

    audio_bytes = path.read_bytes()
    if not audio_bytes:
        st.warning(f"Audio file is empty: {path.name}")
        return

    content_key = short_hash(path.name, str(path.stat().st_mtime_ns), str(len(audio_bytes)), nonce, str(playback_speed))
    audio_b64 = base64.b64encode(audio_bytes).decode("ascii")
    element_id = f"practice_audio_{sanitize_filename_part(content_key)}"
    autoplay_attr = "autoplay" if autoplay else ""
    controls_attr = "controls" if show_native_player else ""
    audio_src = f"data:audio/mpeg;base64,{audio_b64}"
    safe_speed = max(0.5, min(2.0, float(playback_speed or 1.0)))

    component_html = f"""
    <!doctype html>
    <html>
      <head>
        <meta charset="utf-8">
        <style>
          body {{ margin: 0; padding: 0; background: transparent; }}
          audio {{ width: 100%; display: block; }}
          .caption {{
            margin-top: 4px;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
            font-size: 12px;
            color: #6b7280;
          }}
        </style>
      </head>
      <body>
        <audio id="{element_id}" {controls_attr} {autoplay_attr} preload="auto" src="{audio_src}"></audio>
        <div class="caption">{html.escape(label)} · {safe_speed:.2f}x · audio key: {content_key}</div>
        <script>
          const el = document.getElementById("{element_id}");
          if (el) {{
            el.playbackRate = {safe_speed};
            el.defaultPlaybackRate = {safe_speed};
          }}
        </script>
      </body>
    </html>
    """
    components.html(component_html, height=76, scrolling=False)


def render_simultaneous_audio_player(
    path: Path,
    label: str,
    autoplay: bool = False,
    show_native_player: bool = True,
    nonce: str = "",
    playback_speed: float = 1.0,
    loop: bool = False,
) -> None:
    """Render simultaneous-practice audio in an iframe-safe HTML component.

    Phase 2D.1 originally used st.markdown(unsafe_allow_html=True) for this
    custom audio tag. On some Streamlit/browser combinations, the long base64
    data URL can be escaped and shown as raw HTML text. Using components.html
    renders the audio element directly and avoids that raw HTML display issue
    while preserving playback speed and repeat-loop support.
    """
    if not path.exists():
        st.warning(f"Audio file not found: {path.name}")
        return

    audio_bytes = path.read_bytes()
    if not audio_bytes:
        st.warning(f"Audio file is empty: {path.name}")
        return

    content_key = short_hash(
        path.name,
        str(path.stat().st_mtime_ns),
        str(len(audio_bytes)),
        nonce,
        str(playback_speed),
        str(loop),
    )
    audio_b64 = base64.b64encode(audio_bytes).decode("ascii")
    element_id = f"sim_audio_{sanitize_filename_part(content_key)}"
    autoplay_attr = "autoplay" if autoplay else ""
    controls_attr = "controls" if show_native_player else ""
    loop_attr = "loop" if loop else ""
    safe_speed = max(0.5, min(2.0, float(playback_speed or 1.0)))
    audio_src = f"data:audio/mpeg;base64,{audio_b64}"

    component_html = f"""
    <!doctype html>
    <html>
      <head>
        <meta charset="utf-8">
        <style>
          body {{ margin: 0; padding: 0; background: transparent; }}
          audio {{ width: 100%; display: block; }}
          .caption {{
            margin-top: 4px;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
            font-size: 12px;
            color: #6b7280;
          }}
        </style>
      </head>
      <body>
        <audio id="{element_id}" {controls_attr} {autoplay_attr} {loop_attr} preload="auto" src="{audio_src}"></audio>
        <div class="caption">{html.escape(label)} · {safe_speed:.2f}x{(' · loop on' if loop else '')} · audio key: {content_key}</div>
        <script>
          const el = document.getElementById("{element_id}");
          if (el) {{
            el.playbackRate = {safe_speed};
            el.defaultPlaybackRate = {safe_speed};
          }}
        </script>
      </body>
    </html>
    """
    components.html(component_html, height=76, scrolling=False)



# Phase 2E UI helpers -------------------------------------------------------
def render_text_card(title: str, text: str, meta: str = "", tone: str = "source") -> None:
    """Render a clean text card for practice source/reference text."""
    tone_class = "reference" if tone == "reference" else "source"
    safe_title = html.escape(str(title or ""))
    safe_meta = html.escape(str(meta or ""))
    safe_text = html.escape(str(text or "")).replace("\n", "<br>")
    meta_html = f'<div class="practice-card-meta">{safe_meta}</div>' if safe_meta else ''
    st.markdown(
        f"""
        <div class="practice-card {tone_class}">
          <div class="practice-card-title">{safe_title}</div>
          {meta_html}
          <div class="practice-card-body">{safe_text}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_hidden_reference_card(title: str = "Reference hidden", prompt: str = "Reveal when ready.") -> None:
    st.markdown(
        f"""
        <div class="practice-card hidden-reference">
          <div class="practice-card-title">{html.escape(title)}</div>
          <div class="practice-card-body muted">{html.escape(prompt)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_side_by_side_text(
    source_text: str,
    target_text: str | None,
    source_lang: str,
    target_lang: str,
    speaker: str = "",
    segment_type: str = "",
    revealed: bool = False,
    source_visible: bool = True,
    context_label: str = "Reference interpretation",
) -> None:
    """Responsive source/reference workspace for Consecutive and Simultaneous practice."""
    left, right = st.columns(2, gap="large")
    meta_bits = [b for b in [f"Speaker: {speaker}" if speaker else "", f"Type: {segment_type}" if segment_type else ""] if b]
    meta = " · ".join(meta_bits)
    with left:
        if source_visible:
            render_text_card(f"Source — {str(source_lang).title()}", source_text, meta=meta, tone="source")
        else:
            render_hidden_reference_card("Source transcript hidden", "Audio-only mode is on. Interpret aloud while listening.")
    with right:
        if revealed and target_text:
            render_text_card(f"{context_label} — {str(target_lang).title()}", target_text, meta=meta, tone="reference")
        else:
            render_hidden_reference_card("Reference hidden", "Reveal the reference when you are ready to compare.")



def format_lang_label(lang: str) -> str:
    lang = str(lang or "").strip().lower()
    if lang == "english":
        return "EN"
    if lang == "spanish":
        return "ES"
    return lang[:2].upper() if lang else "—"


def render_compact_segment_meta(
    primary_label: str,
    source_lang: str,
    target_lang: str,
    word_count: int | None = None,
    speaker: str = "",
    segment_type: str = "",
    extra: str = "",
) -> None:
    """Render small, single-line metadata instead of large metric cards."""
    parts = [f"{format_lang_label(source_lang)} → {format_lang_label(target_lang)}"]
    if word_count is not None:
        parts.append(f"{int(word_count)} words")
    if speaker:
        parts.append(str(speaker))
    if segment_type:
        parts.append(str(segment_type))
    if extra:
        parts.append(str(extra))

    safe_primary = html.escape(str(primary_label or ""))
    safe_parts = [html.escape(str(part)) for part in parts if str(part).strip()]
    pills = "".join(f'<span class="compact-meta-pill">{part}</span>' for part in safe_parts)
    st.markdown(
        f"""
        <div class="compact-segment-meta">
          <span class="compact-meta-primary">{safe_primary}</span>
          <span class="compact-meta-pills">{pills}</span>
        </div>
        """,
        unsafe_allow_html=True,
    )



def compact_metadata_row(parts: list[str]) -> str:
    """Return compact metadata HTML for exam-mode metadata rows."""
    clean_parts = [str(part).strip() for part in (parts or []) if str(part).strip()]
    if not clean_parts:
        return ""
    safe_primary = html.escape(clean_parts[0])
    safe_rest = [html.escape(part) for part in clean_parts[1:]]
    pills = "".join(f'<span class="compact-meta-pill">{part}</span>' for part in safe_rest)
    return f"""
    <div class="compact-segment-meta">
      <span class="compact-meta-primary">{safe_primary}</span>
      <span class="compact-meta-pills">{pills}</span>
    </div>
    """


def render_source_reference_cards(
    source_title: str,
    source_text: str,
    target_title: str,
    target_text: str,
    source_meta: str = "",
    target_meta: str = "",
    source_visible: bool = True,
    target_visible: bool = False,
) -> None:
    """Compatibility wrapper for Phase 3 exam-mode side-by-side cards."""
    left, right = st.columns(2, gap="large")
    with left:
        if source_visible:
            render_text_card(source_title, source_text, meta=source_meta, tone="source")
        else:
            render_hidden_reference_card("Source transcript hidden", "Audio-only mode is on. Interpret aloud while listening.")
    with right:
        if target_visible and target_text:
            render_text_card(target_title, target_text, meta=target_meta, tone="reference")
        else:
            render_hidden_reference_card("Reference hidden", "Reference text is hidden during exam mode.")

def inject_phase2e_css() -> None:
    st.markdown(
        """
        <style>
          .practice-card {
            border: 1px solid rgba(49, 51, 63, 0.16);
            border-radius: 14px;
            padding: 1rem 1.05rem;
            margin: 0.35rem 0 0.85rem 0;
            background: rgba(255, 255, 255, 0.72);
            box-shadow: 0 1px 2px rgba(0,0,0,0.04);
          }
          .practice-card.reference { border-left: 5px solid rgba(46, 125, 50, 0.55); }
          .practice-card.source { border-left: 5px solid rgba(30, 136, 229, 0.55); }
          .practice-card.hidden-reference {
            border-style: dashed;
            background: rgba(250, 250, 250, 0.62);
          }
          .practice-card-title {
            font-weight: 700;
            font-size: 0.92rem;
            text-transform: uppercase;
            letter-spacing: 0.035em;
            margin-bottom: 0.35rem;
          }
          .practice-card-meta {
            font-size: 0.82rem;
            color: rgba(49, 51, 63, 0.68);
            margin-bottom: 0.55rem;
          }
          .practice-card-body { font-size: 1rem; line-height: 1.55; }
          .practice-card-body.muted {
            color: rgba(49, 51, 63, 0.58);
            font-style: italic;
          }
          .compact-segment-meta {
            display: flex;
            align-items: center;
            gap: 0.45rem;
            flex-wrap: wrap;
            margin: 0.15rem 0 0.75rem 0;
            padding: 0.35rem 0.45rem;
            border-radius: 999px;
            background: rgba(250, 250, 250, 0.72);
            border: 1px solid rgba(49, 51, 63, 0.10);
            font-size: 0.78rem;
            line-height: 1.2;
          }
          .compact-meta-primary {
            font-weight: 700;
            color: rgba(49, 51, 63, 0.78);
            margin-right: 0.1rem;
          }
          .compact-meta-pills {
            display: inline-flex;
            gap: 0.28rem;
            flex-wrap: wrap;
          }
          .compact-meta-pill {
            display: inline-block;
            padding: 0.14rem 0.42rem;
            border-radius: 999px;
            background: rgba(49, 51, 63, 0.055);
            color: rgba(49, 51, 63, 0.72);
            max-width: 18rem;
            overflow: hidden;
            text-overflow: ellipsis;
            white-space: nowrap;
          }
          @media (max-width: 760px) {
            .practice-card { padding: 0.9rem; }
            .practice-card-body { font-size: 0.96rem; }
          }
        </style>
        """,
        unsafe_allow_html=True,
    )

ERROR_TAGS = [
    "Omission",
    "Number/date/time error",
    "Wrong legal term",
    "Wrong speaker/actor",
    "Register/style issue",
    "Sequence/chronology error",
    "Source-language interference",
    "Memory/retention breakdown",
]


def get_segment_source_and_target(seg: dict[str, Any], direction: str = "Source language first") -> tuple[str, str]:
    """Return (source_lang, target_lang) for a paired segment."""
    if direction == "English → Spanish only":
        return "english", "spanish"
    if direction == "Spanish → English only":
        return "spanish", "english"

    order = seg.get("playback_order")
    if isinstance(order, list) and order and str(order[0]).lower() in {"english", "spanish"}:
        source = str(order[0]).lower()
    else:
        source = str(seg.get("source_language") or seg.get("original_language") or "english").lower()
        if source not in {"english", "spanish"}:
            source = "english"
    target = "spanish" if source == "english" else "english"
    return source, target


def estimate_dynamic_pause_seconds(text: str, ratio: float, min_pause: float, max_pause: float) -> float:
    """Estimate practice pause without creating audio. Audio generator uses actual MP3 duration."""
    words = max(1, len(str(text).split()))
    # Conservative speech estimate: about 145 words/minute.
    estimated_audio_seconds = words / 145 * 60
    return round(max(min_pause, min(max_pause, estimated_audio_seconds * ratio)), 1)


def filter_practice_segments(data: dict[str, Any], direction: str, segment_type_filter: list[str], speaker_filter: list[str]) -> list[dict[str, Any]]:
    paired = data.get("paired_segments", []) or []
    if not isinstance(paired, list):
        return []
    filtered: list[dict[str, Any]] = []
    for idx, seg in enumerate(paired, start=1):
        if not isinstance(seg, dict):
            continue
        source, target = get_segment_source_and_target(seg, direction)
        if not str(seg.get(source, "")).strip() or not str(seg.get(target, "")).strip():
            continue
        if segment_type_filter and str(seg.get("segment_type", "")).strip() not in segment_type_filter:
            continue
        if speaker_filter and str(seg.get("speaker", "")).strip() not in speaker_filter:
            continue
        copy = dict(seg)
        copy["_segment_index"] = idx
        copy["_source_language"] = source
        copy["_target_language"] = target
        filtered.append(copy)
    return filtered


PRACTICE_LOG_FIELDS = [
    "timestamp", "session_id", "script_id", "title", "practice_direction", "practice_set_size",
    "segment_index", "speaker", "segment_type", "source_language", "target_language",
    "score", "error_tags", "notes", "source_text", "target_text",
]


def practice_log_to_csv(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return ""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=PRACTICE_LOG_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        writer.writerow({k: row.get(k, "") for k in PRACTICE_LOG_FIELDS})
    return output.getvalue()


def normalize_log_row(row: dict[str, Any]) -> dict[str, Any]:
    normalized = {k: row.get(k, "") for k in PRACTICE_LOG_FIELDS}
    normalized["score"] = str(normalized.get("score", ""))
    return normalized


def load_persistent_practice_log(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            return [normalize_log_row(row) for row in csv.DictReader(f)]
    except Exception as exc:
        st.warning(f"Could not read persistent practice log: {exc}")
        return []


def append_persistent_practice_log(path: Path, row: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    file_exists = path.exists() and path.stat().st_size > 0
    with path.open("a", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=PRACTICE_LOG_FIELDS, extrasaction="ignore")
        if not file_exists:
            writer.writeheader()
        writer.writerow({k: row.get(k, "") for k in PRACTICE_LOG_FIELDS})


def overwrite_persistent_practice_log(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=PRACTICE_LOG_FIELDS, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in PRACTICE_LOG_FIELDS})


def parse_score(value: Any) -> int | None:
    try:
        score = int(value)
        return score if score in {0, 1, 2, 3} else None
    except Exception:
        return None


def rows_to_dataframe(rows: list[dict[str, Any]]) -> pd.DataFrame:
    if not rows:
        return pd.DataFrame(columns=PRACTICE_LOG_FIELDS)
    df = pd.DataFrame([normalize_log_row(r) for r in rows])
    for col in PRACTICE_LOG_FIELDS:
        if col not in df.columns:
            df[col] = ""
    df["score_numeric"] = pd.to_numeric(df["score"], errors="coerce")
    df["timestamp_dt"] = pd.to_datetime(df["timestamp"], errors="coerce")
    df["date"] = df["timestamp_dt"].dt.date.astype("string")
    df["direction"] = df["source_language"].str.title() + " → " + df["target_language"].str.title()
    return df


def split_error_tags(rows: list[dict[str, Any]]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for row in rows:
        raw = str(row.get("error_tags", ""))
        for tag in [t.strip() for t in raw.split(";") if t.strip()]:
            counts[tag] = counts.get(tag, 0) + 1
    return dict(sorted(counts.items(), key=lambda kv: (-kv[1], kv[0])))




# -----------------------------
# Phase 2C term review helpers
# -----------------------------
TERM_REVIEW_FIELDS = [
    "timestamp", "session_id", "script_id", "title", "source_file",
    "prompt_language", "answer_language", "english", "spanish", "alternatives",
    "score_label", "score_numeric", "notes",
]

TERM_SCORE_OPTIONS = {
    "Easy": 3,
    "Good": 2,
    "Hard": 1,
    "Missed": 0,
}


def normalize_term_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def extract_terms_from_payload(data: dict[str, Any], filename: str) -> list[dict[str, Any]]:
    terms = data.get("terms_used", []) or []
    if not isinstance(terms, list):
        return []
    script_id = safe_script_id(data, Path(filename).stem)
    title = str(data.get("title", ""))
    extracted: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()
    for i, term in enumerate(terms, start=1):
        if not isinstance(term, dict):
            continue
        english = normalize_term_text(term.get("english") or term.get("term") or term.get("source") or "")
        spanish = normalize_term_text(
            term.get("spanish_contextual_choice")
            or term.get("preferred_spanish")
            or term.get("spanish")
            or term.get("target")
            or ""
        )
        alternatives_raw = term.get("spanish_alternatives_from_sheet") or term.get("spanish_alternatives") or []
        if isinstance(alternatives_raw, list):
            alternatives = "; ".join(normalize_term_text(v) for v in alternatives_raw if normalize_term_text(v))
        else:
            alternatives = normalize_term_text(alternatives_raw)
        if not english or not spanish:
            continue
        key = (english.lower(), spanish.lower())
        if key in seen:
            continue
        seen.add(key)
        extracted.append({
            "term_id": f"{script_id}:{i}:{short_hash(english, spanish)}",
            "script_id": script_id,
            "title": title,
            "source_file": filename,
            "english": english,
            "spanish": spanish,
            "alternatives": alternatives,
        })
    return extracted


def load_term_payloads_from_files(files: list[tuple[str, bytes]]) -> tuple[list[tuple[str, dict[str, Any], bytes]], list[str]]:
    payloads: list[tuple[str, dict[str, Any], bytes]] = []
    errors: list[str] = []
    for name, raw in files:
        data, err = load_json_bytes(name, raw)
        if err:
            errors.append(err)
        elif data is not None:
            payloads.append((name, data, raw))
    return payloads, errors


def term_review_to_csv(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return ""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=TERM_REVIEW_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        writer.writerow({k: row.get(k, "") for k in TERM_REVIEW_FIELDS})
    return output.getvalue()


def normalize_term_review_row(row: dict[str, Any]) -> dict[str, Any]:
    normalized = {k: row.get(k, "") for k in TERM_REVIEW_FIELDS}
    normalized["score_numeric"] = str(normalized.get("score_numeric", ""))
    return normalized


def load_term_review_log(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            return [normalize_term_review_row(row) for row in csv.DictReader(f)]
    except Exception as exc:
        st.warning(f"Could not read term review log: {exc}")
        return []


def append_term_review_log(path: Path, row: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    file_exists = path.exists() and path.stat().st_size > 0
    with path.open("a", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=TERM_REVIEW_FIELDS, extrasaction="ignore")
        if not file_exists:
            writer.writeheader()
        writer.writerow({k: row.get(k, "") for k in TERM_REVIEW_FIELDS})


def term_rows_to_dataframe(rows: list[dict[str, Any]]) -> pd.DataFrame:
    if not rows:
        return pd.DataFrame(columns=TERM_REVIEW_FIELDS)
    df = pd.DataFrame([normalize_term_review_row(r) for r in rows])
    for col in TERM_REVIEW_FIELDS:
        if col not in df.columns:
            df[col] = ""
    df["score_numeric"] = pd.to_numeric(df["score_numeric"], errors="coerce")
    df["timestamp_dt"] = pd.to_datetime(df["timestamp"], errors="coerce")
    df["date"] = df["timestamp_dt"].dt.date.astype("string")
    df["term_key"] = df["english"].str.lower().str.strip() + "||" + df["spanish"].str.lower().str.strip()
    return df


def get_term_stats(term_history_rows: list[dict[str, Any]]) -> pd.DataFrame:
    df = term_rows_to_dataframe(term_history_rows)
    if df.empty:
        return pd.DataFrame(columns=["english", "spanish", "attempts", "average_score", "misses", "last_seen"])
    grouped = (
        df.dropna(subset=["score_numeric"])
        .groupby(["english", "spanish"], as_index=False)
        .agg(
            attempts=("score_numeric", "count"),
            average_score=("score_numeric", "mean"),
            misses=("score_numeric", lambda s: int((s <= 1).sum())),
            hard_or_missed=("score_numeric", lambda s: int((s <= 1).sum())),
            last_seen=("timestamp_dt", "max"),
        )
    )
    return grouped


def rank_terms_for_review(terms: list[dict[str, Any]], term_history_rows: list[dict[str, Any]], mode: str) -> list[dict[str, Any]]:
    if not terms:
        return []
    stats = get_term_stats(term_history_rows)
    stats_by_key: dict[tuple[str, str], dict[str, Any]] = {}
    if not stats.empty:
        for row in stats.to_dict("records"):
            stats_by_key[(str(row.get("english", "")).lower().strip(), str(row.get("spanish", "")).lower().strip())] = row

    ranked: list[dict[str, Any]] = []
    for pos, term in enumerate(terms):
        key = (str(term.get("english", "")).lower().strip(), str(term.get("spanish", "")).lower().strip())
        stat = stats_by_key.get(key, {})
        attempts = int(stat.get("attempts") or 0)
        avg = float(stat.get("average_score") if pd.notna(stat.get("average_score")) else 3.0) if stat else 3.0
        misses = int(stat.get("misses") or 0)
        priority = 0
        if attempts == 0:
            priority += 50
        priority += misses * 25
        priority += max(0, int(round((3 - avg) * 10)))
        priority -= min(attempts, 10)
        item = dict(term)
        item["attempts"] = attempts
        item["average_score"] = avg if attempts else None
        item["misses"] = misses
        item["priority"] = priority
        item["original_order"] = pos
        ranked.append(item)

    if mode == "New terms first":
        ranked.sort(key=lambda t: (t.get("attempts", 0) != 0, t.get("original_order", 0)))
    elif mode == "Hard/missed first":
        ranked.sort(key=lambda t: (-int(t.get("priority", 0)), t.get("original_order", 0)))
    elif mode == "Random":
        # stable random-ish shuffle for Streamlit reruns; Start New Session refreshes the seed
        seed = st.session_state.get("term_review_shuffle_seed", "default")
        ranked.sort(key=lambda t: short_hash(seed, t.get("term_id", "")))
    else:
        ranked.sort(key=lambda t: t.get("original_order", 0))
    return ranked



# -----------------------------
# Phase 2C.2 spaced repetition helpers
# -----------------------------
TERM_SRS_FIELDS = [
    "term_key", "english", "spanish", "script_ids", "titles", "source_files",
    "created_at", "last_reviewed_at", "next_review_at", "last_score_label",
    "last_score_numeric", "attempts", "misses", "lapses", "interval_days", "mastery_level",
]


def term_srs_key(term: dict[str, Any]) -> str:
    """Stable key across scripts for the same English/Spanish pair."""
    return short_hash(
        normalize_term_text(term.get("english", "")).lower(),
        normalize_term_text(term.get("spanish", "")).lower(),
    )


def today_date() -> datetime.date:
    return datetime.now().date()


def parse_date_value(value: Any) -> datetime.date | None:
    raw = str(value or "").strip()
    if not raw:
        return None
    try:
        return datetime.fromisoformat(raw.replace("Z", "+00:00")).date()
    except Exception:
        try:
            return datetime.strptime(raw[:10], "%Y-%m-%d").date()
        except Exception:
            return None


def normalize_srs_row(row: dict[str, Any]) -> dict[str, Any]:
    normalized = {k: row.get(k, "") for k in TERM_SRS_FIELDS}
    for numeric in ["last_score_numeric", "attempts", "misses", "lapses", "interval_days", "mastery_level"]:
        normalized[numeric] = str(normalized.get(numeric, "") or "0")
    return normalized


def load_term_srs_state(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            return [normalize_srs_row(row) for row in csv.DictReader(f)]
    except Exception as exc:
        st.warning(f"Could not read term SRS state: {exc}")
        return []


def save_term_srs_state(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=TERM_SRS_FIELDS, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in TERM_SRS_FIELDS})


def srs_rows_to_csv(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return ""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=TERM_SRS_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        writer.writerow({k: row.get(k, "") for k in TERM_SRS_FIELDS})
    return output.getvalue()


def term_srs_to_dataframe(rows: list[dict[str, Any]]) -> pd.DataFrame:
    if not rows:
        return pd.DataFrame(columns=TERM_SRS_FIELDS)
    df = pd.DataFrame([normalize_srs_row(r) for r in rows])
    for col in TERM_SRS_FIELDS:
        if col not in df.columns:
            df[col] = ""
    for col in ["last_score_numeric", "attempts", "misses", "lapses", "interval_days", "mastery_level"]:
        df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)
    df["next_review_date"] = pd.to_datetime(df["next_review_at"], errors="coerce").dt.date
    df["last_reviewed_date"] = pd.to_datetime(df["last_reviewed_at"], errors="coerce").dt.date
    return df


def merge_terms_into_srs_state(existing_rows: list[dict[str, Any]], terms: list[dict[str, Any]]) -> list[dict[str, Any]]:
    now = datetime.now().isoformat(timespec="seconds")
    by_key = {str(r.get("term_key", "")): normalize_srs_row(r) for r in existing_rows if str(r.get("term_key", "")).strip()}
    for term in terms:
        key = term_srs_key(term)
        if not key:
            continue
        script_id = str(term.get("script_id", ""))
        title = str(term.get("title", ""))
        source_file = str(term.get("source_file", ""))
        if key not in by_key:
            by_key[key] = {
                "term_key": key,
                "english": term.get("english", ""),
                "spanish": term.get("spanish", ""),
                "script_ids": script_id,
                "titles": title,
                "source_files": source_file,
                "created_at": now,
                "last_reviewed_at": "",
                "next_review_at": today_date().isoformat(),
                "last_score_label": "",
                "last_score_numeric": "",
                "attempts": "0",
                "misses": "0",
                "lapses": "0",
                "interval_days": "0",
                "mastery_level": "0",
            }
        else:
            row = by_key[key]
            for field, value in [("script_ids", script_id), ("titles", title), ("source_files", source_file)]:
                pieces = [x.strip() for x in str(row.get(field, "")).split(";") if x.strip()]
                if value and value not in pieces:
                    pieces.append(value)
                    row[field] = "; ".join(pieces)
    return list(by_key.values())


def calculate_next_srs_state(row: dict[str, Any], score_label: str) -> dict[str, Any]:
    score = int(TERM_SCORE_OPTIONS.get(score_label, 0))
    attempts = int(float(row.get("attempts", 0) or 0)) + 1
    misses = int(float(row.get("misses", 0) or 0)) + (1 if score <= 0 else 0)
    lapses = int(float(row.get("lapses", 0) or 0)) + (1 if score <= 1 and int(float(row.get("mastery_level", 0) or 0)) >= 2 else 0)
    old_interval = int(float(row.get("interval_days", 0) or 0))
    old_mastery = int(float(row.get("mastery_level", 0) or 0))
    if score <= 0:  # Missed
        interval = 0
        mastery = max(0, old_mastery - 1)
    elif score == 1:  # Hard
        interval = 1 if old_interval <= 1 else max(1, min(3, old_interval // 2))
        mastery = max(0, old_mastery)
    elif score == 2:  # Good
        interval = 3 if old_interval <= 0 else max(3, min(30, int(round(old_interval * 1.8))))
        mastery = min(5, old_mastery + 1)
    else:  # Easy
        interval = 7 if old_interval <= 0 else max(7, min(60, int(round(old_interval * 2.5))))
        mastery = min(5, old_mastery + 2)
    now = datetime.now()
    updated = dict(row)
    updated.update({
        "last_reviewed_at": now.isoformat(timespec="seconds"),
        "next_review_at": (now.date() + timedelta(days=interval)).isoformat(),
        "last_score_label": score_label,
        "last_score_numeric": str(score),
        "attempts": str(attempts),
        "misses": str(misses),
        "lapses": str(lapses),
        "interval_days": str(interval),
        "mastery_level": str(mastery),
    })
    return updated


def rank_terms_for_srs(terms: list[dict[str, Any]], srs_rows: list[dict[str, Any]], mode: str) -> list[dict[str, Any]]:
    if not terms:
        return []
    srs_by_key = {str(r.get("term_key", "")): normalize_srs_row(r) for r in srs_rows}
    today = today_date()
    ranked: list[dict[str, Any]] = []
    for pos, term in enumerate(terms):
        key = term_srs_key(term)
        srs = srs_by_key.get(key, {})
        attempts = int(float(srs.get("attempts", 0) or 0)) if srs else 0
        next_due = parse_date_value(srs.get("next_review_at")) if srs else today
        is_new = attempts == 0
        is_due = is_new or next_due is None or next_due <= today
        overdue_days = 0 if not next_due else max(0, (today - next_due).days)
        score = int(float(srs.get("last_score_numeric", 3) or 3)) if srs else 3
        interval = int(float(srs.get("interval_days", 0) or 0)) if srs else 0
        item = dict(term)
        item.update({
            "term_key": key,
            "srs_attempts": attempts,
            "srs_is_new": is_new,
            "srs_is_due": is_due,
            "srs_next_review_at": srs.get("next_review_at", today.isoformat()) if srs else today.isoformat(),
            "srs_last_score_label": srs.get("last_score_label", "") if srs else "",
            "srs_interval_days": interval,
            "srs_mastery_level": int(float(srs.get("mastery_level", 0) or 0)) if srs else 0,
            "srs_priority": (1000 if is_due else 0) + overdue_days * 10 + (20 if is_new else 0) + max(0, 3 - score) * 8 - min(interval, 30),
            "original_order": pos,
        })
        ranked.append(item)
    if mode == "Due today":
        ranked = [t for t in ranked if t.get("srs_is_due")]
        ranked.sort(key=lambda t: (-int(t.get("srs_priority", 0)), t.get("original_order", 0)))
    elif mode == "New terms":
        ranked = [t for t in ranked if t.get("srs_is_new")]
        ranked.sort(key=lambda t: t.get("original_order", 0))
    elif mode == "Review terms":
        ranked = [t for t in ranked if not t.get("srs_is_new") and t.get("srs_is_due")]
        ranked.sort(key=lambda t: (-int(t.get("srs_priority", 0)), t.get("original_order", 0)))
    elif mode == "All SRS priority":
        ranked.sort(key=lambda t: (-int(t.get("srs_priority", 0)), t.get("original_order", 0)))
    else:
        return rank_terms_for_review(terms, [], mode)
    return ranked




# -----------------------------
# Phase 3A.1 exam mode helpers
# -----------------------------
EXAM_LOG_FIELDS = [
    "timestamp", "session_id", "script_id", "title", "source_file", "exam_mode",
    "item_index", "item_label", "source_language", "target_language", "speaker",
    "segment_type", "source_text", "source_audio", "response_audio", "transcript_text", "transcript_file", "reference_text", "notes",
]


def exam_log_to_csv(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return ""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=EXAM_LOG_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        writer.writerow({k: row.get(k, "") for k in EXAM_LOG_FIELDS})
    return output.getvalue()


def save_audio_input(uploaded_audio: Any, out_path: Path) -> Path | None:
    """Save Streamlit st.audio_input output to disk."""
    if uploaded_audio is None:
        return None
    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        data = uploaded_audio.getvalue()
    except Exception:
        data = uploaded_audio.read()
    if not data:
        return None
    out_path.write_bytes(data)
    return out_path


def ffmpeg_convert_to_mp3(src_path: Path, dst_path: Path) -> Path:
    dst_path.parent.mkdir(parents=True, exist_ok=True)
    if dst_path.exists() and dst_path.stat().st_size > 0:
        return dst_path
    if not shutil.which("ffmpeg"):
        raise FileNotFoundError("ffmpeg was not found on PATH; cannot combine recordings.")
    subprocess.run([
        "ffmpeg", "-y", "-hide_banner", "-loglevel", "error",
        "-i", str(src_path), "-acodec", "libmp3lame", "-q:a", "4", str(dst_path)
    ], check=True)
    return dst_path


def get_audio_duration_seconds(audio_path: Path) -> float | None:
    """Return audio duration in seconds using ffprobe when available."""
    if not audio_path or not audio_path.exists() or not audio_path.is_file():
        return None
    if not shutil.which("ffprobe"):
        return None
    try:
        result = subprocess.run([
            "ffprobe", "-v", "error", "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1", str(audio_path)
        ], capture_output=True, text=True, check=True)
        value = float((result.stdout or "").strip())
        if value > 0:
            return value
    except Exception:
        return None
    return None


# -----------------------------
# Spoken-rate / WPM helpers
# -----------------------------
def count_spoken_words(text: str) -> int:
    """Count likely spoken words in English/Spanish text for WPM estimates."""
    return len(re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9]+(?:[-'][A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9]+)?", str(text or "")))


def calculate_effective_wpm(text: str, audio_path: Path | None, playback_speed: float = 1.0) -> dict[str, Any]:
    """Calculate effective WPM using actual generated audio duration and browser playback speed.

    TTS voice speed/rate is already reflected in the generated MP3 duration.
    Browser playback speed changes the heard duration, so effective_duration = raw_duration / playback_speed.
    """
    words = count_spoken_words(text)
    raw_duration = get_audio_duration_seconds(audio_path) if audio_path else None
    speed = max(0.1, float(playback_speed or 1.0))
    if not raw_duration or raw_duration <= 0 or words <= 0:
        return {
            "words": words,
            "raw_duration_seconds": raw_duration,
            "effective_duration_seconds": None,
            "playback_speed": speed,
            "wpm": None,
        }
    effective_duration = raw_duration / speed
    wpm = (words / effective_duration) * 60.0 if effective_duration > 0 else None
    return {
        "words": words,
        "raw_duration_seconds": raw_duration,
        "effective_duration_seconds": effective_duration,
        "playback_speed": speed,
        "wpm": wpm,
    }


def format_seconds_compact(seconds: float | None) -> str:
    if seconds is None:
        return "—"
    seconds = max(0.0, float(seconds))
    minutes = int(seconds // 60)
    remainder = int(round(seconds % 60))
    if minutes:
        return f"{minutes}:{remainder:02d}"
    return f"{remainder}s"


def render_wpm_estimate(label: str, text: str, audio_path: Path | None, playback_speed: float = 1.0, target_wpm: int = 140) -> None:
    """Render a compact WPM estimate for generated/spoken audio."""
    info = calculate_effective_wpm(text, audio_path, playback_speed)
    words = info.get("words", 0)
    wpm = info.get("wpm")
    if wpm is None:
        st.caption(f"{label} WPM: unavailable until audio duration can be measured with ffprobe. Words: {words}.")
        return
    diff = float(wpm) - float(target_wpm)
    if abs(diff) < 3:
        status = f"near {target_wpm} WPM exam target"
    elif diff > 0:
        status = f"{diff:.0f} WPM above {target_wpm}"
    else:
        status = f"{abs(diff):.0f} WPM below {target_wpm}"
    st.caption(
        f"{label} estimated spoken rate: **{wpm:.0f} WPM** "
        f"({words} words / {format_seconds_compact(info.get('effective_duration_seconds'))} heard time; "
        f"audio {format_seconds_compact(info.get('raw_duration_seconds'))} @ {float(info.get('playback_speed') or 1.0):.2f}x) — {status}."
    )


def wpm_fields_for_text_audio(text: str, audio_path: Path | str | None, playback_speed: float = 1.0) -> dict[str, Any]:
    """Return row-safe WPM metadata for AI summaries/reports."""
    path_obj = Path(audio_path) if audio_path else None
    info = calculate_effective_wpm(text, path_obj, playback_speed)
    if info.get("wpm") is None:
        return {
            "source_word_count": info.get("words", len(str(text or "").split())),
            "source_duration_seconds": "",
            "heard_duration_seconds": "",
            "source_wpm": "",
            "effective_wpm": "",
            "wpm_target": 140,
            "wpm_target_delta": "",
        }
    wpm = float(info.get("wpm"))
    return {
        "source_word_count": int(info.get("words", 0) or 0),
        "source_duration_seconds": round(float(info.get("raw_duration_seconds", 0) or 0), 2),
        "heard_duration_seconds": round(float(info.get("effective_duration_seconds", info.get("heard_duration_seconds", 0)) or 0), 2),
        "source_wpm": round(wpm, 1),
        "effective_wpm": round(wpm, 1),
        "wpm_target": 140,
        "wpm_target_delta": round(wpm - 140.0, 1),
    }


def ffmpeg_concat_mp3s(parts: list[Path], dst_path: Path) -> Path | None:
    valid = [p for p in parts if p and p.exists() and p.stat().st_size > 0]
    if not valid:
        return None
    if not shutil.which("ffmpeg"):
        raise FileNotFoundError("ffmpeg was not found on PATH; cannot combine audio.")
    dst_path.parent.mkdir(parents=True, exist_ok=True)
    list_file = dst_path.with_suffix(".concat.txt")
    list_file.write_text("".join(f"file '{str(p.resolve()).replace("'", "'\\''")}'\n" for p in valid), encoding="utf-8")
    subprocess.run([
        "ffmpeg", "-y", "-hide_banner", "-loglevel", "error",
        "-f", "concat", "-safe", "0", "-i", str(list_file), "-c", "copy", str(dst_path)
    ], check=True)
    with contextlib.suppress(Exception):
        list_file.unlink()
    return dst_path


def make_zip_from_paths(paths: list[Path], zip_path: Path) -> Path | None:
    valid = [p for p in paths if p and p.exists() and p.is_file() and p.stat().st_size > 0]
    if not valid:
        return None
    zip_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for p in valid:
            zf.write(p, arcname=p.name)
    return zip_path

# -----------------------------
# Phase 3B.1 exam transcription helpers
# -----------------------------
TRANSCRIPT_CSV_FIELDS = [
    "timestamp", "session_id", "script_id", "exam_mode", "item_index", "item_label",
    "source_language", "target_language", "speaker", "segment_type", "source_text", "reference_text", "response_audio",
    "transcript_file", "transcript_text",
]


def get_openai_api_key(manual_key: str | None = None) -> str:
    """Return an OpenAI API key from the app field or environment."""
    key = (manual_key or "").strip()
    if key:
        return key
    return os.environ.get("OPENAI_API_KEY", "").strip()


def language_hint_to_api_value(hint: str) -> str | None:
    mapping = {"Auto": None, "English": "en", "Spanish": "es"}
    return mapping.get(hint, None)


def response_to_transcript_text(response: Any) -> str:
    """Normalize OpenAI transcription responses across SDK/model variants."""
    if isinstance(response, str):
        return response.strip()
    text = getattr(response, "text", None)
    if text:
        return str(text).strip()
    if isinstance(response, dict):
        return str(response.get("text", "")).strip()
    with contextlib.suppress(Exception):
        dumped = response.model_dump()
        return str(dumped.get("text", "")).strip()
    return str(response).strip()


def prepare_audio_for_transcription(src_path: Path, work_dir: Path) -> Path:
    """Compress/convert audio when needed so the transcription request stays under API limits."""
    if not src_path.exists() or not src_path.is_file():
        raise FileNotFoundError(f"Recording not found: {src_path}")
    if src_path.stat().st_size <= OPENAI_TRANSCRIBE_LIMIT_BYTES and src_path.suffix.lower() in {".mp3", ".mp4", ".mpeg", ".mpga", ".m4a", ".wav", ".webm"}:
        return src_path
    if not shutil.which("ffmpeg"):
        raise FileNotFoundError("ffmpeg was not found on PATH; cannot compress this recording for transcription.")
    work_dir.mkdir(parents=True, exist_ok=True)
    dst = work_dir / f"{src_path.stem}_transcribe.mp3"
    subprocess.run([
        "ffmpeg", "-y", "-hide_banner", "-loglevel", "error",
        "-i", str(src_path), "-vn", "-ac", "1", "-ar", "16000",
        "-codec:a", "libmp3lame", "-b:a", "64k", str(dst)
    ], check=True)
    if dst.stat().st_size > OPENAI_TRANSCRIBE_LIMIT_BYTES:
        raise ValueError("Compressed recording is still over the 25 MB transcription limit. Split it into smaller parts first.")
    return dst


def transcribe_audio_openai(audio_path: Path, model: str, api_key: str, language_hint: str | None = None) -> str:
    """Transcribe one audio file with OpenAI speech-to-text."""
    if not api_key:
        raise ValueError("Missing OpenAI API key. Enter one in the app or set OPENAI_API_KEY in your environment.")
    try:
        from openai import OpenAI
    except Exception as exc:
        raise RuntimeError("The openai package is not installed. Run: python -m pip install openai") from exc
    client = OpenAI(api_key=api_key)
    kwargs: dict[str, Any] = {"model": model}
    if language_hint:
        kwargs["language"] = language_hint
    with audio_path.open("rb") as f:
        response = client.audio.transcriptions.create(file=f, **kwargs)
    return response_to_transcript_text(response)


def openai_response_to_dict(response: Any) -> dict[str, Any]:
    """Normalize OpenAI SDK response objects to a plain dictionary."""
    if isinstance(response, dict):
        return response
    with contextlib.suppress(Exception):
        return response.model_dump()
    with contextlib.suppress(Exception):
        return dict(response)
    return {"text": response_to_transcript_text(response)}


def transcribe_audio_openai_with_timestamps(audio_path: Path, api_key: str, language_hint: str | None = None) -> dict[str, Any]:
    """Transcribe one audio file with whisper-1 verbose_json word/segment timestamps."""
    if not api_key:
        raise ValueError("Missing OpenAI API key. Enter one in the app or set OPENAI_API_KEY in your environment.")
    try:
        from openai import OpenAI
    except Exception as exc:
        raise RuntimeError("The openai package is not installed. Run: python -m pip install openai") from exc
    client = OpenAI(api_key=api_key)
    kwargs: dict[str, Any] = {
        "model": "whisper-1",
        "response_format": "verbose_json",
        "timestamp_granularities": ["word", "segment"],
    }
    if language_hint:
        kwargs["language"] = language_hint
    with audio_path.open("rb") as f:
        response = client.audio.transcriptions.create(file=f, **kwargs)
    data = openai_response_to_dict(response)
    # Some SDK/API combinations may not return both granularities. Keep graceful fallbacks.
    data.setdefault("text", response_to_transcript_text(response))
    data.setdefault("words", [])
    data.setdefault("segments", [])
    return data


def normalize_timestamp_words(timestamp_data: dict[str, Any]) -> list[dict[str, Any]]:
    """Return transcript words as [{word,start,end}] with float timestamps."""
    words = timestamp_data.get("words") or []
    normalized: list[dict[str, Any]] = []
    if isinstance(words, list):
        for item in words:
            if not isinstance(item, dict):
                item = openai_response_to_dict(item)
            word = str(item.get("word", item.get("text", ""))).strip()
            if not word:
                continue
            try:
                start = float(item.get("start", 0) or 0)
                end = float(item.get("end", start) or start)
            except Exception:
                continue
            normalized.append({"word": word, "start": start, "end": max(end, start)})
    if normalized:
        return normalized

    # Fallback: approximate words from the full text when word timestamps are absent.
    text = str(timestamp_data.get("text", "") or "").strip()
    parts = text.split()
    if not parts:
        return []
    segments = timestamp_data.get("segments") or []
    total_duration = None
    if isinstance(segments, list) and segments:
        last = segments[-1]
        if not isinstance(last, dict):
            last = openai_response_to_dict(last)
        with contextlib.suppress(Exception):
            total_duration = float(last.get("end", 0) or 0)
    total_duration = total_duration or max(len(parts) * 0.45, 1.0)
    step = total_duration / max(len(parts), 1)
    return [{"word": w, "start": i * step, "end": (i + 1) * step} for i, w in enumerate(parts)]


def words_to_text(words: list[dict[str, Any]]) -> str:
    """Join timestamped words into readable transcript text."""
    return " ".join(str(w.get("word", "")).strip() for w in words if str(w.get("word", "")).strip()).strip()


def format_timecode(seconds: float | int | None) -> str:
    """Format seconds as M:SS."""
    try:
        total = max(0, int(round(float(seconds))))
    except Exception:
        total = 0
    minutes, sec = divmod(total, 60)
    return f"{minutes}:{sec:02d}"


def split_text_into_chunks_from_row(row: dict[str, Any]) -> list[dict[str, Any]]:
    """Build source/reference chunks from saved exam row metadata."""
    # Simultaneous rows can preserve the original paired_segments as JSON for chunk alignment.
    raw_segments = str(row.get("alignment_segments_json", "") or "").strip()
    if raw_segments:
        with contextlib.suppress(Exception):
            segments = json.loads(raw_segments)
            if isinstance(segments, list) and segments:
                out = []
                for pos, seg in enumerate(segments, start=1):
                    if not isinstance(seg, dict):
                        continue
                    source_lang = str(seg.get("_source_language") or row.get("source_language") or "")
                    target_lang = str(seg.get("_target_language") or row.get("target_language") or "")
                    source_text = str(seg.get("_source_text") or seg.get(source_lang, "") or "").strip()
                    target_text = str(seg.get("_target_text") or seg.get(target_lang, "") or "").strip()
                    if source_text or target_text:
                        out.append({
                            "chunk_number": pos,
                            "item_label": f"Chunk {pos} / JSON segment {seg.get('_segment_index', pos)}",
                            "source_language": source_lang,
                            "target_language": target_lang,
                            "speaker": str(seg.get("speaker", row.get("speaker", "")) or ""),
                            "segment_type": str(seg.get("segment_type", row.get("segment_type", "")) or ""),
                            "source_text": source_text,
                            "reference_text": target_text,
                        })
                if out:
                    return out

    source_text = str(row.get("source_text", "") or "").strip()
    reference_text = str(row.get("reference_text", "") or "").strip()
    source_parts = [p.strip() for p in re.split(r"\n\s*\n", source_text) if p.strip()] or ([source_text] if source_text else [])
    reference_parts = [p.strip() for p in re.split(r"\n\s*\n", reference_text) if p.strip()] or ([reference_text] if reference_text else [])
    count = max(len(source_parts), len(reference_parts), 1)
    chunks = []
    for i in range(count):
        chunks.append({
            "chunk_number": i + 1,
            "item_label": str(row.get("item_label") or f"Item {row.get('item_index', '')}"),
            "source_language": str(row.get("source_language", "")),
            "target_language": str(row.get("target_language", "")),
            "speaker": str(row.get("speaker", "")),
            "segment_type": str(row.get("segment_type", "")),
            "source_text": source_parts[i] if i < len(source_parts) else "",
            "reference_text": reference_parts[i] if i < len(reference_parts) else "",
        })
    return chunks


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except Exception:
        return default


def get_exam_row_playback_speed(row: dict[str, Any]) -> float:
    """Return the browser playback speed used during the exam recording."""
    speed = _safe_float(row.get("playback_speed", 1.0), 1.0)
    return max(0.25, speed)


def estimate_source_chunk_windows(row: dict[str, Any], source_duration: float | None) -> list[dict[str, Any]]:
    """Estimate source timing windows for a saved exam row's source chunks by word share.

    Kept as a fallback for older exam rows or when exact chunk timing cannot be
    prepared. Returned source_start/source_end values are in heard-time seconds,
    so browser playback speed is already accounted for.
    """
    chunks = split_text_into_chunks_from_row(row)
    if not chunks:
        return []
    word_counts = [max(1, len(str(c.get("source_text", "")).split())) for c in chunks]
    total_words = max(sum(word_counts), 1)
    raw_duration = source_duration or max(total_words * 0.42, len(chunks) * 2.0)
    playback_speed = get_exam_row_playback_speed(row)
    duration = raw_duration / playback_speed
    cursor = 0.0
    windows = []
    for chunk, wc in zip(chunks, word_counts):
        seg_duration = duration * (wc / total_words)
        start = cursor
        end = duration if len(windows) == len(chunks) - 1 else cursor + seg_duration
        cursor = end
        item = dict(chunk)
        item.update({
            "source_start": start,
            "source_end": end,
            "source_raw_start": start * playback_speed,
            "source_raw_end": end * playback_speed,
            "source_word_count": wc,
            "source_duration_seconds": round(seg_duration, 3),
            "heard_duration_seconds": round(seg_duration, 3),
            "source_wpm": round((wc / seg_duration) * 60.0, 1) if seg_duration > 0 else "",
            "effective_wpm": round((wc / seg_duration) * 60.0, 1) if seg_duration > 0 else "",
            "wpm_target": 140,
            "wpm_target_delta": round(((wc / seg_duration) * 60.0) - 140.0, 1) if seg_duration > 0 else "",
            "timing_method": "word-share estimate",
            "playback_speed": playback_speed,
        })
        windows.append(item)
    return windows


def build_exact_source_chunk_windows(
    row: dict[str, Any],
    source_duration: float | None,
    timing_cache_dir: Path,
    generator_path: Path | None = None,
) -> tuple[list[dict[str, Any]], Path | None]:
    """Build a per-chunk source timing map from actual generated chunk audio durations.

    This is more accurate than word-count allocation for simultaneous full-passage
    exams. It generates/reuses one small source MP3 per JSON chunk using the same
    exam voice/rate metadata saved with the recording, measures each chunk, scales
    the total to the actual full-passage source MP3 duration when available, and
    finally converts raw MP3 time to heard time by dividing by playback speed.
    """
    chunks = split_text_into_chunks_from_row(row)
    if not chunks:
        return [], None

    playback_speed = get_exam_row_playback_speed(row)
    raw_source_duration = source_duration or 0.0
    voice = str(row.get("tts_voice") or "").strip() or DEFAULT_PRACTICE_VOICES.get(str(row.get("source_language", "english")).lower(), "en-US-JennyNeural")
    rate = str(row.get("tts_rate") or "").strip() or "+0%"
    script_id = str(row.get("script_id") or "exam")
    session_id = str(row.get("session_id") or "session")
    item_index = str(row.get("item_index") or "item")

    timing_cache_dir.mkdir(parents=True, exist_ok=True)
    raw_durations: list[float] = []
    chunk_audio_files: list[str] = []

    # Consecutive rows are already one source segment, so the exact source file
    # duration is the best timing map. For simultaneous rows, use per-chunk TTS.
    if len(chunks) == 1:
        raw_durations = [raw_source_duration or max(len(str(chunks[0].get("source_text", "")).split()) * 0.42, 1.0)]
        chunk_audio_files = [str(row.get("source_audio", ""))]
    else:
        for chunk in chunks:
            source_text = str(chunk.get("source_text", "") or "").strip()
            if not source_text:
                raw_durations.append(0.0)
                chunk_audio_files.append("")
                continue
            chunk_no = chunk.get("chunk_number", len(raw_durations) + 1)
            chunk_audio = get_exam_audio_path(
                timing_cache_dir,
                script_id,
                f"{item_index}_{chunk_no}",
                "timing_source_chunk",
                str(row.get("source_language") or chunk.get("source_language") or "source"),
                source_text,
                voice,
                rate,
            )
            try:
                asyncio.run(generate_practice_audio_file(source_text, voice, rate, chunk_audio, generator_path or DEFAULT_GENERATOR_PATH))
            except RuntimeError:
                # If an event loop is already active, fall back to the older estimate path.
                return estimate_source_chunk_windows(row, source_duration), None
            except Exception:
                return estimate_source_chunk_windows(row, source_duration), None
            duration = get_audio_duration_seconds(chunk_audio)
            if not duration or duration <= 0:
                duration = max(len(source_text.split()) * 0.42, 1.0)
            raw_durations.append(float(duration))
            chunk_audio_files.append(str(chunk_audio))

    measured_total = sum(d for d in raw_durations if d > 0)
    if measured_total <= 0:
        return estimate_source_chunk_windows(row, source_duration), None

    # Edge TTS can insert slightly different leading/trailing silence when chunks
    # are generated separately. Scale the chunk map to the actual full-passage
    # source MP3 if it was measured successfully.
    scale = (raw_source_duration / measured_total) if raw_source_duration and raw_source_duration > 0 else 1.0

    windows: list[dict[str, Any]] = []
    raw_cursor = 0.0
    for chunk, raw_duration, chunk_audio_file in zip(chunks, raw_durations, chunk_audio_files):
        adjusted_raw_duration = max(0.0, raw_duration * scale)
        raw_start = raw_cursor
        raw_end = raw_cursor + adjusted_raw_duration
        raw_cursor = raw_end
        item = dict(chunk)
        item.update({
            "source_raw_start": raw_start,
            "source_raw_end": raw_end,
            "source_start": raw_start / playback_speed,
            "source_end": raw_end / playback_speed,
            "source_word_count": max(1, len(str(chunk.get("source_text", "")).split())),
            "source_duration_seconds": round(max(0.0, (raw_end / playback_speed) - (raw_start / playback_speed)), 3),
            "heard_duration_seconds": round(max(0.0, (raw_end / playback_speed) - (raw_start / playback_speed)), 3),
            "source_wpm": round((max(1, len(str(chunk.get("source_text", "")).split())) / max(0.001, ((raw_end / playback_speed) - (raw_start / playback_speed)))) * 60.0, 1),
            "effective_wpm": round((max(1, len(str(chunk.get("source_text", "")).split())) / max(0.001, ((raw_end / playback_speed) - (raw_start / playback_speed)))) * 60.0, 1),
            "wpm_target": 140,
            "wpm_target_delta": round(((max(1, len(str(chunk.get("source_text", "")).split())) / max(0.001, ((raw_end / playback_speed) - (raw_start / playback_speed)))) * 60.0) - 140.0, 1),
            "timing_method": "exact chunk timing map",
            "playback_speed": playback_speed,
            "tts_voice": voice,
            "tts_rate": rate,
            "timing_chunk_audio": chunk_audio_file,
        })
        windows.append(item)

    timing_map_path = timing_cache_dir / f"{sanitize_filename_part(script_id)}_{sanitize_filename_part(session_id)}_{sanitize_filename_part(item_index)}_source_timing_map.json"
    timing_map_payload = {
        "script_id": script_id,
        "session_id": session_id,
        "item_index": item_index,
        "timing_method": "exact chunk timing map",
        "tts_voice": voice,
        "tts_rate": rate,
        "playback_speed": playback_speed,
        "raw_source_duration_seconds": raw_source_duration,
        "measured_chunk_total_seconds": measured_total,
        "scale_factor_to_full_source": scale,
        "chunks": windows,
    }
    timing_map_path.write_text(json.dumps(timing_map_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return windows, timing_map_path


def align_timestamp_words_to_source_windows(words: list[dict[str, Any]], windows: list[dict[str, Any]], lag_seconds: float = 3.0, padding_seconds: float = 1.5) -> list[dict[str, Any]]:
    """Place timestamped response words into estimated source chunk windows using lag/padding."""
    aligned: list[dict[str, Any]] = []
    for window in windows:
        source_start = float(window.get("source_start", 0) or 0)
        source_end = float(window.get("source_end", source_start) or source_start)
        response_start = max(0.0, source_start + float(lag_seconds) - float(padding_seconds))
        response_end = max(response_start, source_end + float(lag_seconds) + float(padding_seconds))
        selected = []
        for word in words:
            mid = (float(word.get("start", 0) or 0) + float(word.get("end", 0) or 0)) / 2
            if response_start <= mid <= response_end:
                selected.append(word)
        text = words_to_text(selected)
        item = dict(window)
        heard_duration = max(0.001, source_end - source_start)
        source_wc = int(window.get("source_word_count", 0) or max(1, len(str(window.get("source_text", "")).split())))
        source_wpm_val = round((source_wc / heard_duration) * 60.0, 1)
        item.update({
            "response_window_start": response_start,
            "response_window_end": response_end,
            "your_transcript_aligned": text,
            "word_count_aligned": len(selected),
            "possible_gap": "yes" if not text.strip() else "",
            "source_word_count": source_wc,
            "source_duration_seconds": round(heard_duration, 3),
            "heard_duration_seconds": round(heard_duration, 3),
            "source_wpm": window.get("source_wpm", source_wpm_val),
            "effective_wpm": window.get("effective_wpm", source_wpm_val),
            "wpm_target": 140,
            "wpm_target_delta": window.get("wpm_target_delta", round(source_wpm_val - 140.0, 1)),
        })
        aligned.append(item)
    return aligned


TIMESTAMP_ALIGNMENT_CSV_FIELDS = [
    "script_id", "session_id", "exam_mode", "item_index", "chunk_number", "item_label",
    "source_language", "target_language", "speaker", "segment_type",
    "source_start", "source_end", "source_raw_start", "source_raw_end", "response_window_start", "response_window_end",
    "timing_method", "playback_speed", "tts_voice", "tts_rate",
    "source_text", "reference_text", "your_transcript_aligned", "word_count_aligned", "possible_gap",
    "source_word_count", "source_duration_seconds", "heard_duration_seconds", "source_wpm", "effective_wpm", "wpm_target", "wpm_target_delta",
]


def timestamp_alignment_rows_to_csv(rows: list[dict[str, Any]]) -> str:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=TIMESTAMP_ALIGNMENT_CSV_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        writer.writerow({k: row.get(k, "") for k in TIMESTAMP_ALIGNMENT_CSV_FIELDS})
    return output.getvalue()


def timestamp_alignment_rows_to_txt(rows: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for row in rows:
        parts.append(
            f"{row.get('item_label', 'Item')} — chunk {row.get('chunk_number', '')}\n"
            f"Source window: {format_timecode(row.get('source_start'))}–{format_timecode(row.get('source_end'))} | "
            f"Response window: {format_timecode(row.get('response_window_start'))}–{format_timecode(row.get('response_window_end'))} | "
            f"Timing: {row.get('timing_method', '')} @ {row.get('playback_speed', '1.0')}x\n"
            f"Direction: {row.get('source_language', '')} → {row.get('target_language', '')} | "
            f"Speaker: {row.get('speaker', '')} | Type: {row.get('segment_type', '')}\n\n"
            f"SOURCE:\n{str(row.get('source_text', '')).strip()}\n\n"
            f"REFERENCE:\n{str(row.get('reference_text', '')).strip()}\n\n"
            f"YOUR TIMESTAMP-ALIGNED TRANSCRIPT:\n{str(row.get('your_transcript_aligned', '')).strip() or '[No matching speech detected in this window]'}\n"
        )
    return ("\n" + "-" * 80 + "\n\n").join(parts).strip() + "\n"


def build_timestamp_alignment_docx(rows: list[dict[str, Any]], out_path: Path) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise RuntimeError("python-docx is not installed. Run: python -m pip install python-docx") from exc
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    with contextlib.suppress(Exception):
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(8)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Timestamp-Anchored Exam Alignment")
    run.bold = True
    run.font.size = Pt(15)
    if rows:
        first = rows[0]
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Script: {first.get('script_id', '')} | Mode: {first.get('exam_mode', '')} | Session: {first.get('session_id', '')}").italic = True
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(["Time", "Source", "Reference", "Your timestamp-aligned transcript"]):
        hdr[i].text = label
        with contextlib.suppress(Exception):
            hdr[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = f"{format_timecode(row.get('source_start'))}–{format_timecode(row.get('source_end'))}\nResp: {format_timecode(row.get('response_window_start'))}–{format_timecode(row.get('response_window_end'))}\n{row.get('timing_method', '')}\n{row.get('playback_speed', '1.0')}x"
        cells[1].text = str(row.get("source_text", "") or "")
        cells[2].text = str(row.get("reference_text", "") or "")
        transcript = str(row.get("your_transcript_aligned", "") or "").strip()
        cells[3].text = transcript or "[No matching speech detected in this window]"
    doc.save(out_path)
    return out_path


def build_timestamp_anchored_alignment_outputs(
    exam_rows: list[dict[str, Any]],
    script_id: str,
    session_id: str,
    out_dir: Path,
    api_key: str,
    language_hint: str | None,
    lag_seconds: float,
    padding_seconds: float,
    use_exact_timing_map: bool = True,
) -> tuple[list[dict[str, Any]], Path, Path, Path]:
    """Transcribe saved responses with whisper-1 timestamps and align them to source windows."""
    out_dir.mkdir(parents=True, exist_ok=True)
    prepared_dir = out_dir / "prepared_audio"
    all_rows: list[dict[str, Any]] = []
    for row in exam_rows:
        response_path = Path(str(row.get("response_audio", "") or ""))
        if not response_path.exists() or not response_path.is_file():
            continue
        prepared = prepare_audio_for_transcription(response_path, prepared_dir)
        cached_timestamp_json = Path(str(row.get("timestamp_json_file", "") or ""))
        if cached_timestamp_json.exists() and cached_timestamp_json.is_file():
            timestamp_data = json.loads(cached_timestamp_json.read_text(encoding="utf-8"))
        else:
            timestamp_data = transcribe_audio_openai_with_timestamps(prepared, api_key, language_hint)
        words = normalize_timestamp_words(timestamp_data)
        source_path = Path(str(row.get("source_audio", "") or ""))
        source_duration = get_audio_duration_seconds(source_path)
        timing_map_path: Path | None = None
        if use_exact_timing_map:
            windows, timing_map_path = build_exact_source_chunk_windows(
                row,
                source_duration,
                out_dir / "source_timing_maps",
                DEFAULT_GENERATOR_PATH,
            )
        else:
            windows = estimate_source_chunk_windows(row, source_duration)
        aligned = align_timestamp_words_to_source_windows(words, windows, lag_seconds=lag_seconds, padding_seconds=padding_seconds)
        cached_timestamp_json = Path(str(row.get("timestamp_json_file", "") or ""))
        if cached_timestamp_json.exists() and cached_timestamp_json.is_file():
            raw_json_path = cached_timestamp_json
        else:
            raw_json_path = out_dir / f"{sanitize_filename_part(script_id)}_{session_id}_{sanitize_filename_part(str(row.get('item_index', 'item')))}_whisper_timestamps.json"
            raw_json_path.write_text(json.dumps(timestamp_data, ensure_ascii=False, indent=2), encoding="utf-8")
        full_transcript_path = out_dir / f"{sanitize_filename_part(script_id)}_{session_id}_{sanitize_filename_part(str(row.get('item_index', 'item')))}_timestamp_full_transcript.txt"
        full_transcript_path.write_text(str(timestamp_data.get("text", "") or words_to_text(words)), encoding="utf-8")
        for item in aligned:
            item.update({
                "script_id": row.get("script_id", script_id),
                "session_id": row.get("session_id", session_id),
                "exam_mode": row.get("exam_mode", ""),
                "item_index": row.get("item_index", ""),
                "timestamp_json_file": str(raw_json_path),
                "timestamp_full_transcript_file": str(full_transcript_path),
                "timing_map_file": str(timing_map_path) if timing_map_path else "",
            })
            all_rows.append(item)
    base = f"{sanitize_filename_part(script_id)}_{session_id}_timestamp_alignment"
    csv_path = out_dir / f"{base}.csv"
    txt_path = out_dir / f"{base}.txt"
    docx_path = out_dir / f"{base}.docx"
    csv_path.write_text(timestamp_alignment_rows_to_csv(all_rows), encoding="utf-8")
    txt_path.write_text(timestamp_alignment_rows_to_txt(all_rows), encoding="utf-8")
    build_timestamp_alignment_docx(all_rows, docx_path)
    return all_rows, txt_path, csv_path, docx_path


def render_timestamp_alignment_review(rows: list[dict[str, Any]]) -> None:
    if not rows:
        return
    st.markdown("#### Timestamp-anchored alignment review")
    labels = [f"{r.get('item_label', 'Item')} · {format_timecode(r.get('source_start'))}–{format_timecode(r.get('source_end'))}" for r in rows]
    idx = st.selectbox("Review aligned chunk", list(range(len(rows))), format_func=lambda i: labels[i], key="timestamp_alignment_review_item")
    row = rows[int(idx)]
    st.markdown(compact_metadata_row([
        str(row.get("item_label") or "Aligned chunk"),
        f"Source {format_timecode(row.get('source_start'))}–{format_timecode(row.get('source_end'))}",
        f"Response {format_timecode(row.get('response_window_start'))}–{format_timecode(row.get('response_window_end'))}",
        "Possible gap" if row.get("possible_gap") else "",
        str(row.get("timing_method") or ""),
    ]), unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3, gap="large")
    with c1:
        render_text_card("Source", str(row.get("source_text", "") or ""), meta=str(row.get("source_language", "")).title(), tone="source")
    with c2:
        render_text_card("Reference", str(row.get("reference_text", "") or ""), meta=str(row.get("target_language", "")).title(), tone="reference")
    with c3:
        transcript = str(row.get("your_transcript_aligned", "") or "").strip()
        if transcript:
            render_text_card("Your transcript", transcript, meta="Whisper timestamp window", tone="reference")
        else:
            render_hidden_reference_card("Possible omission / gap", "No matching speech was detected in this estimated response window.")


def transcript_rows_to_csv(rows: list[dict[str, Any]]) -> str:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=TRANSCRIPT_CSV_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        writer.writerow({k: row.get(k, "") for k in TRANSCRIPT_CSV_FIELDS})
    return output.getvalue()


def transcript_rows_to_txt(rows: list[dict[str, Any]]) -> str:
    chunks: list[str] = []
    for row in rows:
        label = row.get("item_label") or f"Item {row.get('item_index', '')}"
        chunks.append(
            f"{label}\n"
            f"Mode: {row.get('exam_mode', '')}\n"
            f"Direction: {row.get('source_language', '')} → {row.get('target_language', '')}\n"
            f"Speaker: {row.get('speaker', '')}\n"
            f"Type: {row.get('segment_type', '')}\n\n"
            f"Transcript:\n{row.get('transcript_text', '').strip()}\n"
        )
    return "\n" + ("\n" + "-" * 72 + "\n\n").join(chunks)


def build_exam_transcript_outputs(
    exam_rows: list[dict[str, Any]],
    script_id: str,
    session_id: str,
    transcript_dir: Path,
    model: str,
    api_key: str,
    language_hint: str | None,
) -> tuple[list[dict[str, Any]], Path, Path]:
    transcript_dir.mkdir(parents=True, exist_ok=True)
    work_dir = transcript_dir / "prepared_audio"
    transcript_rows: list[dict[str, Any]] = []
    for row in exam_rows:
        response_path = Path(str(row.get("response_audio", "")))
        if not response_path.exists() or not response_path.is_file():
            continue
        prepared = prepare_audio_for_transcription(response_path, work_dir)
        item_index = str(row.get("item_index", "item"))
        enriched = dict(row)
        if str(model).strip() == "whisper-1":
            # Unified transcript pipeline: one Whisper-1 verbose_json call provides
            # both the plain transcript and timestamp data for later review/alignment.
            timestamp_data = transcribe_audio_openai_with_timestamps(prepared, api_key, language_hint)
            text = str(timestamp_data.get("text", "") or words_to_text(normalize_timestamp_words(timestamp_data))).strip()
            timestamp_json_path = transcript_dir / f"{sanitize_filename_part(script_id)}_{session_id}_{sanitize_filename_part(item_index)}_whisper_timestamps.json"
            timestamp_json_path.write_text(json.dumps(timestamp_data, ensure_ascii=False, indent=2), encoding="utf-8")
            enriched["timestamp_json_file"] = str(timestamp_json_path)
            enriched["timestamp_full_transcript_file"] = ""
            enriched["transcription_model"] = "whisper-1"
        else:
            text = transcribe_audio_openai(prepared, model, api_key, language_hint)
            enriched["transcription_model"] = str(model)
        transcript_path = transcript_dir / f"{sanitize_filename_part(script_id)}_{session_id}_{sanitize_filename_part(item_index)}_transcript.txt"
        transcript_path.write_text(text, encoding="utf-8")
        enriched["transcript_text"] = text
        enriched["transcript_file"] = str(transcript_path)
        transcript_rows.append(enriched)

    csv_text = transcript_rows_to_csv(transcript_rows)
    csv_path = transcript_dir / f"{sanitize_filename_part(script_id)}_{session_id}_exam_transcripts.csv"
    csv_path.write_text(csv_text, encoding="utf-8")

    txt_text = transcript_rows_to_txt(transcript_rows)
    txt_path = transcript_dir / f"{sanitize_filename_part(script_id)}_{session_id}_exam_transcripts.txt"
    txt_path.write_text(txt_text, encoding="utf-8")
    return transcript_rows, txt_path, csv_path



def merge_exam_rows_with_transcripts(exam_rows: list[dict[str, Any]], transcript_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Return exam rows enriched with transcript_text/transcript_file where available."""
    by_index = {str(r.get("item_index", "")): dict(r) for r in exam_rows}
    for tr in transcript_rows or []:
        key = str(tr.get("item_index", ""))
        if key in by_index:
            by_index[key].update(tr)
        else:
            by_index[key] = dict(tr)
    def sort_key(k: str):
        if k == "full":
            return (-1, 0)
        try:
            return (0, int(k))
        except Exception:
            return (1, k)
    return [by_index[k] for k in sorted(by_index, key=sort_key)]


def structured_exam_report_rows(exam_rows: list[dict[str, Any]], transcript_rows: list[dict[str, Any]] | None = None) -> list[dict[str, Any]]:
    """Normalize report rows for in-app review, CSV, TXT, and DOCX exports."""
    rows = merge_exam_rows_with_transcripts(exam_rows, transcript_rows or [])
    report_rows: list[dict[str, Any]] = []
    for row in rows:
        report_rows.append({
            "timestamp": row.get("timestamp", ""),
            "session_id": row.get("session_id", ""),
            "script_id": row.get("script_id", ""),
            "title": row.get("title", ""),
            "source_file": row.get("source_file", ""),
            "exam_mode": row.get("exam_mode", ""),
            "item_index": row.get("item_index", ""),
            "item_label": row.get("item_label", ""),
            "source_language": row.get("source_language", ""),
            "target_language": row.get("target_language", ""),
            "speaker": row.get("speaker", ""),
            "segment_type": row.get("segment_type", ""),
            "source_text": row.get("source_text", ""),
            "reference_text": row.get("reference_text", ""),
            "your_transcript": row.get("transcript_text", ""),
            "notes": row.get("notes", ""),
            "response_audio": row.get("response_audio", ""),
            "transcript_file": row.get("transcript_file", ""),
            "playback_speed": row.get("playback_speed", ""),
            "source_word_count": row.get("source_word_count", ""),
            "source_duration_seconds": row.get("source_duration_seconds", ""),
            "heard_duration_seconds": row.get("heard_duration_seconds", ""),
            "source_wpm": row.get("source_wpm", ""),
            "effective_wpm": row.get("effective_wpm", ""),
            "wpm_target": row.get("wpm_target", ""),
            "wpm_target_delta": row.get("wpm_target_delta", ""),
        })
    return report_rows


def timestamp_alignment_rows_to_structured_report_rows(alignment_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Convert timestamp-aligned chunks into structured report rows for loose chunk-based review.

    This keeps simultaneous-mode review readable by showing Source | Reference |
    Your nearby transcript per source chunk instead of one full transcript block.
    """
    report_rows: list[dict[str, Any]] = []
    for row in alignment_rows or []:
        label = str(row.get("item_label") or "Item").strip()
        chunk = str(row.get("chunk_number") or "").strip()
        if chunk and f"Chunk {chunk}" not in label:
            label = f"{label} — Chunk {chunk}"
        transcript = str(row.get("your_transcript_aligned", "") or "").strip()
        notes = str(row.get("alignment_note", "") or "").strip()
        if row.get("source_start") != "" or row.get("source_end") != "":
            notes = (notes + " | " if notes else "") + f"Loose timing window: {row.get('source_start', '')}–{row.get('source_end', '')}s source, {row.get('review_start', '')}–{row.get('review_end', '')}s response"
        report_rows.append({
            "timestamp": row.get("timestamp", ""),
            "session_id": row.get("session_id", ""),
            "script_id": row.get("script_id", ""),
            "title": row.get("title", ""),
            "source_file": row.get("source_file", ""),
            "exam_mode": row.get("exam_mode", ""),
            "item_index": row.get("item_index", ""),
            "item_label": label,
            "source_language": row.get("source_language", ""),
            "target_language": row.get("target_language", ""),
            "speaker": row.get("speaker", ""),
            "segment_type": row.get("segment_type", ""),
            "source_text": row.get("source_text", ""),
            "reference_text": row.get("reference_text", ""),
            "your_transcript": transcript,
            "notes": notes,
            "response_audio": row.get("response_audio", ""),
            "transcript_file": row.get("timestamp_full_transcript_file", ""),
            "playback_speed": row.get("playback_speed", ""),
            "source_word_count": row.get("source_word_count", ""),
            "source_duration_seconds": row.get("source_duration_seconds", ""),
            "heard_duration_seconds": row.get("heard_duration_seconds", ""),
            "source_wpm": row.get("source_wpm", ""),
            "effective_wpm": row.get("effective_wpm", ""),
            "wpm_target": row.get("wpm_target", ""),
            "wpm_target_delta": row.get("wpm_target_delta", ""),
        })
    return report_rows


def choose_structured_review_rows(exam_rows: list[dict[str, Any]], transcript_rows: list[dict[str, Any]] | None = None, alignment_rows: list[dict[str, Any]] | None = None) -> list[dict[str, Any]]:
    """Prefer loose timestamp-aligned rows when available; otherwise use plain transcript rows."""
    if alignment_rows:
        converted = timestamp_alignment_rows_to_structured_report_rows(alignment_rows)
        if converted:
            return converted
    return structured_exam_report_rows(exam_rows, transcript_rows or [])


EXAM_REPORT_CSV_FIELDS = [
    "timestamp", "session_id", "script_id", "title", "source_file", "exam_mode", "item_index", "item_label",
    "source_language", "target_language", "speaker", "segment_type", "source_text", "reference_text",
    "your_transcript", "notes", "response_audio", "transcript_file",
    "playback_speed", "source_word_count", "source_duration_seconds", "heard_duration_seconds",
    "source_wpm", "effective_wpm", "wpm_target", "wpm_target_delta",
]


def exam_report_rows_to_csv(report_rows: list[dict[str, Any]]) -> str:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=EXAM_REPORT_CSV_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for row in report_rows:
        writer.writerow({k: row.get(k, "") for k in EXAM_REPORT_CSV_FIELDS})
    return output.getvalue()


def exam_report_rows_to_txt(report_rows: list[dict[str, Any]]) -> str:
    chunks: list[str] = []
    if report_rows:
        first = report_rows[0]
        chunks.append(
            f"Exam Transcript Report\n"
            f"Script: {first.get('script_id', '')} — {first.get('title', '')}\n"
            f"Mode: {first.get('exam_mode', '')}\n"
            f"Session: {first.get('session_id', '')}\n"
            f"Generated: {datetime.now().isoformat(timespec='seconds')}\n"
        )
    for row in report_rows:
        chunks.append(
            f"\n{'=' * 80}\n"
            f"{row.get('item_label') or 'Item'}\n"
            f"Direction: {row.get('source_language', '')} → {row.get('target_language', '')}\n"
            f"Speaker: {row.get('speaker', '')}\n"
            f"Type: {row.get('segment_type', '')}\n\n"
            f"SOURCE:\n{str(row.get('source_text', '')).strip()}\n\n"
            f"REFERENCE INTERPRETATION:\n{str(row.get('reference_text', '')).strip()}\n\n"
            f"YOUR TRANSCRIPT:\n{str(row.get('your_transcript', '')).strip()}\n\n"
            f"NOTES:\n{str(row.get('notes', '')).strip()}\n"
        )
    return "\n".join(chunks).strip() + "\n"


def build_exam_report_docx(report_rows: list[dict[str, Any]], out_path: Path) -> Path:
    """Create a structured DOCX report. Requires python-docx."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise RuntimeError("python-docx is not installed. Run: python -m pip install python-docx") from exc

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    with contextlib.suppress(Exception):
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(9)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Exam Transcript Report")
    run.bold = True
    run.font.size = Pt(16)

    first = report_rows[0] if report_rows else {}
    meta = [
        ("Script ID", first.get("script_id", "")),
        ("Title", first.get("title", "")),
        ("Mode", first.get("exam_mode", "")),
        ("Session", first.get("session_id", "")),
        ("Generated", datetime.now().isoformat(timespec="seconds")),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in meta:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)
        with contextlib.suppress(Exception):
            cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    for row in report_rows:
        heading = doc.add_paragraph()
        hr = heading.add_run(str(row.get("item_label") or "Item"))
        hr.bold = True
        hr.font.size = Pt(12)
        details = doc.add_paragraph()
        details.add_run(
            f"Direction: {row.get('source_language', '')} → {row.get('target_language', '')} | "
            f"Speaker: {row.get('speaker', '')} | Type: {row.get('segment_type', '')}"
        ).italic = True

        seg_table = doc.add_table(rows=1, cols=3)
        seg_table.style = "Table Grid"
        hdr = seg_table.rows[0].cells
        hdr[0].text = "Source"
        hdr[1].text = "Reference interpretation"
        hdr[2].text = "Your transcript"
        for cell in hdr:
            with contextlib.suppress(Exception):
                cell.paragraphs[0].runs[0].bold = True
        cells = seg_table.add_row().cells
        cells[0].text = str(row.get("source_text", "") or "")
        cells[1].text = str(row.get("reference_text", "") or "")
        cells[2].text = str(row.get("your_transcript", "") or "")

        notes = doc.add_paragraph()
        notes.add_run("Notes: ").bold = True
        notes.add_run(str(row.get("notes", "") or ""))
        doc.add_paragraph()

    doc.save(out_path)
    return out_path


def build_structured_exam_report_outputs(
    exam_rows: list[dict[str, Any]],
    transcript_rows: list[dict[str, Any]],
    script_id: str,
    session_id: str,
    report_dir: Path,
    alignment_rows: list[dict[str, Any]] | None = None,
) -> tuple[list[dict[str, Any]], Path, Path, Path]:
    report_dir.mkdir(parents=True, exist_ok=True)
    report_rows = choose_structured_review_rows(exam_rows, transcript_rows, alignment_rows)
    base = f"{sanitize_filename_part(script_id)}_{session_id}_structured_exam_report"
    csv_path = report_dir / f"{base}.csv"
    txt_path = report_dir / f"{base}.txt"
    docx_path = report_dir / f"{base}.docx"
    csv_path.write_text(exam_report_rows_to_csv(report_rows), encoding="utf-8")
    txt_path.write_text(exam_report_rows_to_txt(report_rows), encoding="utf-8")
    build_exam_report_docx(report_rows, docx_path)
    return report_rows, txt_path, csv_path, docx_path


def _read_first_existing_text_path(paths: list[str]) -> str:
    """Read the first valid text file path in a list."""
    for value in paths or []:
        value = str(value or "").strip()
        if not value:
            continue
        with contextlib.suppress(Exception):
            path = Path(value)
            if path.exists() and path.is_file():
                return path.read_text(encoding="utf-8").strip()
    return ""


def _join_unique_text_blocks(rows: list[dict[str, Any]], field: str) -> str:
    """Join unique text blocks in order, avoiding exact duplicate adjacent text."""
    blocks: list[str] = []
    seen: set[str] = set()
    for row in rows or []:
        value = str(row.get(field, "") or "").strip()
        if not value:
            continue
        norm = re.sub(r"\s+", " ", value).strip().lower()
        if norm in seen:
            continue
        seen.add(norm)
        blocks.append(value)
    return "\n\n".join(blocks).strip()


def _plain_full_transcript_from_rows(report_rows: list[dict[str, Any]]) -> str:
    """Prefer the saved full transcript file; fall back to non-empty transcript rows."""
    candidate_paths: list[str] = []
    for row in report_rows or []:
        candidate_paths.extend([
            str(row.get("timestamp_full_transcript_file", "") or ""),
            str(row.get("transcript_file", "") or ""),
        ])
    file_text = _read_first_existing_text_path(candidate_paths)
    if file_text:
        return file_text
    # Fallback: for consecutive rows this is naturally one response per row.
    # For simultaneous alignment rows, this fallback may still be chunked, so it
    # is used only when the saved full transcript is unavailable.
    return _join_unique_text_blocks(report_rows, "your_transcript")


def render_chunk_formatted_full_transcript(report_rows: list[dict[str, Any]]) -> None:
    """Display a full 3-column review without forcing the transcript into timestamp slices.

    The detailed chunk/timestamp rows remain available in an expander. This view
    intentionally shows the user's transcript as one continuous transcript because
    timestamp windows can drift during simultaneous interpreting.
    """
    full_transcript = _plain_full_transcript_from_rows(report_rows)
    if not full_transcript:
        render_hidden_reference_card("No transcript yet", "Transcribe saved exam recordings before building the structured report.")
        return

    full_source = _join_unique_text_blocks(report_rows, "source_text")
    full_reference = _join_unique_text_blocks(report_rows, "reference_text")

    st.markdown("##### Full transcript review")
    st.caption("This view shows the complete transcript as continuous text. Timestamp/chunk rows are kept below only as optional detail, because simultaneous interpretation can lag or lead the source.")
    c1, c2, c3 = st.columns(3, gap="large")
    with c1:
        render_text_card("Source", full_source, meta="Full source text", tone="source")
    with c2:
        render_text_card("Reference interpretation", full_reference, meta="Full reference text", tone="reference")
    with c3:
        render_text_card("Your full transcript", full_transcript, meta="Continuous transcript", tone="reference")

    with st.expander("Show loose chunk transcript details", expanded=False):
        transcript_parts: list[str] = []
        for i, row in enumerate(report_rows, start=1):
            transcript = str(row.get("your_transcript", row.get("your_transcript_aligned", "")) or "").strip()
            if not transcript:
                continue
            label = str(row.get("item_label") or f"Item {i}").strip()
            timing = ""
            if str(row.get("source_start", "")).strip() or str(row.get("source_end", "")).strip():
                timing = f" · source {format_timecode(row.get('source_start'))}–{format_timecode(row.get('source_end'))}"
            wpm = row.get("effective_wpm", row.get("source_wpm", ""))
            wpm_text = f" · {wpm} WPM" if str(wpm).strip() else ""
            transcript_parts.append(f"**{label}{timing}{wpm_text}**\n\n{transcript}")
        if transcript_parts:
            st.markdown("\n\n---\n\n".join(transcript_parts))
        else:
            st.caption("No chunk-level transcript details available.")


def render_exam_report_review(report_rows: list[dict[str, Any]]) -> None:
    """Show a compact side-by-side review of source/reference/transcript in Exam Mode."""
    if not report_rows:
        return
    st.markdown("#### Structured transcript review")
    if len(report_rows) > 1:
        view_mode = st.radio(
            "Structured review display",
            ["Full transcript view", "Selected item detail"],
            horizontal=True,
            key="exam_structured_review_display_mode",
        )
        if view_mode == "Full transcript view":
            render_chunk_formatted_full_transcript(report_rows)
            with st.expander("Show source/reference alignment rows", expanded=False):
                for row in report_rows:
                    st.markdown(compact_metadata_row([
                        str(row.get("item_label") or "Item"),
                        f"{str(row.get('source_language', ''))[:2].upper()} → {str(row.get('target_language', ''))[:2].upper()}",
                        str(row.get("speaker") or ""),
                        str(row.get("segment_type") or ""),
                        (f"{row.get('effective_wpm')} WPM" if str(row.get("effective_wpm", "")).strip() else ""),
                    ]), unsafe_allow_html=True)
                    c1, c2, c3 = st.columns(3, gap="large")
                    with c1:
                        render_text_card("Source", str(row.get("source_text", "") or ""), meta=str(row.get("source_language", "")).title(), tone="source")
                    with c2:
                        render_text_card("Reference", str(row.get("reference_text", "") or ""), meta=str(row.get("target_language", "")).title(), tone="reference")
                    with c3:
                        transcript = str(row.get("your_transcript", "") or "").strip()
                        if transcript:
                            render_text_card("Your transcript", transcript, meta="Loose chunk-aligned", tone="reference")
                        else:
                            render_hidden_reference_card("No transcript for this item", "")
            return
        labels = [str(r.get("item_label") or f"Item {i+1}") for i, r in enumerate(report_rows)]
        idx = st.selectbox("Review item", list(range(len(report_rows))), format_func=lambda i: labels[i], key="exam_report_review_item")
        rows_to_show = [report_rows[int(idx)]]
    else:
        rows_to_show = report_rows
    for row in rows_to_show:
        st.markdown(compact_metadata_row([
            str(row.get("item_label") or "Item"),
            f"{str(row.get('source_language', ''))[:2].upper()} → {str(row.get('target_language', ''))[:2].upper()}",
            str(row.get("speaker") or ""),
            str(row.get("segment_type") or ""),
            (f"{row.get('effective_wpm')} WPM" if str(row.get("effective_wpm", "")).strip() else ""),
        ]), unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3, gap="large")
        with c1:
            render_text_card("Source", str(row.get("source_text", "") or ""), meta=str(row.get("source_language", "")).title(), tone="source")
        with c2:
            render_text_card("Reference", str(row.get("reference_text", "") or ""), meta=str(row.get("target_language", "")).title(), tone="reference")
        with c3:
            transcript = str(row.get("your_transcript", "") or "").strip()
            if transcript:
                render_text_card("Your transcript", transcript, meta="OpenAI transcript", tone="reference")
            else:
                render_hidden_reference_card("No transcript yet", "Transcribe saved exam recordings before building the structured report.")



# -----------------------------
# Phase 3C.1 rule-based comparison helpers
# -----------------------------
REVIEW_CSV_FIELDS = [
    "severity", "category", "item_label", "direction", "speaker", "segment_type",
    "expected", "found", "details", "source_text", "reference_text", "your_transcript",
]


def normalize_for_rule_compare(text: str) -> str:
    """Lowercase text with light punctuation removal for deterministic comparison."""
    text = unicodedata.normalize("NFKD", str(text or ""))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = text.lower()
    text = re.sub(r"[^\w\s$%.-]", " ", text, flags=re.UNICODE)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def normalize_for_matching(text: str) -> str:
    """Compatibility helper used by Drill Builder and Drill Studio matching logic."""
    return normalize_for_rule_compare(text)


def phrase_present(phrase: str, text: str) -> bool:
    phrase_n = normalize_for_rule_compare(phrase)
    text_n = normalize_for_rule_compare(text)
    if not phrase_n:
        return False
    if phrase_n in text_n:
        return True
    # Helpful fallback for multi-word terms: require most meaningful tokens.
    tokens = [t for t in phrase_n.split() if len(t) > 2]
    if len(tokens) >= 2:
        hits = sum(1 for t in tokens if re.search(rf"\b{re.escape(t)}\b", text_n))
        return hits >= max(2, int(len(tokens) * 0.75))
    return any(re.search(rf"\b{re.escape(t)}\b", text_n) for t in tokens)


# Bilingual normalization helpers for rule-based review.
EN_NUM_WORDS = {
    "zero": 0, "one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6, "seven": 7,
    "eight": 8, "nine": 9, "ten": 10, "eleven": 11, "twelve": 12, "thirteen": 13,
    "fourteen": 14, "fifteen": 15, "sixteen": 16, "seventeen": 17, "eighteen": 18,
    "nineteen": 19, "twenty": 20, "thirty": 30, "forty": 40, "fifty": 50, "sixty": 60,
    "seventy": 70, "eighty": 80, "ninety": 90,
}
ES_NUM_WORDS = {
    "cero": 0, "uno": 1, "una": 1, "un": 1, "dos": 2, "tres": 3, "cuatro": 4, "cinco": 5,
    "seis": 6, "siete": 7, "ocho": 8, "nueve": 9, "diez": 10, "once": 11, "doce": 12,
    "trece": 13, "catorce": 14, "quince": 15, "dieciseis": 16, "dieciséis": 16, "diecisiete": 17,
    "dieciocho": 18, "diecinueve": 19, "veinte": 20, "veintiuno": 21, "veintiuna": 21,
    "veintidos": 22, "veintidós": 22, "veintitres": 23, "veintitrés": 23, "veinticuatro": 24,
    "veinticinco": 25, "veintiseis": 26, "veintiséis": 26, "veintisiete": 27, "veintiocho": 28,
    "veintinueve": 29, "treinta": 30, "cuarenta": 40, "cincuenta": 50, "sesenta": 60,
    "setenta": 70, "ochenta": 80, "noventa": 90,
    "cien": 100, "ciento": 100, "doscientos": 200, "doscientas": 200, "trescientos": 300,
    "trescientas": 300, "cuatrocientos": 400, "cuatrocientas": 400, "quinientos": 500,
    "quinientas": 500, "seiscientos": 600, "seiscientas": 600, "setecientos": 700,
    "setecientas": 700, "ochocientos": 800, "ochocientas": 800, "novecientos": 900,
    "novecientas": 900, "mil": 1000,
}
MONTHS = {
    "january": 1, "jan": 1, "enero": 1,
    "february": 2, "feb": 2, "febrero": 2,
    "march": 3, "mar": 3, "marzo": 3,
    "april": 4, "apr": 4, "abril": 4,
    "may": 5, "mayo": 5,
    "june": 6, "jun": 6, "junio": 6,
    "july": 7, "jul": 7, "julio": 7,
    "august": 8, "aug": 8, "agosto": 8,
    "september": 9, "sep": 9, "sept": 9, "septiembre": 9, "setiembre": 9,
    "october": 10, "oct": 10, "octubre": 10,
    "november": 11, "nov": 11, "noviembre": 11,
    "december": 12, "dec": 12, "diciembre": 12,
}


def _strip_ordinals(s: str) -> str:
    return re.sub(r"\b(\d{1,2})(st|nd|rd|th)\b", r"\1", str(s or ""), flags=re.IGNORECASE)


def _words_to_small_number(tokens: list[str]) -> int | None:
    """Very small bilingual number parser for fact-checking transcripts, not legal grading."""
    norm = [normalize_for_rule_compare(t) for t in tokens if normalize_for_rule_compare(t) and normalize_for_rule_compare(t) not in {"and", "y", "de", "del"}]
    if not norm:
        return None
    total = 0
    current = 0
    found = False
    for tok in norm:
        val = EN_NUM_WORDS.get(tok, ES_NUM_WORDS.get(tok))
        if val is None:
            return None
        found = True
        if val == 1000:
            current = max(current, 1) * 1000
            total += current
            current = 0
        elif val == 100 and current and current < 10:
            current *= 100
        else:
            current += val
    total += current
    return total if found else None


def numeric_variants(value: str) -> set[str]:
    """Return normalized variants for a number or amount, including simple bilingual word forms found in transcripts."""
    raw = normalize_for_rule_compare(value)
    variants = {raw} if raw else set()
    digits = re.findall(r"\d+", raw.replace(",", ""))
    for d in digits:
        variants.add(str(int(d)))
    return {v for v in variants if v}


def transcript_has_number(value: str, transcript: str) -> tuple[bool, str]:
    tr_n = normalize_for_rule_compare(transcript)
    for var in numeric_variants(value):
        if re.search(rf"\b{re.escape(var)}\b", tr_n):
            return True, var
    # Compare simple number words in transcript against numeric expected values.
    expected_nums = {int(v) for v in numeric_variants(value) if v.isdigit()}
    if expected_nums:
        words = tr_n.split()
        for span in range(1, 5):
            for i in range(0, max(0, len(words) - span + 1)):
                val = _words_to_small_number(words[i:i+span])
                if val in expected_nums:
                    return True, " ".join(words[i:i+span])
    return False, "Not clearly detected"


def canonical_date(value: str) -> tuple[int | None, int | None, int | None]:
    raw = normalize_for_rule_compare(_strip_ordinals(value))
    # English: October 15 2026 / Oct 15
    m = re.search(r"\b([a-z]+)\s+(\d{1,2})(?:\s+(\d{4}))?\b", raw)
    if m and m.group(1) in MONTHS:
        return (MONTHS[m.group(1)], int(m.group(2)), int(m.group(3)) if m.group(3) else None)
    # Spanish: 15 de octubre de 2026
    m = re.search(r"\b(\d{1,2})\s+(?:de\s+)?([a-z]+)(?:\s+de\s+(\d{4}))?\b", raw)
    if m and m.group(2) in MONTHS:
        return (MONTHS[m.group(2)], int(m.group(1)), int(m.group(3)) if m.group(3) else None)
    # Numeric: 10/15/2026 or 15/10/2026; keep flexible by accepting either day/month order later.
    m = re.search(r"\b(\d{1,2})[/-](\d{1,2})(?:[/-](\d{2,4}))?\b", raw)
    if m:
        a, b = int(m.group(1)), int(m.group(2))
        y = int(m.group(3)) if m.group(3) else None
        if y and y < 100:
            y += 2000
        # Ambiguous; return month/day when plausible U.S. format.
        return (a if a <= 12 else b, b if a <= 12 else a, y)
    # Word-day Spanish/English: quince de octubre / october fifteen
    toks = raw.split()
    for i, tok in enumerate(toks):
        if tok in MONTHS:
            # look around for day as digit or word
            window = toks[max(0, i-3): min(len(toks), i+4)]
            for j in range(len(window)):
                if window[j].isdigit():
                    return (MONTHS[tok], int(window[j]), None)
                val = _words_to_small_number([window[j]])
                if val and 1 <= val <= 31:
                    return (MONTHS[tok], val, None)
    return (None, None, None)


def transcript_has_date(value: str, transcript: str) -> tuple[bool, str]:
    month, day, year = canonical_date(value)
    if not month or not day:
        # Fallback to phrase match.
        return (phrase_present(value, transcript), value if phrase_present(value, transcript) else "Not clearly detected")
    tr_n = normalize_for_rule_compare(_strip_ordinals(transcript))
    # Search all date-like snippets from the transcript and compare canonically.
    candidates = []
    candidates += re.findall(r"\b[a-z]+\s+\d{1,2}(?:\s+\d{4})?\b", tr_n)
    candidates += re.findall(r"\b\d{1,2}\s+(?:de\s+)?[a-z]+(?:\s+de\s+\d{4})?\b", tr_n)
    candidates += re.findall(r"\b\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?\b", tr_n)
    # Also test short windows for word-day dates.
    words = tr_n.split()
    candidates += [" ".join(words[i:i+4]) for i in range(max(0, len(words)-3))]
    for cand in candidates:
        cm, cd, cy = canonical_date(cand)
        if cm == month and cd == day and (year is None or cy is None or cy == year):
            return True, cand
    return False, "Not clearly detected"


def fuzzy_name_present(name: str, transcript: str) -> tuple[bool, str, str]:
    """Return (present, found_text, note). Allows ASR spelling variants for names."""
    name_n = normalize_for_rule_compare(name)
    tr_n = normalize_for_rule_compare(transcript)
    if not name_n:
        return False, "", ""
    if name_n in tr_n:
        return True, name, "exact/normalized match"
    name_tokens = [t for t in name_n.split() if len(t) > 1]
    tr_tokens = [t for t in tr_n.split() if len(t) > 1]
    if not name_tokens or not tr_tokens:
        return False, "Not clearly detected", ""
    # Require either a close full-name window or each token fuzzy-present somewhere nearby.
    n = len(name_tokens)
    best_ratio = 0.0
    best_window = ""
    for i in range(0, max(1, len(tr_tokens) - n + 1)):
        window = " ".join(tr_tokens[i:i+n])
        ratio = difflib.SequenceMatcher(None, name_n, window).ratio()
        if ratio > best_ratio:
            best_ratio, best_window = ratio, window
    if best_ratio >= 0.82:
        return True, best_window, "likely present; transcript spelling may vary"
    hits = 0
    matched = []
    for nt in name_tokens:
        token_hit = ""
        for tt in tr_tokens:
            if nt == tt or difflib.SequenceMatcher(None, nt, tt).ratio() >= 0.84:
                token_hit = tt
                break
        if token_hit:
            hits += 1
            matched.append(token_hit)
    if hits >= max(1, len(name_tokens) - 1):
        return True, " ".join(matched), "likely present; partial/fuzzy name match"
    return False, "Not clearly detected", "transcription may spell names unpredictably; verify manually"


def extract_numbers_dates_names(text: str) -> dict[str, list[str]]:
    """Extract stable high-risk items from source/reference text for soft rule checks."""
    raw = str(text or "")
    numbers = set()
    # Money/percent/decimals/integers and common case-style numbers.
    for m in re.findall(r"\$\s?\d[\d,]*(?:\.\d+)?|\b\d+(?:[,.]\d+)*(?:\.\d+)?\s?%?\b", raw):
        cleaned = re.sub(r"\s+", "", m).strip(".,;:")
        if cleaned:
            numbers.add(cleaned)
    case_numbers = set(re.findall(r"\b[A-Z]{1,4}\s?-?\d{2,4}-?\d{1,8}\b", raw))
    dates = set()
    month_names = r"January|February|March|April|May|June|July|August|September|October|November|December|Jan\.?|Feb\.?|Mar\.?|Apr\.?|Jun\.?|Jul\.?|Aug\.?|Sep\.?|Sept\.?|Oct\.?|Nov\.?|Dec\.?|Enero|Febrero|Marzo|Abril|Mayo|Junio|Julio|Agosto|Septiembre|Setiembre|Octubre|Noviembre|Diciembre"
    for m in re.findall(rf"\b(?:{month_names})\s+\d{{1,2}}(?:st|nd|rd|th)?(?:,\s*\d{{4}})?\b", raw, flags=re.IGNORECASE):
        dates.add(m.strip())
    for m in re.findall(rf"\b\d{{1,2}}\s+de\s+(?:{month_names})(?:\s+de\s+\d{{4}})?\b", raw, flags=re.IGNORECASE):
        dates.add(m.strip())
    for m in re.findall(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", raw):
        dates.add(m.strip())
    names = set()
    # Conservative proper-noun sequences. Avoid all-caps role labels.
    for m in re.findall(r"\b[A-Z][a-záéíóúñü]+(?:\s+[A-Z][a-záéíóúñü]+){1,3}\b", raw):
        if len(m.split()) >= 2:
            names.add(m.strip())
    return {
        "numbers": sorted(numbers),
        "case_numbers": sorted(case_numbers),
        "dates": sorted(dates),
        "names": sorted(names),
    }

def build_terms_for_row(payload: dict[str, Any], row: dict[str, Any]) -> list[dict[str, str]]:
    """Return terms from JSON terms_used that appear in the row source/reference."""
    terms = payload.get("terms_used") or []
    source = str(row.get("source_text", "") or "")
    reference = str(row.get("reference_text", "") or "")
    out: list[dict[str, str]] = []
    for term in terms:
        if not isinstance(term, dict):
            continue
        english = str(term.get("english", "") or "").strip()
        spanish = str(term.get("spanish_contextual_choice", "") or "").strip()
        # Some preferred entries include alternatives with slashes. Check the most concrete pieces too.
        candidates = [english, spanish]
        for val in list(candidates):
            candidates.extend([part.strip() for part in re.split(r"/|;|,", val) if part.strip()])
        if any(c and (phrase_present(c, source) or phrase_present(c, reference)) for c in candidates):
            out.append({"english": english, "spanish": spanish})
    return out


def add_review_issue(issues: list[dict[str, Any]], row: dict[str, Any], severity: str, category: str, expected: str, found: str = "", details: str = "") -> None:
    issues.append({
        "severity": severity,
        "category": category,
        "item_label": row.get("item_label", ""),
        "direction": f"{row.get('source_language', '')} → {row.get('target_language', '')}",
        "speaker": row.get("speaker", ""),
        "segment_type": row.get("segment_type", ""),
        "expected": expected,
        "found": found,
        "details": details,
        "source_text": row.get("source_text", ""),
        "reference_text": row.get("reference_text", ""),
        "your_transcript": row.get("your_transcript", row.get("your_transcript_aligned", "")),
    })


def compare_exam_rows_rule_based(payload: dict[str, Any], rows: list[dict[str, Any]], use_aligned_transcript: bool = False) -> list[dict[str, Any]]:
    """Deterministic checks for likely omissions and high-risk legal details."""
    issues: list[dict[str, Any]] = []
    for row in rows:
        transcript = str(row.get("your_transcript_aligned" if use_aligned_transcript else "your_transcript", "") or "").strip()
        if not transcript:
            add_review_issue(issues, row, "high", "Likely omission", "Any interpretation", "[No transcript]", "No transcript was available for this item/window.")
            continue
        reference = str(row.get("reference_text", "") or "")
        source = str(row.get("source_text", "") or "")
        ref_words = len(reference.split())
        tr_words = len(transcript.split())
        if ref_words >= 12 and tr_words <= max(3, int(ref_words * 0.25)):
            add_review_issue(issues, row, "high", "Likely omission", f"Reference length about {ref_words} words", f"Transcript about {tr_words} words", "Transcript is much shorter than the reference for this item/window.")
        elif ref_words >= 20 and tr_words <= int(ref_words * 0.45):
            add_review_issue(issues, row, "medium", "Possible omission", f"Reference length about {ref_words} words", f"Transcript about {tr_words} words", "Transcript is substantially shorter than the reference.")

        # Terms: use target-language preferred term for the expected output when available.
        for term in build_terms_for_row(payload, row):
            target_lang = str(row.get("target_language", "") or "").lower()
            expected = term.get("spanish") if target_lang.startswith("span") else term.get("english")
            alternate = term.get("english") if target_lang.startswith("span") else term.get("spanish")
            expected = str(expected or alternate or "").strip()
            if not expected:
                continue
            term_parts = [expected] + [p.strip() for p in re.split(r"/|;", expected) if p.strip()]
            if not any(phrase_present(part, transcript) for part in term_parts):
                add_review_issue(issues, row, "medium", "Legal term check", expected, "Not detected", f"Term from JSON terms_used may be missing or rendered differently. Source term: {term.get('english', '')}")

        # Numbers/dates/names/case numbers: check source and reference because either side can carry a critical item.
        # These are intentionally soft review prompts, not confirmed errors. ASR often spells names oddly, and
        # interpreted dates/numbers can switch between English, Spanish, digits, and words.
        extracted = extract_numbers_dates_names(source + "\n" + reference)
        for expected in extracted.get("case_numbers", []):
            if expected:
                present, found = transcript_has_number(expected, transcript)
                if not present and not phrase_present(expected, transcript):
                    add_review_issue(
                        issues, row, "medium", "Case number check", expected, "Not clearly detected",
                        "Automated review prompt: verify this case number manually. Transcript spelling/formatting may differ."
                    )
        for expected in extracted.get("numbers", []):
            if expected:
                present, found = transcript_has_number(expected, transcript)
                if not present:
                    add_review_issue(
                        issues, row, "low", "Number / amount check", expected, found,
                        "Automated review prompt: the number/amount was not clearly detected after digit/word normalization. Verify manually."
                    )
        for expected in extracted.get("dates", []):
            if expected:
                present, found = transcript_has_date(expected, transcript)
                if not present:
                    add_review_issue(
                        issues, row, "low", "Date check", expected, found,
                        "Automated review prompt: date formats are normalized across English/Spanish where possible, but verify manually."
                    )
        for expected in extracted.get("names", []):
            if expected:
                present, found, note = fuzzy_name_present(expected, transcript)
                if not present:
                    add_review_issue(
                        issues, row, "low", "Name check", expected, found,
                        "Automated review prompt: name was not clearly detected. ASR spelling can be unreliable; verify manually. " + note
                    )
    severity_order = {"high": 0, "medium": 1, "low": 2}
    return sorted(issues, key=lambda x: (severity_order.get(str(x.get("severity", "low")), 9), str(x.get("item_label", "")), str(x.get("category", ""))))


def review_issues_to_csv(issues: list[dict[str, Any]]) -> str:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=REVIEW_CSV_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for issue in issues:
        writer.writerow({k: issue.get(k, "") for k in REVIEW_CSV_FIELDS})
    return output.getvalue()


def review_issues_to_txt(issues: list[dict[str, Any]]) -> str:
    chunks = ["Rule-Based Exam Review List (Automated Prompts)", f"Generated: {datetime.now().isoformat(timespec='seconds')}", ""]
    if not issues:
        chunks.append("No rule-based review prompts were detected.")
    for issue in issues:
        chunks.append(
            f"{issue.get('severity', '').upper()} — {issue.get('category', '')}\n"
            f"Item: {issue.get('item_label', '')}\n"
            f"Direction: {issue.get('direction', '')}\n"
            f"Expected: {issue.get('expected', '')}\n"
            f"Found: {issue.get('found', '')}\n"
            f"Details: {issue.get('details', '')}\n"
        )
    return "\n".join(chunks).strip() + "\n"


def build_rule_review_docx(issues: list[dict[str, Any]], out_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise RuntimeError("python-docx is not installed. Run: python -m pip install python-docx") from exc
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    with contextlib.suppress(Exception):
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(9)
    title = doc.add_paragraph()
    r = title.add_run("Rule-Based Exam Review List (Automated Prompts)")
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    if not issues:
        doc.add_paragraph("No rule-based review prompts were detected.")
    for issue in issues:
        p = doc.add_paragraph()
        h = p.add_run(f"{str(issue.get('severity', '')).upper()} — {issue.get('category', '')}")
        h.bold = True
        doc.add_paragraph(f"Item: {issue.get('item_label', '')}")
        doc.add_paragraph(f"Direction: {issue.get('direction', '')} | Speaker: {issue.get('speaker', '')} | Type: {issue.get('segment_type', '')}")
        doc.add_paragraph(f"Expected: {issue.get('expected', '')}")
        doc.add_paragraph(f"Found: {issue.get('found', '')}")
        doc.add_paragraph(f"Details: {issue.get('details', '')}")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Source"
        hdr[1].text = "Reference"
        hdr[2].text = "Your transcript"
        cells = table.add_row().cells
        cells[0].text = str(issue.get("source_text", "") or "")
        cells[1].text = str(issue.get("reference_text", "") or "")
        cells[2].text = str(issue.get("your_transcript", "") or "")
        doc.add_paragraph()
    doc.save(out_path)
    return out_path


def build_rule_based_review_outputs(payload: dict[str, Any], rows: list[dict[str, Any]], script_id: str, session_id: str, out_dir: Path, use_aligned_transcript: bool = False) -> tuple[list[dict[str, Any]], Path, Path, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    issues = compare_exam_rows_rule_based(payload, rows, use_aligned_transcript=use_aligned_transcript)
    suffix = "timestamp_aligned" if use_aligned_transcript else "structured"
    base = f"{sanitize_filename_part(script_id)}_{session_id}_rule_review_{suffix}"
    csv_path = out_dir / f"{base}.csv"
    txt_path = out_dir / f"{base}.txt"
    docx_path = out_dir / f"{base}.docx"
    csv_path.write_text(review_issues_to_csv(issues), encoding="utf-8")
    txt_path.write_text(review_issues_to_txt(issues), encoding="utf-8")
    build_rule_review_docx(issues, docx_path)
    return issues, txt_path, csv_path, docx_path


def render_rule_review_summary(issues: list[dict[str, Any]]) -> None:
    if not issues:
        st.success("No rule-based review prompts detected.")
        return
    high = sum(1 for i in issues if str(i.get("severity", "")).lower() == "high")
    med = sum(1 for i in issues if str(i.get("severity", "")).lower() == "medium")
    low = sum(1 for i in issues if str(i.get("severity", "")).lower() == "low")
    c1, c2, c3 = st.columns(3)
    c1.metric("High", high)
    c2.metric("Medium", med)
    c3.metric("Low", low)
    st.dataframe([{k: i.get(k, "") for k in ["severity", "category", "item_label", "expected", "details"]} for i in issues], use_container_width=True, hide_index=True)


# -----------------------------
# Phase 3C.2A AI-assisted feedback helpers
# -----------------------------
AI_FEEDBACK_CSV_FIELDS = [
    "item_label", "review_mode", "context_policy", "model", "depth", "overall_confidence",
    "summary", "strengths", "issues", "suggested_drill", "source_text", "reference_text", "your_transcript",
]


def get_row_transcript_for_feedback(row: dict[str, Any]) -> str:
    return str(row.get("your_transcript_aligned", row.get("your_transcript", "")) or "").strip()


def is_simultaneous_feedback_rows(rows: list[dict[str, Any]]) -> bool:
    if not rows:
        return False
    mode_values = " ".join(str(r.get("exam_mode", "")) for r in rows).lower()
    if "simultaneous" in mode_values:
        return True
    return any(str(r.get("timing_method", "")).strip() for r in rows) and len(rows) > 1


def feedback_window_for_row(rows: list[dict[str, Any]], index: int, tolerance: str) -> dict[str, Any]:
    """Build strict or soft context around a row for AI review.

    Consecutive review stays narrow. Simultaneous review intentionally includes
    neighboring chunks and transcript text so lag/lead drift does not become a
    false omission.
    """
    row = rows[index]
    tolerance = str(tolerance or "Conservative")
    if not is_simultaneous_feedback_rows(rows):
        return {
            "row": row,
            "context_policy": "Strict per-segment review",
            "source_context": str(row.get("source_text", "") or ""),
            "reference_context": str(row.get("reference_text", "") or ""),
            "transcript_context": get_row_transcript_for_feedback(row),
            "neighbor_labels": "",
        }
    radius = {"Strict": 0, "Balanced": 1, "Conservative": 2}.get(tolerance, 2)
    start = max(0, index - radius)
    end = min(len(rows), index + radius + 1)
    window_rows = rows[start:end]
    def join_field(field: str) -> str:
        parts = []
        for i, r in enumerate(window_rows, start=start):
            label = str(r.get("item_label", f"Item {i+1}") or f"Item {i+1}")
            val = str(r.get(field, "") or "").strip()
            if val:
                parts.append(f"[{label}] {val}")
        return "\n\n".join(parts)
    transcript_parts = []
    for i, r in enumerate(window_rows, start=start):
        label = str(r.get("item_label", f"Item {i+1}") or f"Item {i+1}")
        val = get_row_transcript_for_feedback(r)
        if val:
            transcript_parts.append(f"[{label}] {val}")
    return {
        "row": row,
        "context_policy": f"Soft simultaneous window: {tolerance} ({start+1}-{end} of {len(rows)})",
        "source_context": join_field("source_text"),
        "reference_context": join_field("reference_text"),
        "transcript_context": "\n\n".join(transcript_parts),
        "neighbor_labels": ", ".join(str(r.get("item_label", "")) for r in window_rows),
    }


def issues_for_feedback_row(rule_issues: list[dict[str, Any]], row: dict[str, Any]) -> list[dict[str, Any]]:
    label = str(row.get("item_label", "") or "")
    out = []
    for issue in rule_issues or []:
        if str(issue.get("item_label", "") or "") == label:
            out.append({k: issue.get(k, "") for k in ["severity", "category", "expected", "found", "details"]})
    return out


def ai_feedback_prompt_payload(payload: dict[str, Any], window: dict[str, Any], rule_issues: list[dict[str, Any]], depth: str) -> str:
    row = window["row"]
    terms = payload.get("terms_used") or []
    compact_terms = []
    for term in terms[:80]:
        if isinstance(term, dict):
            compact_terms.append({
                "english": term.get("english", ""),
                "spanish_contextual_choice": term.get("spanish_contextual_choice", ""),
            })
    data = {
        "review_purpose": "Interpreter practice feedback only, not an official score or certification decision.",
        "depth": depth,
        "item": {
            "item_label": row.get("item_label", ""),
            "source_language": row.get("source_language", ""),
            "target_language": row.get("target_language", ""),
            "speaker": row.get("speaker", ""),
            "segment_type": row.get("segment_type", ""),
            "source_start": row.get("source_start", ""),
            "source_end": row.get("source_end", ""),
            "response_window_start": row.get("response_window_start", ""),
            "response_window_end": row.get("response_window_end", ""),
            "wpm": row.get("effective_wpm", ""),
        },
        "context_policy": window.get("context_policy", ""),
        "alignment_warning": "For simultaneous mode, transcript timing may lag or lead the source. Do not flag omissions if the idea appears in the neighboring context window.",
        "source_context": window.get("source_context", ""),
        "reference_context": window.get("reference_context", ""),
        "your_transcript_context": window.get("transcript_context", ""),
        "rule_based_flags_for_this_item": issues_for_feedback_row(rule_issues, row),
        "rule_based_flags_warning": "Automated flags are review prompts, not confirmed errors. They may include false positives from bilingual date formats, number words, ASR spelling of names, or timestamp drift. Do not treat them as proven mistakes without support from the source/reference/transcript context.",
        "terms_used_reference": compact_terms,
        "return_format": {
            "summary": "brief study-focused summary",
            "overall_confidence": "High/Medium/Low",
            "strengths": ["..."],
            "issues": [
                {"type": "Possible omission | Meaning shift | Legal terminology | Number/date/name | Register | Other", "severity": "High/Medium/Low", "confidence": "High/Medium/Low", "evidence": "short quote or description", "comment": "why it matters", "suggested_review": "what to practice"}
            ],
            "suggested_drill": "one practical drill",
        },
    }
    return json.dumps(data, ensure_ascii=False, indent=2)


def parse_ai_feedback_json(text: str) -> dict[str, Any]:
    raw = str(text or "").strip()
    if not raw:
        return {"summary": "No AI feedback text returned.", "overall_confidence": "Low", "strengths": [], "issues": [], "suggested_drill": ""}
    # Strip common markdown fences.
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE).strip()
    raw = re.sub(r"\s*```$", "", raw).strip()
    with contextlib.suppress(Exception):
        parsed = json.loads(raw)
        if isinstance(parsed, dict):
            return parsed
    # Try extracting first JSON object.
    m = re.search(r"\{.*\}", raw, flags=re.DOTALL)
    if m:
        with contextlib.suppress(Exception):
            parsed = json.loads(m.group(0))
            if isinstance(parsed, dict):
                return parsed
    return {
        "summary": raw[:1200],
        "overall_confidence": "Low",
        "strengths": [],
        "issues": [{"type": "Unstructured feedback", "severity": "Low", "confidence": "Low", "evidence": "", "comment": raw, "suggested_review": "Review this item manually."}],
        "suggested_drill": "Review this item manually and compare source, reference, and transcript side by side.",
    }


def openai_generate_ai_feedback(model: str, api_key: str, prompt_payload: str, depth: str) -> dict[str, Any]:
    if not api_key:
        raise ValueError("Missing OpenAI API key. Enter one in the app or set OPENAI_API_KEY in your environment.")
    try:
        from openai import OpenAI
    except Exception as exc:
        raise RuntimeError("The openai package is not installed. Run: python -m pip install openai") from exc
    client = OpenAI(api_key=api_key)
    instructions = (
        "You are an interpreter-training reviewer. Provide careful study feedback for legal interpreting practice. "
        "Do not give an official exam score. Be conservative with omissions, especially in simultaneous mode where transcripts may drift by a few seconds. "
        "Only identify high-confidence issues when the meaning is clearly absent, distorted, or terminologically risky. "
        "Return only valid JSON matching the requested structure."
    )
    text = ""
    # Prefer the current Responses API; fall back for older SDK installs.
    try:
        response = client.responses.create(
            model=model,
            instructions=instructions,
            input=prompt_payload,
        )
        text = str(getattr(response, "output_text", "") or response_to_transcript_text(response)).strip()
    except Exception:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": instructions},
                {"role": "user", "content": prompt_payload},
            ],
            temperature=0.2 if str(depth).lower().startswith("detailed") else 0.1,
        )
        text = str(response.choices[0].message.content or "").strip()
    return parse_ai_feedback_json(text)


def build_ai_feedback_rows(
    payload: dict[str, Any],
    rows: list[dict[str, Any]],
    rule_issues: list[dict[str, Any]],
    model: str,
    api_key: str,
    depth: str,
    tolerance: str,
    selected_index: int | None = None,
) -> list[dict[str, Any]]:
    indices = [selected_index] if selected_index is not None else list(range(len(rows)))
    output: list[dict[str, Any]] = []
    for idx in indices:
        if idx is None or idx < 0 or idx >= len(rows):
            continue
        window = feedback_window_for_row(rows, idx, tolerance)
        row = window["row"]
        prompt_payload = ai_feedback_prompt_payload(payload, window, rule_issues, depth)
        feedback = openai_generate_ai_feedback(model, api_key, prompt_payload, depth)
        issues = feedback.get("issues", [])
        if not isinstance(issues, list):
            issues = []
        strengths = feedback.get("strengths", [])
        if not isinstance(strengths, list):
            strengths = [str(strengths)] if strengths else []
        output.append({
            "item_label": row.get("item_label", f"Item {idx+1}"),
            "review_mode": "simultaneous soft-window" if is_simultaneous_feedback_rows(rows) else "consecutive strict segment",
            "context_policy": window.get("context_policy", ""),
            "model": model,
            "depth": depth,
            "overall_confidence": feedback.get("overall_confidence", ""),
            "summary": feedback.get("summary", ""),
            "strengths": json.dumps(strengths, ensure_ascii=False),
            "issues": json.dumps(issues, ensure_ascii=False),
            "suggested_drill": feedback.get("suggested_drill", ""),
            "source_text": row.get("source_text", ""),
            "reference_text": row.get("reference_text", ""),
            "your_transcript": get_row_transcript_for_feedback(row),
            "_issues_list": issues,
            "_strengths_list": strengths,
        })
    return output


def ai_feedback_rows_to_csv(rows: list[dict[str, Any]]) -> str:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=AI_FEEDBACK_CSV_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        writer.writerow({k: row.get(k, "") for k in AI_FEEDBACK_CSV_FIELDS})
    return output.getvalue()


def ai_feedback_rows_to_txt(rows: list[dict[str, Any]]) -> str:
    parts = ["AI-Assisted Interpreter Practice Feedback", f"Generated: {datetime.now().isoformat(timespec='seconds')}", "Not an official exam score.", ""]
    for row in rows:
        parts.append(f"{row.get('item_label', 'Item')} — {row.get('review_mode', '')}")
        parts.append(f"Model: {row.get('model', '')} | Depth: {row.get('depth', '')} | Confidence: {row.get('overall_confidence', '')}")
        parts.append(f"Context policy: {row.get('context_policy', '')}")
        parts.append(f"Summary: {row.get('summary', '')}")
        strengths = row.get("_strengths_list") or []
        if strengths:
            parts.append("Strengths:")
            parts.extend([f"- {s}" for s in strengths])
        issues = row.get("_issues_list") or []
        if issues:
            parts.append("Issues / review items:")
            for issue in issues:
                if isinstance(issue, dict):
                    parts.append(f"- {issue.get('severity', '')} / {issue.get('confidence', '')} — {issue.get('type', '')}: {issue.get('comment', '')} Review: {issue.get('suggested_review', '')}")
                else:
                    parts.append(f"- {issue}")
        else:
            parts.append("Issues / review items: none identified.")
        if row.get("suggested_drill"):
            parts.append(f"Suggested drill: {row.get('suggested_drill')}")
        parts.append("\nSOURCE:\n" + str(row.get("source_text", "") or ""))
        parts.append("\nREFERENCE:\n" + str(row.get("reference_text", "") or ""))
        parts.append("\nYOUR TRANSCRIPT:\n" + str(row.get("your_transcript", "") or ""))
        parts.append("\n" + "-" * 80 + "\n")
    return "\n".join(parts).strip() + "\n"


def build_ai_feedback_docx(rows: list[dict[str, Any]], out_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise RuntimeError("python-docx is not installed. Run: python -m pip install python-docx") from exc
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    with contextlib.suppress(Exception):
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(9)
    title = doc.add_paragraph()
    run = title.add_run("AI-Assisted Interpreter Practice Feedback")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("Not an official exam score. Use this as a study aid.")
    for row in rows:
        p = doc.add_paragraph()
        h = p.add_run(str(row.get("item_label", "Item")))
        h.bold = True
        h.font.size = Pt(12)
        doc.add_paragraph(f"Mode: {row.get('review_mode', '')} | Model: {row.get('model', '')} | Confidence: {row.get('overall_confidence', '')}")
        doc.add_paragraph(f"Context policy: {row.get('context_policy', '')}")
        doc.add_paragraph(f"Summary: {row.get('summary', '')}")
        strengths = row.get("_strengths_list") or []
        if strengths:
            doc.add_paragraph("Strengths:")
            for s in strengths:
                doc.add_paragraph(str(s), style=None)
        issues = row.get("_issues_list") or []
        doc.add_paragraph("Issues / review items:")
        if issues:
            for issue in issues:
                if isinstance(issue, dict):
                    doc.add_paragraph(f"{issue.get('severity', '')} / {issue.get('confidence', '')} — {issue.get('type', '')}: {issue.get('comment', '')} Suggested review: {issue.get('suggested_review', '')}")
                else:
                    doc.add_paragraph(str(issue))
        else:
            doc.add_paragraph("None identified.")
        if row.get("suggested_drill"):
            doc.add_paragraph(f"Suggested drill: {row.get('suggested_drill')}")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Source"
        hdr[1].text = "Reference"
        hdr[2].text = "Your transcript"
        cells = table.add_row().cells
        cells[0].text = str(row.get("source_text", "") or "")
        cells[1].text = str(row.get("reference_text", "") or "")
        cells[2].text = str(row.get("your_transcript", "") or "")
        doc.add_paragraph()
    doc.save(out_path)
    return out_path


def build_ai_feedback_outputs(payload: dict[str, Any], rows: list[dict[str, Any]], rule_issues: list[dict[str, Any]], script_id: str, session_id: str, out_dir: Path, model: str, api_key: str, depth: str, tolerance: str, selected_index: int | None = None) -> tuple[list[dict[str, Any]], Path, Path, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    feedback_rows = build_ai_feedback_rows(payload, rows, rule_issues, model, api_key, depth, tolerance, selected_index=selected_index)
    scope = "selected" if selected_index is not None else "all"
    base = f"{sanitize_filename_part(script_id)}_{session_id}_ai_feedback_{scope}"
    csv_path = out_dir / f"{base}.csv"
    txt_path = out_dir / f"{base}.txt"
    docx_path = out_dir / f"{base}.docx"
    csv_path.write_text(ai_feedback_rows_to_csv(feedback_rows), encoding="utf-8")
    txt_path.write_text(ai_feedback_rows_to_txt(feedback_rows), encoding="utf-8")
    build_ai_feedback_docx(feedback_rows, docx_path)
    return feedback_rows, txt_path, csv_path, docx_path


def render_ai_feedback_summary(rows: list[dict[str, Any]]) -> None:
    if not rows:
        return
    st.markdown("#### AI feedback review")
    labels = [str(r.get("item_label", f"Item {i+1}")) for i, r in enumerate(rows)]
    idx = st.selectbox("Review AI feedback item", list(range(len(rows))), format_func=lambda i: labels[i], key="ai_feedback_review_item")
    row = rows[int(idx)]
    st.markdown(compact_metadata_row([row.get("review_mode", ""), row.get("context_policy", ""), f"Confidence: {row.get('overall_confidence', '')}", row.get("model", "")]), unsafe_allow_html=True)
    st.write(str(row.get("summary", "") or ""))
    issues = row.get("_issues_list") or []
    if issues:
        st.dataframe(issues, use_container_width=True, hide_index=True)
    else:
        st.success("No AI issue items returned for this segment/chunk.")
    if row.get("suggested_drill"):
        st.info(str(row.get("suggested_drill")))


# -----------------------------
# Phase 3C.2B full exam summary + study plan helpers
# -----------------------------
AI_EXAM_SUMMARY_CSV_FIELDS = [
    "model", "depth", "overall_confidence", "overall_summary", "overall_strengths",
    "repeated_patterns", "top_review_priorities", "legal_terms_to_review", "numbers_dates_names_notes",
    "omission_lag_notes", "wpm_notes", "suggested_drills", "suggested_wpm_progression", "terms_to_add_to_srs",
]


def summarize_issue_counts(rule_issues: list[dict[str, Any]], ai_feedback_rows: list[dict[str, Any]]) -> dict[str, Any]:
    rule_by_cat: dict[str, int] = {}
    rule_by_severity: dict[str, int] = {}
    for issue in rule_issues or []:
        cat = str(issue.get("category", "Other") or "Other")
        sev = str(issue.get("severity", "Unknown") or "Unknown")
        rule_by_cat[cat] = rule_by_cat.get(cat, 0) + 1
        rule_by_severity[sev] = rule_by_severity.get(sev, 0) + 1
    ai_by_type: dict[str, int] = {}
    ai_by_severity: dict[str, int] = {}
    for row in ai_feedback_rows or []:
        issues = row.get("_issues_list")
        if not isinstance(issues, list):
            with contextlib.suppress(Exception):
                issues = json.loads(str(row.get("issues", "[]") or "[]"))
        if not isinstance(issues, list):
            issues = []
        for issue in issues:
            if isinstance(issue, dict):
                typ = str(issue.get("type", "Other") or "Other")
                sev = str(issue.get("severity", "Unknown") or "Unknown")
            else:
                typ = "Other"
                sev = "Unknown"
            ai_by_type[typ] = ai_by_type.get(typ, 0) + 1
            ai_by_severity[sev] = ai_by_severity.get(sev, 0) + 1
    return {
        "rule_issue_count": len(rule_issues or []),
        "rule_by_category": rule_by_cat,
        "rule_by_severity": rule_by_severity,
        "ai_feedback_item_count": len(ai_feedback_rows or []),
        "ai_issue_count": sum(ai_by_type.values()),
        "ai_by_type": ai_by_type,
        "ai_by_severity": ai_by_severity,
    }


def wpm_summary_from_rows(rows: list[dict[str, Any]]) -> dict[str, Any]:
    values: list[float] = []
    for row in rows or []:
        for key in ["effective_wpm", "wpm", "source_wpm"]:
            val = row.get(key, "")
            with contextlib.suppress(Exception):
                f = float(str(val).replace("wpm", "").strip())
                if f > 0:
                    values.append(f)
                    break
    if not values:
        return {"count": 0, "average_wpm": "", "min_wpm": "", "max_wpm": "", "target_wpm": 140}
    avg = sum(values) / len(values)
    return {
        "count": len(values),
        "average_wpm": round(avg, 1),
        "min_wpm": round(min(values), 1),
        "max_wpm": round(max(values), 1),
        "target_wpm": 140,
        "distance_from_target": round(avg - 140, 1),
    }


def compact_ai_feedback_for_exam_summary(ai_feedback_rows: list[dict[str, Any]], limit: int = 60) -> list[dict[str, Any]]:
    compact = []
    for row in (ai_feedback_rows or [])[:limit]:
        issues = row.get("_issues_list")
        if not isinstance(issues, list):
            with contextlib.suppress(Exception):
                issues = json.loads(str(row.get("issues", "[]") or "[]"))
        if not isinstance(issues, list):
            issues = []
        compact.append({
            "item_label": row.get("item_label", ""),
            "confidence": row.get("overall_confidence", ""),
            "summary": row.get("summary", ""),
            "suggested_drill": row.get("suggested_drill", ""),
            "issues": issues[:8],
        })
    return compact


def ai_exam_summary_prompt_payload(
    payload: dict[str, Any],
    rows: list[dict[str, Any]],
    rule_issues: list[dict[str, Any]],
    ai_feedback_rows: list[dict[str, Any]],
    depth: str,
) -> str:
    terms = payload.get("terms_used") or []
    compact_terms = []
    for term in terms[:100]:
        if isinstance(term, dict):
            compact_terms.append({
                "english": term.get("english", ""),
                "spanish_contextual_choice": term.get("spanish_contextual_choice", ""),
            })
    sample_rows = []
    for row in (rows or [])[:80]:
        sample_rows.append({
            "item_label": row.get("item_label", ""),
            "source_language": row.get("source_language", ""),
            "target_language": row.get("target_language", ""),
            "speaker": row.get("speaker", ""),
            "segment_type": row.get("segment_type", ""),
            "wpm": row.get("effective_wpm", row.get("wpm", "")),
            "source_text": str(row.get("source_text", "") or "")[:900],
            "reference_text": str(row.get("reference_text", "") or "")[:900],
            "your_transcript": str(get_row_transcript_for_feedback(row) or "")[:900],
        })
    compact_rule_issues = [
        {k: issue.get(k, "") for k in ["severity", "category", "item_label", "expected", "found", "details"]}
        for issue in (rule_issues or [])[:120]
    ]
    data = {
        "review_purpose": "Full-exam interpreter practice summary and study plan only. Not an official score, certification decision, or pass/fail result.",
        "depth": depth,
        "script": {
            "script_id": payload.get("script_id", ""),
            "title": payload.get("title", ""),
            "format": payload.get("format", ""),
            "category": payload.get("category", ""),
        },
        "important_instruction": "For simultaneous mode, timing-aligned chunks may drift. Do not over-count possible omissions caused only by lag/lead. Look for repeated patterns and high-confidence issues across neighboring context, not isolated timestamp mismatches.",
        "issue_counts": summarize_issue_counts(rule_issues, ai_feedback_rows),
        "wpm_summary": wpm_summary_from_rows(rows),
        "terms_used_reference": compact_terms,
        "rule_based_review_items": compact_rule_issues,
        "ai_segment_feedback": compact_ai_feedback_for_exam_summary(ai_feedback_rows),
        "aligned_exam_rows_sample": sample_rows,
        "return_format": {
            "overall_summary": "paragraph summary of performance",
            "overall_confidence": "High/Medium/Low confidence in this study summary",
            "overall_strengths": ["..."],
            "repeated_patterns": [{"pattern": "...", "evidence": "...", "priority": "High/Medium/Low"}],
            "top_review_priorities": [{"priority": "...", "why_it_matters": "...", "practice_action": "..."}],
            "legal_terms_to_review": [{"term": "...", "recommended_rendering": "...", "reason": "..."}],
            "numbers_dates_names_notes": ["..."],
            "omission_lag_notes": ["..."],
            "wpm_notes": "comment on practice speed compared with 140 WPM exam target",
            "suggested_drills": [{"drill": "...", "duration": "...", "instructions": "..."}],
            "suggested_wpm_progression": ["..."],
            "terms_to_add_to_srs": ["..."],
        },
    }
    return json.dumps(data, ensure_ascii=False, indent=2)


def openai_generate_exam_summary(model: str, api_key: str, prompt_payload: str, depth: str) -> dict[str, Any]:
    if not api_key:
        raise ValueError("Missing OpenAI API key. Enter one in the app or set OPENAI_API_KEY in your environment.")
    try:
        from openai import OpenAI
    except Exception as exc:
        raise RuntimeError("The openai package is not installed. Run: python -m pip install openai") from exc
    client = OpenAI(api_key=api_key)
    instructions = (
        "You are an interpreter-training coach. Create a full-exam study summary from structured transcript rows, soft automated rule-based review prompts, and per-segment AI feedback. "
        "Do not issue an official grade or pass/fail result. Be conservative with simultaneous omissions because timing alignment may drift. "
        "Focus on repeated patterns, practical drills, legal terminology, numbers/dates/names, omissions, lag management, and WPM progression. Treat rule-based prompts as items to verify, not confirmed errors, because bilingual date formats, number words, ASR name spelling, and timing drift can create false positives. "
        "Return only valid JSON matching the requested structure."
    )
    text = ""
    try:
        response = client.responses.create(
            model=model,
            instructions=instructions,
            input=prompt_payload,
        )
        text = str(getattr(response, "output_text", "") or response_to_transcript_text(response)).strip()
    except Exception:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": instructions},
                {"role": "user", "content": prompt_payload},
            ],
            temperature=0.2 if str(depth).lower().startswith("detailed") else 0.1,
        )
        text = str(response.choices[0].message.content or "").strip()
    return parse_ai_feedback_json(text)




def ensure_list(value: Any) -> list[Any]:
    """Return value as a list for report/drill builders.

    This top-level helper is intentionally permissive because some AI outputs
    return strings where the app expects lists.
    """
    if isinstance(value, list):
        return value
    if value in (None, ""):
        return []
    if isinstance(value, tuple):
        return list(value)
    return [value]

def normalize_exam_summary(summary: dict[str, Any], model: str, depth: str) -> dict[str, Any]:
    def ensure_list(value: Any) -> list[Any]:
        if isinstance(value, list):
            return value
        if value in (None, ""):
            return []
        return [value]
    out = dict(summary or {})
    out["model"] = model
    out["depth"] = depth
    out["overall_summary"] = str(out.get("overall_summary", out.get("summary", "")) or "")
    out["overall_confidence"] = str(out.get("overall_confidence", "") or "")
    for key in ["overall_strengths", "repeated_patterns", "top_review_priorities", "legal_terms_to_review", "numbers_dates_names_notes", "omission_lag_notes", "suggested_drills", "suggested_wpm_progression", "terms_to_add_to_srs"]:
        out[key] = ensure_list(out.get(key, []))
    out["wpm_notes"] = str(out.get("wpm_notes", "") or "")
    return out


def ai_exam_summary_to_csv(summary: dict[str, Any]) -> str:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=AI_EXAM_SUMMARY_CSV_FIELDS, extrasaction="ignore")
    writer.writeheader()
    row = {}
    for field in AI_EXAM_SUMMARY_CSV_FIELDS:
        val = summary.get(field, "")
        if isinstance(val, (list, dict)):
            row[field] = json.dumps(val, ensure_ascii=False)
        else:
            row[field] = val
    writer.writerow(row)
    return output.getvalue()


def ai_exam_summary_to_txt(summary: dict[str, Any]) -> str:
    parts = ["Full AI Exam Summary + Study Plan", f"Generated: {datetime.now().isoformat(timespec='seconds')}", "Not an official exam score.", ""]
    parts.append(f"Model: {summary.get('model', '')} | Depth: {summary.get('depth', '')} | Confidence: {summary.get('overall_confidence', '')}")
    parts.append("\nOverall summary:\n" + str(summary.get("overall_summary", "") or ""))
    def add_list(title: str, items: list[Any]) -> None:
        parts.append(f"\n{title}:")
        if not items:
            parts.append("- None listed.")
            return
        for item in items:
            if isinstance(item, dict):
                pieces = [f"{k}: {v}" for k, v in item.items() if v not in (None, "")]
                parts.append("- " + "; ".join(pieces))
            else:
                parts.append("- " + str(item))
    add_list("Overall strengths", summary.get("overall_strengths", []))
    add_list("Repeated patterns", summary.get("repeated_patterns", []))
    add_list("Top review priorities", summary.get("top_review_priorities", []))
    add_list("Legal terms to review", summary.get("legal_terms_to_review", []))
    add_list("Numbers / dates / names notes", summary.get("numbers_dates_names_notes", []))
    add_list("Omission / lag notes", summary.get("omission_lag_notes", []))
    parts.append("\nWPM notes:\n" + str(summary.get("wpm_notes", "") or ""))
    add_list("Suggested drills", summary.get("suggested_drills", []))
    add_list("Suggested WPM progression", summary.get("suggested_wpm_progression", []))
    add_list("Terms to add to SRS", summary.get("terms_to_add_to_srs", []))
    return "\n".join(parts).strip() + "\n"


def build_ai_exam_summary_docx(summary: dict[str, Any], out_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise RuntimeError("python-docx is not installed. Run: python -m pip install python-docx") from exc
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    with contextlib.suppress(Exception):
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)
    title = doc.add_paragraph()
    r = title.add_run("Full AI Exam Summary + Study Plan")
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph("Not an official exam score. Use this as a study plan.")
    doc.add_paragraph(f"Model: {summary.get('model', '')} | Depth: {summary.get('depth', '')} | Confidence: {summary.get('overall_confidence', '')}")
    doc.add_heading("Overall summary", level=1)
    doc.add_paragraph(str(summary.get("overall_summary", "") or ""))
    def add_section(title: str, items: list[Any]) -> None:
        doc.add_heading(title, level=1)
        if not items:
            doc.add_paragraph("None listed.")
            return
        for item in items:
            if isinstance(item, dict):
                pieces = [f"{k}: {v}" for k, v in item.items() if v not in (None, "")]
                doc.add_paragraph("; ".join(pieces))
            else:
                doc.add_paragraph(str(item))
    add_section("Overall strengths", summary.get("overall_strengths", []))
    add_section("Repeated patterns", summary.get("repeated_patterns", []))
    add_section("Top review priorities", summary.get("top_review_priorities", []))
    add_section("Legal terms to review", summary.get("legal_terms_to_review", []))
    add_section("Numbers / dates / names notes", summary.get("numbers_dates_names_notes", []))
    add_section("Omission / lag notes", summary.get("omission_lag_notes", []))
    doc.add_heading("WPM notes", level=1)
    doc.add_paragraph(str(summary.get("wpm_notes", "") or ""))
    add_section("Suggested drills", summary.get("suggested_drills", []))
    add_section("Suggested WPM progression", summary.get("suggested_wpm_progression", []))
    add_section("Terms to add to SRS", summary.get("terms_to_add_to_srs", []))
    doc.save(out_path)
    return out_path


def build_ai_exam_summary_outputs(payload: dict[str, Any], rows: list[dict[str, Any]], rule_issues: list[dict[str, Any]], ai_feedback_rows: list[dict[str, Any]], script_id: str, session_id: str, out_dir: Path, model: str, api_key: str, depth: str) -> tuple[dict[str, Any], Path, Path, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    prompt_payload = ai_exam_summary_prompt_payload(payload, rows, rule_issues, ai_feedback_rows, depth)
    raw_summary = openai_generate_exam_summary(model, api_key, prompt_payload, depth)
    summary = normalize_exam_summary(raw_summary, model, depth)
    base = f"{sanitize_filename_part(script_id)}_{session_id}_ai_exam_summary"
    csv_path = out_dir / f"{base}.csv"
    txt_path = out_dir / f"{base}.txt"
    docx_path = out_dir / f"{base}.docx"
    csv_path.write_text(ai_exam_summary_to_csv(summary), encoding="utf-8")
    txt_path.write_text(ai_exam_summary_to_txt(summary), encoding="utf-8")
    build_ai_exam_summary_docx(summary, docx_path)
    return summary, txt_path, csv_path, docx_path


def render_ai_exam_summary(summary: dict[str, Any]) -> None:
    if not summary:
        return
    st.markdown("#### Full AI exam summary review")
    st.markdown(compact_metadata_row([summary.get("model", ""), summary.get("depth", ""), f"Confidence: {summary.get('overall_confidence', '')}", "Study plan, not official score"]), unsafe_allow_html=True)
    st.write(str(summary.get("overall_summary", "") or ""))
    cols = st.columns(3)
    cols[0].metric("Strengths", len(summary.get("overall_strengths", []) or []))
    cols[1].metric("Review priorities", len(summary.get("top_review_priorities", []) or []))
    cols[2].metric("Suggested drills", len(summary.get("suggested_drills", []) or []))
    with st.expander("Top review priorities", expanded=True):
        items = summary.get("top_review_priorities", []) or []
        if items:
            st.dataframe(items, use_container_width=True, hide_index=True)
        else:
            st.write("None listed.")
    with st.expander("Suggested drills and WPM progression", expanded=True):
        drills = summary.get("suggested_drills", []) or []
        if drills:
            st.dataframe(drills, use_container_width=True, hide_index=True)
        if summary.get("wpm_notes"):
            st.info(str(summary.get("wpm_notes")))
        progression = summary.get("suggested_wpm_progression", []) or []
        if progression:
            st.write("WPM progression:")
            for item in progression:
                st.write(f"- {item}")
    with st.expander("Legal terms / SRS candidates", expanded=False):
        terms = summary.get("legal_terms_to_review", []) or []
        if terms:
            st.dataframe(terms, use_container_width=True, hide_index=True)
        srs_terms = summary.get("terms_to_add_to_srs", []) or []
        if srs_terms:
            st.write("Terms to add to SRS:")
            st.write(", ".join(str(t) for t in srs_terms))


# -----------------------------
# Phase 3C.2C Personalized Drill Builder helpers
# -----------------------------
PERSONALIZED_DRILL_FIELDS = [
    "priority", "drill_type", "source", "item_label", "focus", "direction",
    "prompt", "target", "instructions", "suggested_reps", "wpm_target",
    "related_terms", "evidence", "source_text", "reference_text", "your_transcript",
]


def _as_text_list(value: Any) -> list[str]:
    if value in (None, ""):
        return []
    if isinstance(value, list):
        out = []
        for item in value:
            if isinstance(item, dict):
                pieces = [str(v) for v in item.values() if str(v).strip()]
                out.append("; ".join(pieces))
            else:
                out.append(str(item))
        return [x.strip() for x in out if x and x.strip()]
    if isinstance(value, dict):
        return ["; ".join(str(v) for v in value.values() if str(v).strip())]
    raw = str(value).strip()
    if not raw:
        return []
    # Split only obvious list-ish text, not every comma in legal phrases.
    parts = re.split(r"\n+|\s*;\s*", raw)
    return [x.strip(" -•\t") for x in parts if x.strip(" -•\t")]


def _term_display(term: dict[str, Any]) -> str:
    eng = str(term.get("english", "") or "").strip()
    spa = str(term.get("spanish", term.get("spanish_contextual_choice", "")) or "").strip()
    if eng and spa:
        return f"{eng} → {spa}"
    return eng or spa


def terms_used_lookup(payload: dict[str, Any]) -> list[dict[str, Any]]:
    terms = []
    for term in payload.get("terms_used") or []:
        if not isinstance(term, dict):
            continue
        eng = str(term.get("english", "") or "").strip()
        spa = str(term.get("spanish_contextual_choice", term.get("spanish", "")) or "").strip()
        terms.append({
            "english": eng,
            "spanish": spa,
            "spanish_contextual_choice": spa,
            "script_id": payload.get("script_id", ""),
            "title": payload.get("title", ""),
            "source_file": payload.get("source_file", ""),
        })
    return terms


def match_terms_from_text(payload: dict[str, Any], text_value: str) -> list[dict[str, Any]]:
    hay = normalize_for_matching(text_value)
    matched = []
    for term in terms_used_lookup(payload):
        candidates = [term.get("english", ""), term.get("spanish", "")]
        if any(c and normalize_for_matching(c) in hay for c in candidates):
            matched.append(term)
    return matched


def make_drill(priority: int, drill_type: str, source: str, item_label: str, focus: str,
               direction: str, prompt: str, target: str, instructions: str,
               suggested_reps: str = "3", wpm_target: str = "", related_terms: str = "",
               evidence: str = "", source_text: str = "", reference_text: str = "", your_transcript: str = "") -> dict[str, Any]:
    return {
        "priority": str(priority),
        "drill_type": drill_type,
        "source": source,
        "item_label": item_label,
        "focus": focus,
        "direction": direction,
        "prompt": prompt,
        "target": target,
        "instructions": instructions,
        "suggested_reps": suggested_reps,
        "wpm_target": wpm_target,
        "related_terms": related_terms,
        "evidence": evidence,
        "source_text": source_text,
        "reference_text": reference_text,
        "your_transcript": your_transcript,
    }


def drill_dedupe_key(drill: dict[str, Any]) -> str:
    return short_hash(
        normalize_for_matching(drill.get("drill_type", "")),
        normalize_for_matching(drill.get("focus", "")),
        normalize_for_matching(drill.get("prompt", ""))[:200],
        normalize_for_matching(drill.get("target", ""))[:200],
    )


def build_personalized_drills(payload: dict[str, Any], rows: list[dict[str, Any]],
                              rule_issues: list[dict[str, Any]], ai_feedback_rows: list[dict[str, Any]],
                              ai_summary: dict[str, Any], target_wpm: int = 140,
                              include_rule_flags: bool = True, include_ai_feedback: bool = True,
                              include_summary: bool = True, include_terms: bool = True,
                              max_drills: int = 60) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    drills: list[dict[str, Any]] = []
    row_by_label = {str(r.get("item_label", "") or ""): r for r in rows or []}
    terms_to_srs: list[dict[str, Any]] = []

    def add_terms_from_text(text_value: str) -> str:
        found = match_terms_from_text(payload, text_value)
        for term in found:
            terms_to_srs.append(term)
        return "; ".join(_term_display(t) for t in found[:6])

    if include_rule_flags:
        for issue in rule_issues or []:
            category = str(issue.get("category", "Review") or "Review")
            expected = str(issue.get("expected", "") or "").strip()
            label = str(issue.get("item_label", "") or "")
            row = row_by_label.get(label, {})
            sev = str(issue.get("severity", "Medium") or "Medium").lower()
            priority = {"high": 1, "medium": 2, "low": 3}.get(sev, 2)
            details = str(issue.get("details", "") or "")
            source_text = str(row.get("source_text", issue.get("source_text", "")) or "")
            reference_text = str(row.get("reference_text", issue.get("reference_text", "")) or "")
            transcript = get_row_transcript_for_feedback(row) if row else str(issue.get("your_transcript", "") or "")
            related_terms = add_terms_from_text(" ".join([expected, details, source_text, reference_text]))
            if "legal term" in category.lower():
                drills.append(make_drill(
                    priority, "Legal term retrieval", "Rule-based review", label, expected or category,
                    "Bidirectional", expected or details, related_terms or "Preferred term from JSON terms_used",
                    "Hear or read the prompt, produce the preferred equivalent aloud, then use it in one full legal sentence.",
                    "5", related_terms=related_terms, evidence=details,
                    source_text=source_text, reference_text=reference_text, your_transcript=transcript,
                ))
            elif any(x in category.lower() for x in ["number", "date", "name", "case"]):
                drills.append(make_drill(
                    priority, "Critical fact drill", "Rule-based review", label, expected or category,
                    "Source → target", source_text or expected, reference_text or expected,
                    "Repeat the source once, interpret aloud, and verify every number/date/name/case number. Treat this as a manual review prompt, not a confirmed error.",
                    "4", related_terms=related_terms, evidence=details,
                    source_text=source_text, reference_text=reference_text, your_transcript=transcript,
                ))
            else:
                drills.append(make_drill(
                    priority, "Omission repair", "Rule-based review", label, expected or category,
                    "Source → target", source_text, reference_text,
                    "Interpret this item again, then compare against the reference and your transcript. Focus on preserving every legal idea and concrete fact.",
                    "3", related_terms=related_terms, evidence=details,
                    source_text=source_text, reference_text=reference_text, your_transcript=transcript,
                ))

    if include_ai_feedback:
        for fb in ai_feedback_rows or []:
            label = str(fb.get("item_label", "") or "")
            row = row_by_label.get(label, {})
            source_text = str(fb.get("source_text", row.get("source_text", "")) or "")
            reference_text = str(fb.get("reference_text", row.get("reference_text", "")) or "")
            transcript = str(fb.get("your_transcript", get_row_transcript_for_feedback(row) if row else "") or "")
            suggested = str(fb.get("suggested_drill", "") or "").strip()
            summary = str(fb.get("summary", "") or "").strip()
            issues_raw = fb.get("issues", "")
            issues_text = issues_raw if isinstance(issues_raw, str) else json.dumps(issues_raw, ensure_ascii=False)
            related_terms = add_terms_from_text(" ".join([summary, issues_text, suggested, source_text, reference_text]))
            focus = summary[:180] if summary else label or "AI feedback item"
            drills.append(make_drill(
                2, "AI feedback repair", "AI per-item feedback", label, focus,
                "Source → target", source_text, reference_text,
                suggested or "Redo this item aloud, then compare your response against the reference and the AI study feedback.",
                "3", related_terms=related_terms, evidence=issues_text[:800],
                source_text=source_text, reference_text=reference_text, your_transcript=transcript,
            ))

    if include_summary and ai_summary:
        for item in ai_summary.get("suggested_drills", []) or []:
            text_item = item if isinstance(item, str) else "; ".join(f"{k}: {v}" for k, v in item.items() if v not in (None, ""))
            related_terms = add_terms_from_text(str(text_item))
            drills.append(make_drill(
                1, "Study plan drill", "Full AI summary", "Whole exam", str(text_item)[:180],
                "Mixed", str(text_item), "", "Complete this drill as written in the AI study plan, then log the result in practice history.",
                "1", related_terms=related_terms, evidence=str(text_item),
            ))
        for item in ai_summary.get("suggested_wpm_progression", []) or []:
            text_item = str(item)
            drills.append(make_drill(
                2, "WPM progression", "Full AI summary", "Whole exam", "Speed building", "Source-only",
                f"Repeat the source-only track at staged speeds toward {target_wpm} WPM.", "",
                text_item or f"Practice once below {target_wpm} WPM, once near target, and increase only when critical facts remain stable.",
                "1", wpm_target=str(target_wpm), evidence=text_item,
            ))
        for t in ai_summary.get("terms_to_add_to_srs", []) or []:
            matched = match_terms_from_text(payload, str(t))
            if matched:
                terms_to_srs.extend(matched)
            else:
                drills.append(make_drill(
                    2, "Term review", "Full AI summary", "Whole exam", str(t), "Bidirectional",
                    str(t), "", "Add this term to your personal review list and practice both directions.", "5", evidence=str(t),
                ))

    if include_terms:
        # Add a few high-yield script terms that appear in feedback contexts or are generally important.
        feedback_blob = " ".join(
            [json.dumps(rule_issues or [], ensure_ascii=False), json.dumps(ai_feedback_rows or [], ensure_ascii=False), json.dumps(ai_summary or {}, ensure_ascii=False)]
        )
        matched_terms = match_terms_from_text(payload, feedback_blob)
        if not matched_terms and not drills:
            matched_terms = terms_used_lookup(payload)[:10]
        for term in matched_terms[:12]:
            terms_to_srs.append(term)
            drills.append(make_drill(
                3, "Term SRS card", "terms_used", "Whole script", _term_display(term), "Bidirectional",
                term.get("english", ""), term.get("spanish", ""),
                "Practice English → Spanish, then Spanish → English. Mark as Hard/Missed in Term Review if retrieval is slow.",
                "5", related_terms=_term_display(term), evidence="From JSON terms_used.",
            ))

    # Dedupe and rank.
    deduped: dict[str, dict[str, Any]] = {}
    for drill in drills:
        key = drill_dedupe_key(drill)
        if key not in deduped:
            deduped[key] = drill
        else:
            old = deduped[key]
            if int(drill.get("priority", 9) or 9) < int(old.get("priority", 9) or 9):
                deduped[key] = drill
    final_drills = list(deduped.values())
    final_drills.sort(key=lambda d: (int(d.get("priority", 9) or 9), str(d.get("drill_type", "")), str(d.get("focus", ""))))
    final_drills = final_drills[:max(1, int(max_drills or 60))]

    # Dedupe terms for SRS.
    by_term_key = {}
    for term in terms_to_srs:
        key = term_srs_key(term)
        if key:
            by_term_key[key] = term
    return final_drills, list(by_term_key.values())


def personalized_drills_to_csv(drills: list[dict[str, Any]]) -> str:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=PERSONALIZED_DRILL_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for drill in drills:
        writer.writerow({k: drill.get(k, "") for k in PERSONALIZED_DRILL_FIELDS})
    return output.getvalue()


def personalized_drills_to_txt(drills: list[dict[str, Any]]) -> str:
    parts = ["Personalized Drill Plan", f"Generated: {datetime.now().isoformat(timespec='seconds')}", ""]
    if not drills:
        parts.append("No drill items were generated.")
        return "\n".join(parts).strip() + "\n"
    for i, drill in enumerate(drills, start=1):
        parts.append(f"{i}. [{drill.get('drill_type')}] {drill.get('focus')}")
        parts.append(f"   Priority: {drill.get('priority')} | Source: {drill.get('source')} | Item: {drill.get('item_label')}")
        if drill.get("direction"):
            parts.append(f"   Direction: {drill.get('direction')}")
        if drill.get("prompt"):
            parts.append(f"   Prompt: {drill.get('prompt')}")
        if drill.get("target"):
            parts.append(f"   Target/reference: {drill.get('target')}")
        if drill.get("instructions"):
            parts.append(f"   Instructions: {drill.get('instructions')}")
        if drill.get("related_terms"):
            parts.append(f"   Related terms: {drill.get('related_terms')}")
        parts.append("")
    return "\n".join(parts).strip() + "\n"


def build_personalized_drills_docx(drills: list[dict[str, Any]], out_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise RuntimeError("python-docx is not installed. Run: python -m pip install python-docx") from exc
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    with contextlib.suppress(Exception):
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)
    p = doc.add_paragraph()
    r = p.add_run("Personalized Drill Plan")
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph("Generated from Exam Mode rule-based prompts, AI feedback, summary items, and JSON terms_used. Use as a study guide, not an official score.")
    if not drills:
        doc.add_paragraph("No drill items were generated.")
    for i, drill in enumerate(drills, start=1):
        doc.add_heading(f"{i}. {drill.get('drill_type', '')}: {drill.get('focus', '')}", level=1)
        doc.add_paragraph(f"Priority: {drill.get('priority', '')} | Source: {drill.get('source', '')} | Item: {drill.get('item_label', '')}")
        doc.add_paragraph(f"Direction: {drill.get('direction', '')} | Suggested reps: {drill.get('suggested_reps', '')} | WPM target: {drill.get('wpm_target', '')}")
        if drill.get("prompt"):
            doc.add_paragraph("Prompt:")
            doc.add_paragraph(str(drill.get("prompt", "")))
        if drill.get("target"):
            doc.add_paragraph("Target / reference:")
            doc.add_paragraph(str(drill.get("target", "")))
        if drill.get("instructions"):
            doc.add_paragraph("Instructions:")
            doc.add_paragraph(str(drill.get("instructions", "")))
        if drill.get("related_terms"):
            doc.add_paragraph("Related terms: " + str(drill.get("related_terms", "")))
        if drill.get("evidence"):
            doc.add_paragraph("Why this drill was created:")
            doc.add_paragraph(str(drill.get("evidence", ""))[:1200])
    doc.save(out_path)
    return out_path


def build_personalized_drill_outputs(payload: dict[str, Any], rows: list[dict[str, Any]], rule_issues: list[dict[str, Any]],
                                     ai_feedback_rows: list[dict[str, Any]], ai_summary: dict[str, Any],
                                     script_id: str, session_id: str, out_dir: Path, target_wpm: int = 140,
                                     include_rule_flags: bool = True, include_ai_feedback: bool = True,
                                     include_summary: bool = True, include_terms: bool = True,
                                     max_drills: int = 60) -> tuple[list[dict[str, Any]], list[dict[str, Any]], Path, Path, Path, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    drills, srs_terms = build_personalized_drills(
        payload, rows, rule_issues, ai_feedback_rows, ai_summary, target_wpm,
        include_rule_flags, include_ai_feedback, include_summary, include_terms, max_drills,
    )
    base = f"{sanitize_filename_part(script_id)}_{session_id}_personalized_drills"
    csv_path = out_dir / f"{base}.csv"
    txt_path = out_dir / f"{base}.txt"
    docx_path = out_dir / f"{base}.docx"
    json_path = out_dir / f"{base}.json"
    csv_path.write_text(personalized_drills_to_csv(drills), encoding="utf-8")
    txt_path.write_text(personalized_drills_to_txt(drills), encoding="utf-8")
    build_personalized_drills_docx(drills, docx_path)
    save_interactive_drill_json(build_interactive_drill_json(payload, drills, srs_terms, script_id, session_id, target_wpm), json_path)
    return drills, srs_terms, txt_path, csv_path, docx_path, json_path


# -----------------------------
# Phase 3C.2E Rubric-based key-unit feedback + remedial JSON
# -----------------------------
RUBRIC_KEY_UNIT_FIELDS = [
    "priority", "item_label", "key_unit", "category", "status", "error_code", "point_impact",
    "source_text", "reference_text", "your_transcript", "recommended_drill", "evidence",
]

RUBRIC_ERROR_CODE_LABELS = {
    "O": "Omission / possible omission",
    "D": "Distortion / meaning shift",
    "L": "Legal terminology",
    "N": "Number / date / time / amount",
    "P": "Proper name / place / case number",
    "A": "Actor / relationship / who-did-what",
    "S": "Sequence / end-of-sentence completion",
    "R": "Register / style / false cognate",
    "T": "Timing / lag / delivery",
    "M": "Manual review prompt",
}


def rubric_error_code_from_text(category: str, text_value: str = "") -> str:
    blob = normalize_for_matching(f"{category} {text_value}")
    if any(x in blob for x in ["legal term", "terminology", "term", "charge", "burden", "reasonable doubt", "standard"]):
        return "L"
    if any(x in blob for x in ["number", "date", "time", "amount", "money", "dollar", "wpm", "age", "percent"]):
        return "N"
    if any(x in blob for x in ["name", "place", "street", "case number", "proper noun", "location"]):
        return "P"
    if any(x in blob for x in ["actor", "relationship", "who", "victim", "defendant", "prosecutor", "judge"]):
        return "A"
    if any(x in blob for x in ["omission", "missing", "not present", "dropped"]):
        return "O"
    if any(x in blob for x in ["distortion", "meaning shift", "incorrect", "wrong", "opposite", "mistranslation"]):
        return "D"
    if any(x in blob for x in ["sequence", "final clause", "unfinished", "end of sentence", "fade"]):
        return "S"
    if any(x in blob for x in ["register", "false cognate", "awkward", "style"]):
        return "R"
    if any(x in blob for x in ["lag", "timing", "speed", "behind", "lead"]):
        return "T"
    return "M"


def rubric_status_from_text(severity: str = "", text_value: str = "") -> tuple[str, str]:
    blob = normalize_for_matching(f"{severity} {text_value}")
    if any(x in blob for x in ["high", "omitted", "missing", "absent", "not rendered", "major"]):
        return "Possible omission/distortion", "High"
    if any(x in blob for x in ["partial", "medium", "weakened", "incomplete", "possible"]):
        return "Partial / needs review", "Medium"
    if any(x in blob for x in ["low", "spelling", "asr", "manual review"]):
        return "Manual review prompt", "Low"
    return "Needs review", "Medium"


def make_rubric_key_unit(priority: int, item_label: str, key_unit: str, category: str, status: str,
                         error_code: str, point_impact: str, source_text: str, reference_text: str,
                         your_transcript: str, recommended_drill: str, evidence: str) -> dict[str, Any]:
    return {
        "priority": str(priority),
        "item_label": item_label,
        "key_unit": key_unit,
        "category": category,
        "status": status,
        "error_code": error_code,
        "point_impact": point_impact,
        "source_text": source_text,
        "reference_text": reference_text,
        "your_transcript": your_transcript,
        "recommended_drill": recommended_drill,
        "evidence": evidence,
    }


def build_rubric_key_units(payload: dict[str, Any], rows: list[dict[str, Any]], rule_issues: list[dict[str, Any]],
                           ai_feedback_rows: list[dict[str, Any]], ai_summary: dict[str, Any],
                           max_units: int = 80) -> list[dict[str, Any]]:
    row_by_label = {str(r.get("item_label", "") or ""): r for r in rows or []}
    units: list[dict[str, Any]] = []

    for issue in rule_issues or []:
        label = str(issue.get("item_label", "") or "")
        row = row_by_label.get(label, {})
        category = str(issue.get("category", "Review prompt") or "Review prompt")
        expected = str(issue.get("expected", "") or issue.get("key_unit", "") or category).strip()
        details = str(issue.get("details", "") or issue.get("evidence", "") or "")
        status, impact = rubric_status_from_text(str(issue.get("severity", "") or ""), details)
        code = rubric_error_code_from_text(category, f"{expected} {details}")
        priority = {"High": 1, "Medium": 2, "Low": 3}.get(impact, 2)
        source_text = str(row.get("source_text", issue.get("source_text", "")) or "")
        reference_text = str(row.get("reference_text", issue.get("reference_text", "")) or "")
        transcript = get_row_transcript_for_feedback(row) if row else str(issue.get("your_transcript", "") or "")
        rec = "Redo this key unit aloud and verify that the message-bearing fact/term is preserved."
        if code == "L":
            rec = "Term automation: say the legal term both directions, then use it in the full sentence."
        elif code in {"N", "P"}:
            rec = "Critical fact recall: repeat and interpret the number/date/name/case detail without changing it."
        elif code == "O":
            rec = "Omission repair: reinterpret the source chunk and deliberately complete the missing idea."
        units.append(make_rubric_key_unit(priority, label, expected, category, status, code, impact, source_text, reference_text, transcript, rec, details))

    for fb in ai_feedback_rows or []:
        label = str(fb.get("item_label", "") or "")
        row = row_by_label.get(label, {})
        source_text = str(fb.get("source_text", row.get("source_text", "")) or "")
        reference_text = str(fb.get("reference_text", row.get("reference_text", "")) or "")
        transcript = str(fb.get("your_transcript", get_row_transcript_for_feedback(row) if row else "") or "")
        issues_raw = fb.get("issues", "")
        if isinstance(issues_raw, str):
            issue_items = _as_text_list(issues_raw) or ([issues_raw] if issues_raw.strip() else [])
        elif isinstance(issues_raw, list):
            issue_items = _as_text_list(issues_raw)
        else:
            issue_items = _as_text_list(issues_raw)
        summary = str(fb.get("summary", "") or "").strip()
        if not issue_items and summary:
            issue_items = [summary]
        for issue_text in issue_items[:4]:
            code = rubric_error_code_from_text("AI feedback", issue_text)
            status, impact = rubric_status_from_text("", issue_text)
            priority = {"High": 1, "Medium": 2, "Low": 3}.get(impact, 2)
            key_unit = str(issue_text).strip()[:220] or summary[:220] or label
            rec = str(fb.get("suggested_drill", "") or "Redo this source/reference unit and compare against the AI study feedback.")
            units.append(make_rubric_key_unit(priority, label, key_unit, "AI feedback", status, code, impact, source_text, reference_text, transcript, rec, str(issue_text)[:1000]))

    # Add high-value terms when no other units are available, or when terms appear in feedback text.
    feedback_blob = json.dumps(rule_issues or [], ensure_ascii=False) + " " + json.dumps(ai_feedback_rows or [], ensure_ascii=False) + " " + json.dumps(ai_summary or {}, ensure_ascii=False)
    matched_terms = match_terms_from_text(payload, feedback_blob)
    if not units:
        matched_terms = terms_used_lookup(payload)[:20]
    for term in matched_terms[:20]:
        ku = _term_display(term)
        units.append(make_rubric_key_unit(
            3, "Whole script", ku, "Legal terminology", "Drill candidate", "L", "Medium",
            str(term.get("english", "") or ""), str(term.get("spanish", "") or ""), "",
            "Term automation: practice English → Spanish, Spanish → English, then embed the term in a full legal sentence.",
            "From JSON terms_used and/or feedback context.",
        ))

    # Dedupe and cap.
    deduped: dict[str, dict[str, Any]] = {}
    for u in units:
        key = short_hash(u.get("item_label", ""), u.get("key_unit", ""), u.get("error_code", ""))
        if key not in deduped or int(u.get("priority", 9) or 9) < int(deduped[key].get("priority", 9) or 9):
            deduped[key] = u
    final = list(deduped.values())
    final.sort(key=lambda x: (int(x.get("priority", 9) or 9), str(x.get("error_code", "")), str(x.get("key_unit", ""))))
    return final[:max(1, int(max_units or 80))]


def rubric_key_units_to_csv(units: list[dict[str, Any]]) -> str:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=RUBRIC_KEY_UNIT_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for unit in units:
        writer.writerow({k: unit.get(k, "") for k in RUBRIC_KEY_UNIT_FIELDS})
    return output.getvalue()


def rubric_key_units_to_txt(units: list[dict[str, Any]]) -> str:
    parts = ["Rubric-Based Key Unit Review", f"Generated: {datetime.now().isoformat(timespec='seconds')}", ""]
    parts.append("Use these as review prompts, not official exam scoring. Key units focus on message-bearing facts, legal terms, names, dates, places, amounts, and meaning shifts.")
    parts.append("")
    for i, u in enumerate(units or [], start=1):
        code = u.get("error_code", "M")
        parts.append(f"{i}. [{code} - {RUBRIC_ERROR_CODE_LABELS.get(code, 'Review')}] {u.get('key_unit', '')}")
        parts.append(f"   Status: {u.get('status', '')} | Point impact: {u.get('point_impact', '')} | Item: {u.get('item_label', '')}")
        if u.get("recommended_drill"):
            parts.append(f"   Drill: {u.get('recommended_drill')}")
        if u.get("reference_text"):
            parts.append(f"   Reference: {str(u.get('reference_text'))[:350]}")
        if u.get("your_transcript"):
            parts.append(f"   Your transcript: {str(u.get('your_transcript'))[:350]}")
        parts.append("")
    return "\n".join(parts).strip() + "\n"


def rubric_units_to_drills(units: list[dict[str, Any]], target_wpm: int = 140) -> list[dict[str, Any]]:
    drills: list[dict[str, Any]] = []
    for u in units or []:
        code = str(u.get("error_code", "M") or "M")
        priority = int(u.get("priority", 2) or 2)
        key_unit = str(u.get("key_unit", "") or "").strip()
        source_text = str(u.get("source_text", "") or "").strip()
        reference_text = str(u.get("reference_text", "") or "").strip()
        transcript = str(u.get("your_transcript", "") or "").strip()
        label = str(u.get("item_label", "") or "")
        if code == "L":
            drill_type = "Key-unit legal term drill"
            instructions = "Practice the legal term both directions, then interpret the full source sentence without pausing."
            prompt = source_text or key_unit
            target = reference_text or key_unit
            reps = "5"
        elif code in {"N", "P"}:
            drill_type = "Key-unit critical fact drill"
            instructions = "Listen, repeat the critical fact, interpret it, and verify exact date/number/name/place preservation."
            prompt = source_text or key_unit
            target = reference_text or key_unit
            reps = "4"
        elif code == "O":
            drill_type = "Key-unit omission repair"
            instructions = "Redo the unit and deliberately complete the omitted idea. Then compare to the reference."
            prompt = source_text or key_unit
            target = reference_text
            reps = "3"
        else:
            drill_type = "Key-unit repair drill"
            instructions = str(u.get("recommended_drill", "") or "Redo this unit and preserve the message-bearing key unit.")
            prompt = source_text or key_unit
            target = reference_text
            reps = "3"
        drills.append(make_drill(
            priority, drill_type, "Rubric key-unit review", label, key_unit[:180], "Source → target",
            prompt, target, instructions, reps, wpm_target=str(target_wpm),
            related_terms=key_unit if code == "L" else "", evidence=str(u.get("evidence", "") or ""),
            source_text=source_text, reference_text=reference_text, your_transcript=transcript,
        ))
    return drills


def remedial_terms_from_drills(payload: dict[str, Any], drills: list[dict[str, Any]], limit: int = 40) -> list[dict[str, Any]]:
    blob = " ".join(json.dumps(d, ensure_ascii=False) for d in drills or [])
    terms = match_terms_from_text(payload, blob)
    if not terms:
        terms = terms_used_lookup(payload)[:min(limit, 20)]
    out = []
    seen = set()
    for term in terms:
        key = term_srs_key(term)
        if not key or key in seen:
            continue
        seen.add(key)
        out.append({
            "english": term.get("english", ""),
            "spanish_contextual_choice": term.get("spanish", term.get("spanish_contextual_choice", "")),
            "spanish_alternatives_from_sheet": [term.get("spanish", term.get("spanish_contextual_choice", ""))] if term.get("spanish", term.get("spanish_contextual_choice", "")) else [],
            "drill_note": "Selected from rubric/key-unit feedback for remedial practice.",
        })
        if len(out) >= limit:
            break
    return out


def build_remedial_practice_json(payload: dict[str, Any], units: list[dict[str, Any]], drills: list[dict[str, Any]],
                                 script_id: str, session_id: str, target_wpm: int = 140,
                                 max_segments: int = 30) -> dict[str, Any]:
    selected_drills = [d for d in drills or [] if str(d.get("prompt", "") or d.get("source_text", "")).strip()]
    selected_drills = selected_drills[:max(1, int(max_segments or 30))]
    paired_segments = []
    english_parts = []
    spanish_parts = []
    for idx, d in enumerate(selected_drills, start=1):
        prompt = str(d.get("source_text", "") or d.get("prompt", "") or "").strip()
        target = str(d.get("reference_text", "") or d.get("target", "") or "").strip()
        # If this is a term card, prefer direct EN/ES term pair.
        if not target and "→" in str(d.get("related_terms", "")):
            left, right = str(d.get("related_terms", "")).split("→", 1)
            prompt = prompt or left.strip()
            target = right.strip()
        if not prompt and not target:
            continue
        prompt_lang = infer_drill_language(prompt, "english")
        target_lang = infer_drill_language(target, "spanish" if prompt_lang == "english" else "english")
        if prompt_lang == "spanish" and target_lang == "english":
            english = target
            spanish = prompt
            source_language = "spanish"
            playback_order = ["spanish", "english"]
        else:
            english = prompt if prompt_lang == "english" else target
            spanish = target if target_lang == "spanish" else prompt
            source_language = "english"
            playback_order = ["english", "spanish"]
        english = english.strip() or prompt or target
        spanish = spanish.strip() or target or prompt
        english_parts.append(english)
        spanish_parts.append(spanish)
        paired_segments.append({
            "segment_number": idx,
            "segment_type": normalize_for_matching(d.get("drill_type", "remedial_drill")).replace(" ", "_")[:60] or "remedial_drill",
            "speaker": "DRILL COACH",
            "source_language": source_language,
            "playback_order": playback_order,
            "english": english,
            "spanish": spanish,
            "performance_focus": str(d.get("focus", "") or d.get("instructions", "") or "Key-unit remedial practice")[:500],
            "rubric_error_code": next((u.get("error_code", "") for u in units if normalize_for_matching(u.get("key_unit", "")) in normalize_for_matching(d.get("focus", "") + " " + d.get("prompt", ""))), ""),
        })
    audio_profile = payload.get("audio_profile") if isinstance(payload.get("audio_profile"), dict) else {}
    if not audio_profile:
        audio_profile = {
            "voice_mode": "speaker",
            "default_voices": {"english": "en-US-JennyNeural", "spanish": "es-MX-DaliaNeural"},
            "voice_assignments": {
                "DRILL COACH": {"english": "en-US-MichelleNeural", "spanish": "es-MX-DaliaNeural", "description": "remedial drill prompt voice"},
                "INTERPRETER": {"english": "en-US-AriaNeural", "spanish": "es-MX-DaliaNeural", "description": "reference interpretation voice"},
            },
            "speak_speaker_labels": False,
        }
    elif "DRILL COACH" not in (audio_profile.get("voice_assignments") or {}):
        ap = json.loads(json.dumps(audio_profile))
        ap.setdefault("voice_assignments", {})["DRILL COACH"] = {"english": "en-US-MichelleNeural", "spanish": "es-MX-DaliaNeural", "description": "remedial drill prompt voice"}
        audio_profile = ap
    title = f"Remedial Key-Unit Drill — {payload.get('title', script_id)}"
    perf_targets = []
    for u in units[:12]:
        code = u.get("error_code", "M")
        perf_targets.append(f"[{code}] {u.get('key_unit', '')}: {u.get('recommended_drill', '')}")
    return {
        "script_id": f"{sanitize_filename_part(script_id)}-REMEDIAL-{session_id}",
        "category": "Remedial Drill",
        "subcategory": str(payload.get("subcategory", "Rubric-Based Key Unit Practice") or "Rubric-Based Key Unit Practice"),
        "title": title,
        "format": "simultaneous_remedial_drill_with_flashcards",
        "jurisdiction_context": {
            "source_exam": script_id,
            "source_title": payload.get("title", ""),
            "purpose": "Rubric-based remedial practice built from key-unit review prompts, AI feedback, and exam transcript data.",
            "recommended_use": "Run flashcards first, then source-only or script audio. Interpret each source segment aloud before listening to the reference. Repeat until key units are automatic.",
            "target_wpm": target_wpm,
        },
        "audio_profile": audio_profile,
        "terms_used": remedial_terms_from_drills(payload, drills),
        "performance_targets": perf_targets,
        "rubric_key_units": units,
        "english_script": "\n".join(english_parts),
        "spanish_script": "\n".join(spanish_parts),
        "paired_segments": paired_segments,
    }


def save_rubric_and_remedial_outputs(payload: dict[str, Any], rows: list[dict[str, Any]], rule_issues: list[dict[str, Any]],
                                     ai_feedback_rows: list[dict[str, Any]], ai_summary: dict[str, Any],
                                     drills: list[dict[str, Any]], script_id: str, session_id: str,
                                     out_dir: Path, target_wpm: int = 140) -> tuple[list[dict[str, Any]], Path, Path, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    units = build_rubric_key_units(payload, rows, rule_issues, ai_feedback_rows, ai_summary, max_units=max(len(drills) * 2, 40))
    if not drills:
        drills.extend(rubric_units_to_drills(units, target_wpm=target_wpm))
    base = f"{sanitize_filename_part(script_id)}_{session_id}_rubric_key_units"
    rubric_csv = out_dir / f"{base}.csv"
    rubric_txt = out_dir / f"{base}.txt"
    remedial_json_path = out_dir / f"{sanitize_filename_part(script_id)}_{session_id}_remedial_drill_script.json"
    rubric_csv.write_text(rubric_key_units_to_csv(units), encoding="utf-8")
    rubric_txt.write_text(rubric_key_units_to_txt(units), encoding="utf-8")
    remedial_json = build_remedial_practice_json(payload, units, drills, script_id, session_id, target_wpm=target_wpm)
    remedial_json_path.write_text(json.dumps(remedial_json, ensure_ascii=False, indent=2), encoding="utf-8")
    return units, rubric_txt, rubric_csv, remedial_json_path


# -----------------------------
# Phase 3C.2F Rubric Score + Performance Summary helpers
# -----------------------------
RUBRIC_SCORE_CSV_FIELDS = [
    "generated_at", "script_id", "title", "estimated_score", "likely_level", "passing_target",
    "confidence", "meaning_key_unit_accuracy", "completeness", "legal_terminology",
    "delivery_grammar", "overall_summary", "main_strengths", "major_scoring_hits",
    "important_omissions", "practice_priorities", "bottom_line", "wpm_summary",
]


def _rubric_score_penalty(unit: dict[str, Any]) -> float:
    code = str(unit.get("error_code", "M") or "M").upper()[:1]
    priority = str(unit.get("priority", "Medium") or "Medium").lower()
    confidence = str(unit.get("confidence", "Medium") or "Medium").lower()
    base_by_code = {
        "O": 4.0,  # omission / possible omission
        "D": 4.0,  # distortion / meaning shift
        "A": 3.5,  # actor / relationship
        "N": 3.0,  # number / date / time / amount
        "P": 2.5,  # proper name / place / case number
        "L": 2.5,  # legal terminology
        "S": 2.0,  # sequence / completion
        "R": 1.5,  # register / false cognate
        "T": 1.0,  # timing / delivery
        "M": 1.0,  # manual review
    }
    penalty = base_by_code.get(code, 1.0)
    if "high" in priority:
        penalty *= 1.25
    elif "low" in priority:
        penalty *= 0.65
    if "low" in confidence or "manual" in confidence:
        penalty *= 0.55
    elif "high" in confidence:
        penalty *= 1.10
    return round(penalty, 2)


def _unit_text(unit: dict[str, Any]) -> str:
    pieces = [
        str(unit.get("key_unit", "") or "").strip(),
        str(unit.get("evidence", "") or "").strip(),
        str(unit.get("recommended_drill", "") or "").strip(),
    ]
    return " — ".join([x for x in pieces if x])


def _score_level(score: float, passing_target: int = 70) -> str:
    if score >= max(88, passing_target + 15):
        return "Strong practice performance"
    if score >= passing_target + 5:
        return "Likely passing-range practice performance"
    if score >= passing_target:
        return "Near/at passing-range practice performance"
    if score >= passing_target - 10:
        return "Developing / below passing target"
    return "Developing / needs focused remediation"


def build_rubric_score_summary(payload: dict[str, Any], rows: list[dict[str, Any]], units: list[dict[str, Any]],
                               rule_issues: list[dict[str, Any]], ai_feedback_rows: list[dict[str, Any]],
                               ai_summary: dict[str, Any], target_wpm: int = 140,
                               passing_target: int = 70) -> dict[str, Any]:
    units = units or build_rubric_key_units(payload, rows, rule_issues, ai_feedback_rows, ai_summary, max_units=80)
    score_buckets = {
        "meaning_key_unit_accuracy": 55.0,
        "completeness": 25.0,
        "legal_terminology": 12.0,
        "delivery_grammar": 8.0,
    }
    deductions = {k: 0.0 for k in score_buckets}
    for unit in units:
        code = str(unit.get("error_code", "M") or "M").upper()[:1]
        pen = _rubric_score_penalty(unit)
        if code in {"O", "D", "A"}:
            deductions["meaning_key_unit_accuracy"] += pen
            if code == "O":
                deductions["completeness"] += pen * 0.55
        elif code in {"N", "P", "S"}:
            deductions["meaning_key_unit_accuracy"] += pen * 0.70
            deductions["completeness"] += pen * 0.45
        elif code == "L":
            deductions["legal_terminology"] += pen
            deductions["meaning_key_unit_accuracy"] += pen * 0.35
        elif code in {"R", "T"}:
            deductions["delivery_grammar"] += pen
        else:
            deductions["completeness"] += pen * 0.35
    category_scores = {}
    for key, max_score in score_buckets.items():
        category_scores[key] = round(max(0.0, min(max_score, max_score - deductions[key])), 1)
    total = round(sum(category_scores.values()), 1)
    wpm = wpm_summary_from_rows(rows)
    high_units = [u for u in units if "high" in str(u.get("priority", "") or "").lower()]
    legal_units = [u for u in units if str(u.get("error_code", "") or "").upper().startswith("L")]
    omission_units = [u for u in units if str(u.get("error_code", "") or "").upper().startswith("O")]
    number_name_units = [u for u in units if str(u.get("error_code", "") or "").upper()[:1] in {"N", "P"}]
    strengths = []
    if ai_summary:
        strengths.extend([str(x) for x in ensure_list(ai_summary.get("overall_strengths", []))[:4]])
    if not strengths:
        strengths = [
            "You completed the attempt and produced a usable interpretation transcript for review.",
            "The review data can now identify repeated terminology, completeness, and key-unit accuracy patterns.",
        ]
    major_hits = [_unit_text(u) for u in high_units[:8]] or [_unit_text(u) for u in units[:8]]
    important_omissions = [_unit_text(u) for u in omission_units[:8]]
    priorities = []
    if legal_units:
        priorities.append("Drill legal terminology/key standards until the preferred rendering is automatic.")
    if number_name_units:
        priorities.append("Slow down enough to stabilize names, dates, times, case numbers, and other critical facts.")
    if omission_units:
        priorities.append("Practice end-of-sentence completion and key-unit capture before increasing speed.")
    if wpm.get("average_wpm"):
        avg = float(wpm.get("average_wpm") or 0)
        if avg < target_wpm - 10:
            priorities.append(f"Build speed gradually from about {avg:.0f} WPM toward {target_wpm} WPM while preserving key facts.")
        elif avg > target_wpm + 15:
            priorities.append(f"Consider lowering practice playback closer to {target_wpm} WPM until accuracy stabilizes.")
        else:
            priorities.append(f"You are practicing near the {target_wpm} WPM target; focus on precision rather than more speed.")
    if not priorities:
        priorities = ["Review the rubric key-unit table, then redo the highest-priority segments aloud."]
    if total >= passing_target:
        bottom = "This looks at or above the current practice passing target, but review the high-priority key units before treating it as stable."
    else:
        bottom = "This is a developing practice attempt below the current passing target; the fastest improvement should come from the listed high-priority key units and targeted remedial drills."
    confidence = "Medium"
    if len(units) < 8:
        confidence = "Low"
    elif any(ai_feedback_rows or []) or ai_summary:
        confidence = "Medium-High"
    return {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "script_id": payload.get("script_id", ""),
        "title": payload.get("title", ""),
        "estimated_score": total,
        "likely_level": _score_level(total, passing_target),
        "passing_target": passing_target,
        "confidence": confidence,
        **category_scores,
        "overall_summary": (
            f"Estimated practice score: {total}/100. This estimate is based on rubric/key-unit prompts, automated review flags, "
            "AI feedback when available, and transcript alignment data. It is a study estimate, not an official exam result."
        ),
        "main_strengths": strengths[:8],
        "major_scoring_hits": major_hits[:10],
        "important_omissions": important_omissions[:10],
        "practice_priorities": priorities[:8],
        "bottom_line": bottom,
        "wpm_summary": wpm,
        "rubric_unit_count": len(units),
    }


def rubric_score_summary_to_csv(summary: dict[str, Any]) -> str:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=RUBRIC_SCORE_CSV_FIELDS, extrasaction="ignore")
    writer.writeheader()
    row = {}
    for field in RUBRIC_SCORE_CSV_FIELDS:
        val = summary.get(field, "")
        row[field] = json.dumps(val, ensure_ascii=False) if isinstance(val, (list, dict)) else val
    writer.writerow(row)
    return output.getvalue()


def rubric_score_summary_to_txt(summary: dict[str, Any]) -> str:
    parts = [
        "Rubric Score + Performance Summary",
        f"Generated: {summary.get('generated_at', '')}",
        "Estimated practice score only. Not an official exam result.",
        "",
        f"Estimated score: {summary.get('estimated_score', '')}/100",
        f"Likely level: {summary.get('likely_level', '')}",
        f"Passing target: {summary.get('passing_target', '')}/100",
        f"Confidence: {summary.get('confidence', '')}",
        "",
        "Category breakdown:",
        f"- Meaning / key-unit accuracy: {summary.get('meaning_key_unit_accuracy', '')}/55",
        f"- Completeness: {summary.get('completeness', '')}/25",
        f"- Legal terminology: {summary.get('legal_terminology', '')}/12",
        f"- Delivery / grammar: {summary.get('delivery_grammar', '')}/8",
        "",
        "Overall summary:",
        str(summary.get("overall_summary", "") or ""),
    ]
    def add_list(title: str, items: list[Any]) -> None:
        parts.append(f"\n{title}:")
        if not items:
            parts.append("- None listed.")
            return
        for item in items:
            parts.append("- " + str(item))
    add_list("Main strengths", ensure_list(summary.get("main_strengths", [])))
    add_list("Major scoring hits", ensure_list(summary.get("major_scoring_hits", [])))
    add_list("Important omissions", ensure_list(summary.get("important_omissions", [])))
    add_list("Practice priorities", ensure_list(summary.get("practice_priorities", [])))
    parts.append("\nWPM summary:\n" + json.dumps(summary.get("wpm_summary", {}), ensure_ascii=False, indent=2))
    parts.append("\nBottom line:\n" + str(summary.get("bottom_line", "") or ""))
    return "\n".join(parts).strip() + "\n"


def build_rubric_score_summary_docx(summary: dict[str, Any], out_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise RuntimeError("python-docx is not installed. Run: python -m pip install python-docx") from exc
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    with contextlib.suppress(Exception):
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)
    title = doc.add_paragraph()
    run = title.add_run("Rubric Score + Performance Summary")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("Estimated practice score only. Not an official exam result.")
    doc.add_paragraph(f"Estimated score: {summary.get('estimated_score', '')}/100")
    doc.add_paragraph(f"Likely level: {summary.get('likely_level', '')}")
    doc.add_paragraph(f"Passing target: {summary.get('passing_target', '')}/100 | Confidence: {summary.get('confidence', '')}")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Category"
    table.rows[0].cells[1].text = "Estimate"
    for label, key, denom in [
        ("Meaning / key-unit accuracy", "meaning_key_unit_accuracy", 55),
        ("Completeness", "completeness", 25),
        ("Legal terminology", "legal_terminology", 12),
        ("Delivery / grammar", "delivery_grammar", 8),
    ]:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = f"{summary.get(key, '')}/{denom}"
    doc.add_heading("Overall summary", level=1)
    doc.add_paragraph(str(summary.get("overall_summary", "") or ""))
    def add_section(title: str, items: list[Any]) -> None:
        doc.add_heading(title, level=1)
        if not items:
            doc.add_paragraph("None listed.")
            return
        for item in items:
            doc.add_paragraph(str(item), style=None)
    add_section("Main strengths", ensure_list(summary.get("main_strengths", [])))
    add_section("Major scoring hits", ensure_list(summary.get("major_scoring_hits", [])))
    add_section("Important omissions", ensure_list(summary.get("important_omissions", [])))
    add_section("Practice priorities", ensure_list(summary.get("practice_priorities", [])))
    doc.add_heading("WPM summary", level=1)
    doc.add_paragraph(json.dumps(summary.get("wpm_summary", {}), ensure_ascii=False, indent=2))
    doc.add_heading("Bottom line", level=1)
    doc.add_paragraph(str(summary.get("bottom_line", "") or ""))
    doc.save(out_path)
    return out_path


def save_rubric_score_summary_outputs(payload: dict[str, Any], rows: list[dict[str, Any]], units: list[dict[str, Any]],
                                      rule_issues: list[dict[str, Any]], ai_feedback_rows: list[dict[str, Any]],
                                      ai_summary: dict[str, Any], script_id: str, session_id: str,
                                      out_dir: Path, target_wpm: int = 140,
                                      passing_target: int = 70) -> tuple[dict[str, Any], Path, Path, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    summary = build_rubric_score_summary(payload, rows, units, rule_issues, ai_feedback_rows, ai_summary, target_wpm=target_wpm, passing_target=passing_target)
    base = f"{sanitize_filename_part(script_id)}_{session_id}_rubric_score_summary"
    txt_path = out_dir / f"{base}.txt"
    csv_path = out_dir / f"{base}.csv"
    docx_path = out_dir / f"{base}.docx"
    txt_path.write_text(rubric_score_summary_to_txt(summary), encoding="utf-8")
    csv_path.write_text(rubric_score_summary_to_csv(summary), encoding="utf-8")
    build_rubric_score_summary_docx(summary, docx_path)
    return summary, txt_path, csv_path, docx_path


# -----------------------------
# Phase 3C.3.1 Strong AI rubric + usable remedial JSON
# -----------------------------
# Phase 3C.3.2 rubric guidance used inside the single AI evaluation call.
EXAM_ADMIN_RUBRIC_GUIDANCE = """
Practice scoring should follow the extracted exam-administration rubric:
- Treat the exam as objective/key-unit driven. Start conceptually from 100 points and deduct for key units rendered incorrectly or omitted.
- Prioritize key-unit accuracy over smoothness. Pauses, restarts, and self-corrections matter less than preserving message-bearing content unless they prevent completion.
- High-risk key units include legal standards, charges, allegations, names, titles, dates, times, locations, addresses, figures/amounts, case numbers, exhibit numbers, and other factual/legal details.
- If a key unit contains multiple factual elements, partial credit/partial deductions are appropriate.
- Passing target for practice scoring is 70/100. A performance may sound fluent but still fall below passing if many key units are missed or distorted.
- When grading, separate: meaning/key-unit accuracy, completeness, legal terminology, and delivery/grammar.
- For simultaneous work, note end-of-sentence fadeouts, lag-related omissions, compressed/unfinished ideas, and distortions caused by rushing.
- Self-corrections should be credited when the final rendition preserves the key unit; graders generally use the final rendition.
""".strip()

PROJECT_TERMINOLOGY_PREFERENCES_FOR_AI = """
Apply these project terminology preferences when judging legal terminology and when writing remedial references: officer/police officer = agente; the State/prosecution = la fiscalía; criminal defendant = el acusado; civil defendant = demandado; the Court when it means the judge = el juez; scene = lugar; scene of the crime = lugar de los hechos; probable cause = motivo fundado; hold to answer = consignar; exhibit = elemento de prueba; testimony = declaración; no further questions = no tengo más preguntas; probation = libertad a prueba; parole = libertad preparatoria; probation violation = incumplimiento/incumplimiento de la libertad a prueba; restitution = resarcimiento; judgment/ruling = fallo; restraining order = orden de prohibición; temporary restraining order = orden de prohibición provisional; under the influence = bajo los efectos; factual basis = bases fácticas; no contest = no me opongo; vehicle registration = matrícula de circulación; auto insurance = seguro automovilístico; valid driver’s license = licencia de conducir válida; take judicial notice = tomar conocimiento de oficio; court clerk = actuario.
""".strip()

def compact_review_rows_for_ai(rows: list[dict[str, Any]], max_rows: int = 80) -> list[dict[str, Any]]:
    """Return only the fields the rubric/remedial AI needs."""
    out: list[dict[str, Any]] = []
    for idx, row in enumerate(rows or [], start=1):
        out.append({
            "index": idx,
            "item_label": str(row.get("item_label", row.get("label", f"Item {idx}")) or f"Item {idx}"),
            "source_language": str(row.get("source_language", "") or ""),
            "target_language": str(row.get("target_language", "") or ""),
            "speaker": str(row.get("speaker", "") or ""),
            "segment_type": str(row.get("segment_type", "") or ""),
            "source_text": str(row.get("source_text", "") or "").strip(),
            "reference_text": str(row.get("reference_text", "") or "").strip(),
            "your_transcript": get_row_transcript_for_feedback(row).strip(),
            "effective_wpm": row.get("effective_wpm", row.get("wpm", "")),
            "wpm_delta_from_target": row.get("wpm_delta_from_target", ""),
        })
        if len(out) >= max_rows:
            break
    return out


def rubric_remedial_ai_prompt_payload(payload: dict[str, Any], rows: list[dict[str, Any]], target_wpm: int = 140, max_segments: int = 18) -> str:
    # IMPORTANT: Do not send the original script terms_used into this AI call.
    # The remedial terms_used list should be newly inferred from the actual
    # source/reference/user-attempt comparison plus project terminology preferences.
    data = {
        "task": "Grade a legal simultaneous/consecutive interpreting practice attempt and create a focused remedial JSON script for the user's audio generator.",
        "exam_rubric_guidance": EXAM_ADMIN_RUBRIC_GUIDANCE,
        "project_terminology_preferences": PROJECT_TERMINOLOGY_PREFERENCES_FOR_AI,
        "important_instructions": [
            "Base the score and remedial JSON ONLY on source_text, reference_text, your_transcript, WPM data when available, and project terminology preferences.",
            "Do not rely on, request, infer from, or mention the original source JSON terms_used list. Create a new terms_used list only from actual weak/high-value units found in the comparison.",
            "Use the exam rubric guidance: objective/key-unit accuracy matters more than smoothness; dates, times, names, numbers, charges, legal standards, locations, and message-bearing details are high-risk key units.",
            "Give a short estimated practice grade, not an official exam score.",
            "Identify concrete weak sections from the transcript results: omissions, distortions, unstable legal terms, numbers/dates/names, actor reversals, and unfinished final clauses.",
            "The remedial JSON paired_segments must contain actual weak or high-value sections to re-practice, with the source section and reference rendering. Do not create paired_segments that are only generic term cards.",
            "Create terms_used as flashcards for high-value terms only when those terms are present in the weak/high-value sections or needed to correct the user's attempt.",
            "If the transcript alignment looks imperfect, be fair and say possible issue rather than certain error, but still create useful drills.",
            "Return only valid JSON. No markdown fences.",
        ],
        "score_schema": {
            "estimated_score": "number 0-100",
            "likely_level": "short label",
            "passing_target": 70,
            "confidence": "Low/Medium/Medium-High/High",
            "meaning_key_unit_accuracy": "number out of 55",
            "completeness": "number out of 25",
            "legal_terminology": "number out of 12",
            "delivery_grammar": "number out of 8",
            "overall_summary": "plain English paragraph",
            "main_strengths": ["specific strengths"],
            "major_scoring_hits": ["specific major issue with safer rendering"],
            "important_omissions": ["specific omitted/weakened ideas"],
            "practice_priorities": ["specific practice priorities"],
            "bottom_line": "plain English conclusion",
        },
        "key_unit_finding_schema": {
            "priority": "1 high, 2 medium, 3 lower",
            "item_label": "row/segment label",
            "key_unit": "message-bearing unit or term",
            "category": "Legal terminology | Omission | Distortion | Number/date/name | Actor/relationship | Completion | Delivery",
            "status": "Preserved/Partial/Distorted/Omitted/Needs manual review",
            "error_code": "O/D/L/N/P/A/S/R/T/M",
            "point_impact": "High/Medium/Low",
            "source_text": "exact weak source section",
            "reference_text": "reference Spanish/English rendering",
            "your_transcript": "user attempt excerpt",
            "recommended_drill": "practical drill instruction",
            "evidence": "why this was selected",
        },
        "remedial_json_schema": {
            "script_id": "SOURCEID-CUSTOM-REMEDIAL-DRILL-001",
            "category": "Simultaneous Remedial Drill or Consecutive Remedial Drill",
            "subcategory": "based on source script",
            "title": "Customized Drill ...",
            "format": "simultaneous_remedial_drill_with_flashcards",
            "jurisdiction_context": {"source_practice": "...", "scenario": "...", "purpose": "...", "recommended_use": "..."},
            "audio_profile": "reuse or create compatible profile",
            "terms_used": [{"english": "...", "spanish_contextual_choice": "...", "spanish_alternatives_from_sheet": ["..."], "drill_note": "..."}],
            "performance_targets": ["..."],
            "english_script": "combined English drill source sections",
            "spanish_script": "combined Spanish references",
            "paired_segments": [{"segment_number": 1, "segment_type": "...", "speaker": "PROSECUTOR", "source_language": "english", "playback_order": ["english", "spanish"], "english": "weak/high-value English source section", "spanish": "Spanish reference rendering", "performance_focus": "what to fix"}],
        },
        "output_schema": {
            "rubric_score_summary": "score_schema object",
            "key_unit_findings": ["key_unit_finding_schema objects"],
            "remedial_practice_json": "remedial_json_schema object",
            "usage_summary": "short instructions for using the remedial JSON",
        },
        "script_metadata": {
            "script_id": payload.get("script_id", ""),
            "title": payload.get("title", ""),
            "category": payload.get("category", ""),
            "subcategory": payload.get("subcategory", ""),
            "format": payload.get("format", ""),
            "target_wpm": target_wpm,
            "max_remedial_segments": max_segments,
        },
        "audio_profile": payload.get("audio_profile", {}),
        "review_rows": compact_review_rows_for_ai(rows, max_rows=90),
    }
    return json.dumps(data, ensure_ascii=False, indent=2)


def openai_generate_rubric_remedial(model: str, api_key: str, prompt_payload: str) -> dict[str, Any]:
    if not api_key:
        raise ValueError("Missing OpenAI API key. Enter one in the app or set OPENAI_API_KEY in your environment.")
    try:
        from openai import OpenAI
    except Exception as exc:
        raise RuntimeError("The openai package is not installed. Run: python -m pip install openai") from exc
    client = OpenAI(api_key=api_key)
    instructions = (
        "You are a legal interpreter exam coach and remedial-drill writer. "
        "Use the provided exam-administration rubric guidance to analyze the source transcript, reference interpretation, and user's attempt. "
        "Do not use the original script terms_used as an input; create new remedial terms only from the transcript comparison and project terminology preferences. "
        "Produce a concise rubric-style practice score and a useful remedial JSON script that can be fed directly into an audio generator. "
        "The remedial JSON must include weak/high-value source sections in paired_segments, not just generic vocabulary. "
        "Return only valid JSON."
    )
    text = ""
    try:
        response = client.responses.create(model=model, instructions=instructions, input=prompt_payload)
        text = str(getattr(response, "output_text", "") or response_to_transcript_text(response)).strip()
    except Exception:
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": instructions}, {"role": "user", "content": prompt_payload}],
            temperature=0.15,
        )
        text = str(response.choices[0].message.content or "").strip()
    return parse_ai_feedback_json(text)


def normalize_ai_score_summary(raw_summary: dict[str, Any], payload: dict[str, Any], target_wpm: int = 140, passing_target: int = 70) -> dict[str, Any]:
    s = dict(raw_summary or {})
    def num(key: str, default: float) -> float:
        try:
            return round(float(s.get(key, default)), 1)
        except Exception:
            return round(float(default), 1)
    meaning = num("meaning_key_unit_accuracy", 0)
    comp = num("completeness", 0)
    legal = num("legal_terminology", 0)
    delivery = num("delivery_grammar", 0)
    estimated = num("estimated_score", meaning + comp + legal + delivery)
    return {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "script_id": payload.get("script_id", ""),
        "title": payload.get("title", ""),
        "estimated_score": estimated,
        "likely_level": str(s.get("likely_level", _score_level(estimated, passing_target)) or _score_level(estimated, passing_target)),
        "passing_target": int(s.get("passing_target", passing_target) or passing_target),
        "confidence": str(s.get("confidence", "Medium") or "Medium"),
        "meaning_key_unit_accuracy": meaning,
        "completeness": comp,
        "legal_terminology": legal,
        "delivery_grammar": delivery,
        "overall_summary": str(s.get("overall_summary", "") or f"Estimated practice score: {estimated}/100. This is a study estimate, not an official exam result."),
        "main_strengths": ensure_list(s.get("main_strengths", [])),
        "major_scoring_hits": ensure_list(s.get("major_scoring_hits", [])),
        "important_omissions": ensure_list(s.get("important_omissions", [])),
        "practice_priorities": ensure_list(s.get("practice_priorities", [])),
        "bottom_line": str(s.get("bottom_line", "") or "Use the remedial JSON to drill the listed weak/high-value sections."),
        "wpm_summary": s.get("wpm_summary", {}),
        "rubric_unit_count": int(s.get("rubric_unit_count", 0) or 0),
    }


def normalize_ai_key_units(raw_units: list[Any], rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    units: list[dict[str, Any]] = []
    row_by_label = {str(r.get("item_label", "") or ""): r for r in rows or []}
    for i, item in enumerate(raw_units or [], start=1):
        if not isinstance(item, dict):
            item = {"key_unit": str(item), "evidence": str(item)}
        label = str(item.get("item_label", "") or item.get("label", "") or f"Item {i}")
        row = row_by_label.get(label, {})
        code = str(item.get("error_code", "") or rubric_error_code_from_text(str(item.get("category", "")), str(item.get("key_unit", "")))).upper()[:1] or "M"
        if code not in RUBRIC_ERROR_CODE_LABELS:
            code = "M"
        priority_raw = item.get("priority", 2)
        try:
            priority = int(priority_raw)
        except Exception:
            priority = {"high": 1, "medium": 2, "low": 3}.get(str(priority_raw).lower(), 2)
        units.append(make_rubric_key_unit(
            priority,
            label,
            str(item.get("key_unit", "") or item.get("focus", "") or "Review this key unit")[:500],
            str(item.get("category", "Review") or "Review"),
            str(item.get("status", "Needs review") or "Needs review"),
            code,
            str(item.get("point_impact", "Medium") or "Medium"),
            str(item.get("source_text", row.get("source_text", "")) or ""),
            str(item.get("reference_text", row.get("reference_text", "")) or ""),
            str(item.get("your_transcript", get_row_transcript_for_feedback(row) if row else "") or ""),
            str(item.get("recommended_drill", "Redo this section and compare to the reference.") or "Redo this section and compare to the reference."),
            str(item.get("evidence", "") or ""),
        ))
    units.sort(key=lambda x: (int(x.get("priority", 9) or 9), str(x.get("error_code", ""))))
    return units


def remedial_terms_from_key_units(units: list[dict[str, Any]], limit: int = 40) -> list[dict[str, Any]]:
    """Create fallback remedial terms from AI key-unit findings only.

    This intentionally does not read the original payload terms_used list.
    """
    terms: list[dict[str, Any]] = []
    seen: set[str] = set()
    for u in units or []:
        key = str(u.get("key_unit", "") or "").strip()
        ref = str(u.get("reference_text", "") or "").strip()
        if not key:
            continue
        norm = normalize_for_matching(key)
        if not norm or norm in seen:
            continue
        seen.add(norm)
        terms.append({
            "english": key[:180],
            "spanish_contextual_choice": ref[:220] if ref else "",
            "spanish_alternatives_from_sheet": [ref[:220]] if ref else [],
            "drill_note": str(u.get("recommended_drill", "") or u.get("evidence", "") or "Selected from rubric/key-unit transcript review.")[:260],
        })
        if len(terms) >= limit:
            break
    return terms


def normalize_ai_remedial_json(raw_json: dict[str, Any], payload: dict[str, Any], units: list[dict[str, Any]], script_id: str, session_id: str, target_wpm: int = 140) -> dict[str, Any]:
    r = dict(raw_json or {})
    r.setdefault("script_id", f"{sanitize_filename_part(script_id)}-CUSTOM-REMEDIAL-DRILL-001")
    r.setdefault("category", "Simultaneous Remedial Drill" if "simult" in normalize_for_matching(payload.get("format", "")) else "Remedial Drill")
    r.setdefault("subcategory", str(payload.get("subcategory", "") or payload.get("title", "") or "Rubric-based remedial practice"))
    r.setdefault("title", f"Customized Remedial Drill — {payload.get('title', script_id)}")
    r.setdefault("format", "simultaneous_remedial_drill_with_flashcards")
    r.setdefault("jurisdiction_context", {})
    if not isinstance(r["jurisdiction_context"], dict):
        r["jurisdiction_context"] = {"purpose": str(r["jurisdiction_context"])}
    r["jurisdiction_context"].setdefault("source_practice", f"{script_id} {session_id}")
    r["jurisdiction_context"].setdefault("purpose", "Drill high-value key units that were missed, distorted, or weakened in the exam attempt.")
    r["jurisdiction_context"].setdefault("recommended_use", "Run flashcards first, then source-only audio. Interpret each segment aloud before listening to the reference Spanish. Repeat weak sections until automatic.")
    ap = r.get("audio_profile") if isinstance(r.get("audio_profile"), dict) else payload.get("audio_profile", {})
    if not isinstance(ap, dict) or not ap:
        ap = {"voice_mode": "speaker", "default_voices": {"english": "en-US-EmmaNeural", "spanish": "es-US-PalomaNeural"}, "voice_assignments": {"DRILL COACH": {"english": "en-US-MichelleNeural", "spanish": "es-MX-DaliaNeural", "description": "remedial drill voice"}}, "speak_speaker_labels": False}
    r["audio_profile"] = ap
    if not isinstance(r.get("terms_used"), list) or not r.get("terms_used"):
        r["terms_used"] = remedial_terms_from_key_units(units, limit=40)
    if not isinstance(r.get("performance_targets"), list) or not r.get("performance_targets"):
        r["performance_targets"] = [str(u.get("recommended_drill", "")) for u in units[:8] if str(u.get("recommended_drill", "")).strip()]
    paired = r.get("paired_segments") if isinstance(r.get("paired_segments"), list) else []
    if not paired:
        paired = []
        for idx, u in enumerate([x for x in units if str(x.get("source_text", "")).strip() or str(x.get("reference_text", "")).strip()][:18], start=1):
            paired.append({
                "segment_number": idx,
                "segment_type": normalize_for_matching(u.get("category", "remedial_drill")).replace(" ", "_")[:60] or "remedial_drill",
                "speaker": "DRILL COACH",
                "source_language": "english",
                "playback_order": ["english", "spanish"],
                "english": str(u.get("source_text", "") or u.get("key_unit", "")),
                "spanish": str(u.get("reference_text", "") or ""),
                "performance_focus": str(u.get("recommended_drill", "") or u.get("evidence", "") or "Key-unit remedial practice"),
            })
    clean_paired = []
    for idx, seg in enumerate(paired, start=1):
        if not isinstance(seg, dict):
            continue
        english = str(seg.get("english", "") or seg.get("source_text", "") or "").strip()
        spanish = str(seg.get("spanish", "") or seg.get("reference_text", "") or "").strip()
        if not english and not spanish:
            continue
        seg.setdefault("segment_number", idx)
        seg.setdefault("segment_type", "remedial_drill")
        seg.setdefault("speaker", "DRILL COACH")
        seg.setdefault("source_language", "english")
        seg.setdefault("playback_order", ["english", "spanish"])
        seg["english"] = english
        seg["spanish"] = spanish
        seg.setdefault("performance_focus", "Key-unit remedial practice")
        clean_paired.append(seg)
    r["paired_segments"] = clean_paired
    if not str(r.get("english_script", "")).strip():
        r["english_script"] = "\n".join(str(seg.get("english", "")) for seg in clean_paired if str(seg.get("english", "")).strip())
    if not str(r.get("spanish_script", "")).strip():
        r["spanish_script"] = "\n".join(str(seg.get("spanish", "")) for seg in clean_paired if str(seg.get("spanish", "")).strip())
    return r


def save_ai_rubric_remedial_outputs(payload: dict[str, Any], rows: list[dict[str, Any]], script_id: str, session_id: str,
                                    out_dir: Path, model: str, api_key: str, target_wpm: int = 140,
                                    passing_target: int = 70, max_segments: int = 18) -> tuple[dict[str, Any], list[dict[str, Any]], list[dict[str, Any]], Path, Path, Path, Path, Path, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    prompt = rubric_remedial_ai_prompt_payload(payload, rows, target_wpm=target_wpm, max_segments=max_segments)
    raw = openai_generate_rubric_remedial(model, api_key, prompt)
    summary = normalize_ai_score_summary(raw.get("rubric_score_summary", raw.get("score_summary", {})), payload, target_wpm, passing_target)
    units = normalize_ai_key_units(raw.get("key_unit_findings", raw.get("rubric_key_units", [])), rows)
    if not units:
        units = build_rubric_key_units(payload, rows, [], [], {}, max_units=40)
    summary["rubric_unit_count"] = len(units)
    if not summary.get("wpm_summary"):
        summary["wpm_summary"] = wpm_summary_from_rows(rows)
    remedial = normalize_ai_remedial_json(raw.get("remedial_practice_json", raw.get("remedial_json", {})), payload, units, script_id, session_id, target_wpm)
    drills = rubric_units_to_drills(units, target_wpm=target_wpm)
    base = f"{sanitize_filename_part(script_id)}_{session_id}"
    score_txt = out_dir / f"{base}_rubric_score_summary.txt"
    score_csv = out_dir / f"{base}_rubric_score_summary.csv"
    score_docx = out_dir / f"{base}_rubric_score_summary.docx"
    units_txt = out_dir / f"{base}_rubric_key_units.txt"
    units_csv = out_dir / f"{base}_rubric_key_units.csv"
    remedial_json_path = out_dir / f"{base}_remedial_drill_script.json"
    score_txt.write_text(rubric_score_summary_to_txt(summary), encoding="utf-8")
    score_csv.write_text(rubric_score_summary_to_csv(summary), encoding="utf-8")
    build_rubric_score_summary_docx(summary, score_docx)
    units_txt.write_text(rubric_key_units_to_txt(units), encoding="utf-8")
    units_csv.write_text(rubric_key_units_to_csv(units), encoding="utf-8")
    remedial_json_path.write_text(json.dumps(remedial, ensure_ascii=False, indent=2), encoding="utf-8")
    return summary, units, drills, score_txt, score_csv, score_docx, units_txt, units_csv, remedial_json_path


def render_rubric_score_summary(summary: dict[str, Any]) -> None:
    if not summary:
        return
    st.markdown("#### Rubric score + performance summary")
    st.caption("Estimated practice score only — not an official exam result.")
    cols = st.columns(4)
    cols[0].metric("Estimated score", f"{summary.get('estimated_score', '')}/100")
    cols[1].metric("Passing target", f"{summary.get('passing_target', '')}/100")
    cols[2].metric("Key-unit accuracy", f"{summary.get('meaning_key_unit_accuracy', '')}/55")
    cols[3].metric("Completeness", f"{summary.get('completeness', '')}/25")
    st.markdown(compact_metadata_row([summary.get("likely_level", ""), f"Confidence: {summary.get('confidence', '')}", f"Units reviewed: {summary.get('rubric_unit_count', '')}"]), unsafe_allow_html=True)
    st.write(str(summary.get("overall_summary", "") or ""))
    breakdown = pd.DataFrame([
        {"Category": "Meaning / key-unit accuracy", "Estimate": summary.get("meaning_key_unit_accuracy", ""), "Max": 55},
        {"Category": "Completeness", "Estimate": summary.get("completeness", ""), "Max": 25},
        {"Category": "Legal terminology", "Estimate": summary.get("legal_terminology", ""), "Max": 12},
        {"Category": "Delivery / grammar", "Estimate": summary.get("delivery_grammar", ""), "Max": 8},
    ])
    st.dataframe(breakdown, use_container_width=True, hide_index=True)
    with st.expander("Main strengths", expanded=True):
        for item in ensure_list(summary.get("main_strengths", [])):
            st.write("- " + str(item))
    with st.expander("Major scoring hits", expanded=True):
        for item in ensure_list(summary.get("major_scoring_hits", [])):
            st.write("- " + str(item))
    with st.expander("Important omissions and practice priorities", expanded=True):
        omissions = ensure_list(summary.get("important_omissions", []))
        if omissions:
            st.markdown("**Important omissions / possible omissions**")
            for item in omissions:
                st.write("- " + str(item))
        st.markdown("**Practice priorities**")
        for item in ensure_list(summary.get("practice_priorities", [])):
            st.write("- " + str(item))
    st.info(str(summary.get("bottom_line", "") or ""))


# -----------------------------
# Phase 3C.2D Interactive Drill JSON + Drill Studio helpers
# -----------------------------
DRILL_STUDIO_LOG_FIELDS = [
    "timestamp", "session_id", "drill_set_id", "source_exam", "drill_index", "drill_type",
    "priority", "focus", "direction", "score", "notes", "prompt", "target", "related_terms",
]


def infer_drill_language(text_value: str, fallback: str = "english") -> str:
    raw = str(text_value or "")
    if not raw.strip():
        return fallback
    lowered = f" {raw.lower()} "
    spanish_markers = ["á", "é", "í", "ó", "ú", "ñ", "¿", "¡", " el ", " la ", " los ", " las ", " de ", " que ", " por ", " para ", " usted ", " señor", " señora"]
    english_markers = [" the ", " and ", " you ", " your ", " court ", " judge ", " proof ", " defendant "]
    sp = sum(1 for m in spanish_markers if m in lowered)
    en = sum(1 for m in english_markers if m in lowered)
    if sp > en:
        return "spanish"
    if en > sp:
        return "english"
    return fallback


def build_interactive_drill_json(payload: dict[str, Any], drills: list[dict[str, Any]], srs_terms: list[dict[str, Any]] | None,
                                 script_id: str, session_id: str, target_wpm: int = 140) -> dict[str, Any]:
    created_at = datetime.now().isoformat(timespec="seconds")
    drill_set_id = f"{sanitize_filename_part(script_id)}_{session_id}_interactive_drills"
    normalized_drills: list[dict[str, Any]] = []
    for idx, d in enumerate(drills or [], start=1):
        prompt = str(d.get("prompt", "") or d.get("source_text", "") or d.get("focus", "") or "").strip()
        target = str(d.get("target", "") or d.get("reference_text", "") or "").strip()
        source_text = str(d.get("source_text", "") or "").strip()
        reference_text = str(d.get("reference_text", "") or "").strip()
        your_transcript = str(d.get("your_transcript", "") or "").strip()
        direction = str(d.get("direction", "") or "").strip()
        prompt_lang = infer_drill_language(prompt or source_text, "english")
        target_lang = infer_drill_language(target or reference_text, "spanish" if prompt_lang == "english" else "english")
        drill_type = str(d.get("drill_type", "Review") or "Review")
        dt_norm = normalize_for_matching(drill_type)
        if "term" in dt_norm:
            interaction = "audio_flashcard"
        elif "number" in dt_norm or "date" in dt_norm or "name" in dt_norm:
            interaction = "fact_recall"
        elif "wpm" in dt_norm:
            interaction = "source_replay_progression"
        elif "omission" in dt_norm or "redo" in dt_norm or "repair" in dt_norm:
            interaction = "redo_chunk"
        else:
            interaction = "review_card"
        normalized_drills.append({
            "drill_id": short_hash(drill_set_id, str(idx), drill_type, prompt, target, str(d.get("focus", ""))),
            "drill_index": idx,
            "priority": str(d.get("priority", "") or ""),
            "drill_type": drill_type,
            "interaction_mode": interaction,
            "source": str(d.get("source", "") or ""),
            "item_label": str(d.get("item_label", "") or ""),
            "focus": str(d.get("focus", "") or ""),
            "direction": direction,
            "prompt_language": prompt_lang,
            "target_language": target_lang,
            "prompt": prompt,
            "target": target,
            "instructions": str(d.get("instructions", "") or ""),
            "suggested_reps": str(d.get("suggested_reps", "") or ""),
            "wpm_target": str(d.get("wpm_target", "") or target_wpm or ""),
            "related_terms": str(d.get("related_terms", "") or ""),
            "evidence": str(d.get("evidence", "") or ""),
            "source_text": source_text,
            "reference_text": reference_text,
            "your_transcript": your_transcript,
        })
    return {
        "drill_set_id": drill_set_id,
        "version": "3C.2E",
        "created_at": created_at,
        "source_exam": script_id,
        "source_title": payload.get("title", ""),
        "target_wpm": int(target_wpm or 140),
        "created_from": {"rule_based_review": True, "ai_feedback": True, "full_summary": True, "terms_used": True},
        "srs_term_candidates": srs_terms or [],
        "drills": normalized_drills,
    }


def save_interactive_drill_json(drill_json: dict[str, Any], out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(drill_json, ensure_ascii=False, indent=2), encoding="utf-8")
    return out_path


def load_drill_json_from_bytes(data: bytes) -> dict[str, Any]:
    parsed = json.loads(data.decode("utf-8-sig"))
    if not isinstance(parsed, dict):
        raise ValueError("Drill JSON must be a JSON object.")
    has_drills = isinstance(parsed.get("drills"), list)
    has_terms = isinstance(parsed.get("terms_used"), list)
    has_segments = isinstance(parsed.get("paired_segments"), list)
    if not (has_drills or has_terms or has_segments):
        raise ValueError("Drill JSON must contain 'terms_used', 'paired_segments', or a top-level 'drills' list.")
    return parsed


def drill_studio_log_to_csv(rows: list[dict[str, Any]]) -> str:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=DRILL_STUDIO_LOG_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for row in rows or []:
        writer.writerow({k: row.get(k, "") for k in DRILL_STUDIO_LOG_FIELDS})
    return output.getvalue()


def append_drill_studio_log(path: Path, row: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    exists = path.exists() and path.stat().st_size > 0
    with path.open("a", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=DRILL_STUDIO_LOG_FIELDS, extrasaction="ignore")
        if not exists:
            writer.writeheader()
        writer.writerow({k: row.get(k, "") for k in DRILL_STUDIO_LOG_FIELDS})


def load_drill_studio_log(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            return [{k: row.get(k, "") for k in DRILL_STUDIO_LOG_FIELDS} for row in csv.DictReader(f)]
    except Exception as exc:
        st.warning(f"Could not read Drill Studio history: {exc}")
        return []


def render_drill_text_card(title: str, body: str, muted: str = "") -> None:
    body_html = html.escape(str(body or "—")).replace("\n", "<br>")
    muted_html = html.escape(str(muted or ""))
    st.markdown(
        f"""
        <div style="border:1px solid #ddd; border-radius:16px; padding:16px; background:#fff; margin-bottom:10px;">
          <div style="font-size:0.82rem; color:#666; font-weight:700; margin-bottom:8px; text-transform:uppercase; letter-spacing:.04em;">{html.escape(title)}</div>
          <div style="font-size:1.05rem; line-height:1.45;">{body_html}</div>
          <div style="font-size:0.78rem; color:#777; margin-top:10px;">{muted_html}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def drill_audio_text(drill: dict[str, Any], role: str) -> tuple[str, str]:
    if role == "prompt":
        txt = str(drill.get("prompt") or drill.get("source_text") or drill.get("focus") or "")
        lang = str(drill.get("prompt_language") or infer_drill_language(txt, "english"))
    else:
        txt = str(drill.get("target") or drill.get("reference_text") or "")
        lang = str(drill.get("target_language") or infer_drill_language(txt, "spanish"))
    lang = lang if lang in {"english", "spanish"} else infer_drill_language(txt, "english")
    return txt, lang


def normalize_remedial_term_drills(terms_used: list[dict[str, Any]], drill_set_id: str = "drill_set") -> list[dict[str, Any]]:
    """Turn remedial-script terms_used entries into Drill Studio flashcards."""
    drills: list[dict[str, Any]] = []
    for idx, term in enumerate(terms_used or [], start=1):
        if not isinstance(term, dict):
            continue
        english = str(term.get("english", "") or term.get("source", "") or "").strip()
        spanish = str(term.get("spanish_contextual_choice", "") or term.get("spanish", "") or term.get("target", "") or "").strip()
        if not english and not spanish:
            continue
        alternatives = term.get("spanish_alternatives_from_sheet", [])
        if isinstance(alternatives, list) and alternatives:
            alt_text = "; ".join(str(x) for x in alternatives if str(x).strip())
        else:
            alt_text = ""
        note = str(term.get("drill_note", "") or term.get("note", "") or term.get("performance_focus", "") or "").strip()
        drills.append({
            "drill_id": short_hash(drill_set_id, "term", str(idx), english, spanish),
            "drill_index": idx,
            "priority": str(term.get("priority", "") or ""),
            "drill_type": "Term flashcard",
            "interaction_mode": "audio_flashcard",
            "source": "terms_used",
            "item_label": f"Term {idx}",
            "focus": english or spanish,
            "direction": "English → Spanish",
            "prompt_language": "english",
            "target_language": "spanish",
            "prompt": english,
            "target": spanish,
            "instructions": "Interpret the English term or phrase into Spanish before revealing the answer.",
            "suggested_reps": "3–5",
            "related_terms": alt_text,
            "evidence": note,
            "source_text": english,
            "reference_text": spanish,
            "your_transcript": "",
        })
    return drills


def normalize_remedial_segment_drills(paired_segments: list[dict[str, Any]], drill_set_id: str = "drill_set") -> list[dict[str, Any]]:
    """Turn remedial-script paired_segments entries into Drill Studio segment practice items."""
    drills: list[dict[str, Any]] = []
    for idx, seg in enumerate(paired_segments or [], start=1):
        if not isinstance(seg, dict):
            continue
        source_lang = str(seg.get("source_language", "english") or "english").lower()
        source_lang = "spanish" if source_lang.startswith("span") else "english"
        target_lang = "spanish" if source_lang == "english" else "english"
        source_text = str(seg.get(source_lang, "") or seg.get("source_text", "") or seg.get("english", "") or "").strip()
        reference_text = str(seg.get(target_lang, "") or seg.get("reference_text", "") or seg.get("spanish", "") or "").strip()
        if not source_text and not reference_text:
            continue
        focus = str(seg.get("performance_focus", "") or seg.get("focus", "") or seg.get("segment_type", "") or "").strip()
        speaker = str(seg.get("speaker", "") or "").strip()
        segment_number = seg.get("segment_number", idx)
        label_parts = [f"Segment {segment_number}"]
        if speaker:
            label_parts.append(speaker)
        drills.append({
            "drill_id": short_hash(drill_set_id, "segment", str(idx), source_text, reference_text),
            "drill_index": idx,
            "priority": str(seg.get("priority", "") or ""),
            "drill_type": "Segment drill",
            "interaction_mode": "redo_chunk",
            "source": "paired_segments",
            "item_label": " · ".join(label_parts),
            "focus": focus or "Weak/high-value segment practice",
            "direction": f"{source_lang.title()} → {target_lang.title()}",
            "prompt_language": source_lang,
            "target_language": target_lang,
            "prompt": source_text,
            "target": reference_text,
            "instructions": "Listen to or read the source, interpret it aloud, then reveal the reference rendering and compare.",
            "suggested_reps": "2–4",
            "related_terms": str(seg.get("related_terms", "") or ""),
            "evidence": str(seg.get("rubric_note", "") or seg.get("why_selected", "") or ""),
            "source_text": source_text,
            "reference_text": reference_text,
            "your_transcript": str(seg.get("your_transcript", "") or seg.get("attempt_transcript", "") or ""),
        })
    return drills


def get_drill_studio_collections(drill_json_data: dict[str, Any]) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    """Return term flashcards, segment drills, and legacy/interactive drills from any supported drill JSON."""
    drill_set_id = str(drill_json_data.get("drill_set_id", "drill_set") or drill_json_data.get("script_id", "drill_set") or "drill_set")
    term_drills = normalize_remedial_term_drills([t for t in drill_json_data.get("terms_used", []) if isinstance(t, dict)], drill_set_id)
    segment_drills = normalize_remedial_segment_drills([s for s in drill_json_data.get("paired_segments", []) if isinstance(s, dict)], drill_set_id)
    legacy_drills = [d for d in drill_json_data.get("drills", []) if isinstance(d, dict)]
    return term_drills, segment_drills, legacy_drills


def render_drill_practice_panel(
    drills: list[dict[str, Any]],
    drill_json_data: dict[str, Any],
    drill_set_id: str,
    panel_key: str,
    panel_label: str,
    drill_audio_enabled: bool,
    drill_autoplay_prompt: bool,
    drill_play_answer: bool,
    drill_audio_only: bool,
    drill_audio_speed: float,
    drill_audio_cache_dir: Path,
    speed_mode: str,
    generator_path: Path,
    drill_studio_log_path: Path,
) -> None:
    if not drills:
        st.info(f"No {panel_label.lower()} found in this JSON.")
        return

    idx_key = f"drill_studio_{panel_key}_idx"
    reveal_key = f"drill_studio_{panel_key}_revealed"
    if idx_key not in st.session_state:
        st.session_state[idx_key] = 0
    if reveal_key not in st.session_state:
        st.session_state[reveal_key] = False

    order = st.selectbox("Order", ["Original order", "Priority order", "Random"], index=0, key=f"{panel_key}_order")
    filtered_drills = list(drills)
    if order == "Priority order":
        filtered_drills = sorted(filtered_drills, key=lambda d: (int(str(d.get("priority") or "9") if str(d.get("priority") or "9").isdigit() else 9), int(d.get("drill_index", 9999) or 9999)))
    elif order == "Random":
        seed = st.session_state.get(f"{panel_key}_random_seed") or uuid.uuid4().hex
        st.session_state[f"{panel_key}_random_seed"] = seed
        filtered_drills = sorted(filtered_drills, key=lambda d: short_hash(seed, str(d.get("drill_id", ""))))

    st.session_state[idx_key] = max(0, min(int(st.session_state[idx_key]), len(filtered_drills) - 1))
    current_idx = int(st.session_state[idx_key])
    current_drill = filtered_drills[current_idx]
    st.progress((current_idx + 1) / max(len(filtered_drills), 1), text=f"{panel_label}: {current_idx + 1} of {len(filtered_drills)}")
    st.markdown(compact_metadata_row([
        current_drill.get("item_label", f"Item {current_idx + 1}"),
        current_drill.get("drill_type", "Review"),
        current_drill.get("direction", ""),
        current_drill.get("focus", ""),
    ]), unsafe_allow_html=True)

    prompt_text, prompt_lang = drill_audio_text(current_drill, "prompt")
    answer_text, answer_lang = drill_audio_text(current_drill, "answer")
    prompt_voice = DEFAULT_PRACTICE_VOICES.get(prompt_lang, "en-US-JennyNeural")
    answer_voice = DEFAULT_PRACTICE_VOICES.get(answer_lang, "es-MX-DaliaNeural")
    audio_key_base = short_hash(drill_set_id, panel_key, str(current_drill.get("drill_id", current_idx)), prompt_text, answer_text, str(current_idx))

    if drill_audio_enabled and prompt_text.strip():
        prompt_audio_path = get_practice_audio_path(drill_audio_cache_dir, drill_set_id, audio_key_base, "prompt", prompt_lang, prompt_text, prompt_voice, get_rate_for_language(speed_mode, prompt_lang))
        try:
            asyncio.run(generate_practice_audio_file(prompt_text, prompt_voice, get_rate_for_language(speed_mode, prompt_lang), prompt_audio_path, generator_path))
            render_audio_player(prompt_audio_path, f"Prompt audio · {prompt_lang.title()}", autoplay=drill_autoplay_prompt, nonce=f"{panel_key}-prompt-{audio_key_base}-{st.session_state.get(panel_key + '_prompt_nonce','')}", playback_speed=float(drill_audio_speed))
        except Exception as exc:
            st.warning(f"Could not prepare prompt audio: {exc}")
        if st.button("↻ Replay prompt", key=f"{panel_key}_replay_prompt_{audio_key_base}"):
            st.session_state[f"{panel_key}_prompt_nonce"] = uuid.uuid4().hex[:8]
            st.rerun()

    prompt_visible = not drill_audio_only or st.session_state.get(f"{panel_key}_show_prompt_{audio_key_base}", False)
    if prompt_visible:
        render_drill_text_card("Prompt / source", prompt_text or current_drill.get("focus", ""), current_drill.get("instructions", ""))
    else:
        st.info("🎧 Audio-only prompt: listen first, perform the drill aloud, then reveal or show the text.")
        if st.button("Show prompt text", key=f"{panel_key}_show_prompt_btn_{audio_key_base}"):
            st.session_state[f"{panel_key}_show_prompt_{audio_key_base}"] = True
            st.rerun()

    if current_drill.get("your_transcript") or current_drill.get("evidence") or current_drill.get("related_terms"):
        with st.expander("Practice context", expanded=False):
            if current_drill.get("your_transcript"):
                render_drill_text_card("Your transcript", current_drill.get("your_transcript", ""), "From the exam attempt, when available")
            if current_drill.get("evidence"):
                render_drill_text_card("Why this item matters", current_drill.get("evidence", ""), current_drill.get("source", ""))
            if current_drill.get("related_terms"):
                render_drill_text_card("Related terms / alternatives", current_drill.get("related_terms", ""), "")

    if not st.session_state[reveal_key]:
        if st.button("👁 Reveal answer / reference", type="primary", key=f"{panel_key}_reveal_{audio_key_base}"):
            st.session_state[reveal_key] = True
            st.rerun()
    else:
        render_drill_text_card("Answer / reference", answer_text or current_drill.get("reference_text", "—"), current_drill.get("focus", ""))
        if drill_audio_enabled and (answer_text or current_drill.get("reference_text")):
            answer_source = answer_text or str(current_drill.get("reference_text", ""))
            answer_audio_path = get_practice_audio_path(drill_audio_cache_dir, drill_set_id, audio_key_base, "answer", answer_lang, answer_source, answer_voice, get_rate_for_language(speed_mode, answer_lang))
            try:
                asyncio.run(generate_practice_audio_file(answer_source, answer_voice, get_rate_for_language(speed_mode, answer_lang), answer_audio_path, generator_path))
                render_audio_player(answer_audio_path, f"Answer audio · {answer_lang.title()}", autoplay=drill_play_answer, nonce=f"{panel_key}-answer-{audio_key_base}-{st.session_state.get(panel_key + '_answer_nonce','')}", playback_speed=float(drill_audio_speed))
            except Exception as exc:
                st.warning(f"Could not prepare answer audio: {exc}")
            if st.button("↻ Replay answer", key=f"{panel_key}_replay_answer_{audio_key_base}"):
                st.session_state[f"{panel_key}_answer_nonce"] = uuid.uuid4().hex[:8]
                st.rerun()

        st.markdown("#### Score this item")
        score = st.radio("Result", ["Easy", "Good", "Hard", "Missed"], horizontal=True, key=f"{panel_key}_score_{audio_key_base}")
        notes = st.text_area("Notes", key=f"{panel_key}_notes_{audio_key_base}", height=80)
        c1, c2, c3 = st.columns([1, 1, 2])
        with c1:
            if st.button("✅ Save + next", key=f"{panel_key}_save_next_{audio_key_base}"):
                row = {
                    "timestamp": datetime.now().isoformat(timespec="seconds"),
                    "session_id": st.session_state.drill_studio_session_id,
                    "drill_set_id": drill_set_id,
                    "source_exam": drill_json_data.get("source_exam", drill_json_data.get("script_id", "")),
                    "drill_index": current_drill.get("drill_index", current_idx + 1),
                    "drill_type": current_drill.get("drill_type", ""),
                    "priority": current_drill.get("priority", ""),
                    "focus": current_drill.get("focus", ""),
                    "direction": current_drill.get("direction", ""),
                    "score": score,
                    "notes": notes,
                    "prompt": prompt_text,
                    "target": answer_text,
                    "related_terms": current_drill.get("related_terms", ""),
                }
                st.session_state.drill_studio_session_rows.append(row)
                try:
                    append_drill_studio_log(drill_studio_log_path, row)
                except Exception as exc:
                    st.warning(f"Could not save Drill Studio history: {exc}")
                st.session_state[idx_key] = min(current_idx + 1, len(filtered_drills) - 1)
                st.session_state[reveal_key] = False
                st.rerun()
        with c2:
            if st.button("Skip", key=f"{panel_key}_skip_{audio_key_base}"):
                st.session_state[idx_key] = min(current_idx + 1, len(filtered_drills) - 1)
                st.session_state[reveal_key] = False
                st.rerun()
        with c3:
            if st.button("Start new drill session", key=f"{panel_key}_new_session"):
                st.session_state.drill_studio_session_id = f"drill-session-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{uuid.uuid4().hex[:8]}"
                st.session_state.drill_studio_session_rows = []
                st.session_state[idx_key] = 0
                st.session_state[reveal_key] = False
                st.rerun()

    n1, n2, n3 = st.columns(3)
    with n1:
        if st.button("⬅ Previous", key=f"{panel_key}_prev"):
            st.session_state[idx_key] = max(0, current_idx - 1)
            st.session_state[reveal_key] = False
            st.rerun()
    with n2:
        if st.button("➡ Next", key=f"{panel_key}_next"):
            st.session_state[idx_key] = min(len(filtered_drills) - 1, current_idx + 1)
            st.session_state[reveal_key] = False
            st.rerun()
    with n3:
        if st.session_state.drill_studio_session_rows:
            st.download_button("Download session CSV", data=drill_studio_log_to_csv(st.session_state.drill_studio_session_rows).encode("utf-8"), file_name=f"{sanitize_filename_part(drill_set_id)}_drill_session.csv", mime="text/csv", key=f"{panel_key}_download_session")


def render_personalized_drills(drills: list[dict[str, Any]], srs_terms: list[dict[str, Any]] | None = None) -> None:
    if not drills:
        st.info("No personalized drills have been generated yet.")
        return
    c1, c2, c3 = st.columns(3)
    c1.metric("Drills", len(drills))
    c2.metric("High priority", sum(1 for d in drills if str(d.get("priority", "")) == "1"))
    c3.metric("SRS term candidates", len(srs_terms or []))
    st.dataframe(
        [{k: d.get(k, "") for k in ["priority", "drill_type", "source", "item_label", "focus", "suggested_reps", "related_terms"]} for d in drills],
        use_container_width=True,
        hide_index=True,
    )
    with st.expander("Detailed drill cards", expanded=False):
        for i, drill in enumerate(drills, start=1):
            st.markdown(f"**{i}. {drill.get('drill_type', '')}: {html.escape(str(drill.get('focus', '')))}**")
            st.caption(compact_metadata_row([f"Priority {drill.get('priority', '')}", drill.get("source", ""), drill.get("item_label", ""), drill.get("direction", "")]), unsafe_allow_html=True)
            if drill.get("prompt"):
                st.write("Prompt:", drill.get("prompt"))
            if drill.get("target"):
                st.write("Target/reference:", drill.get("target"))
            if drill.get("instructions"):
                st.info(str(drill.get("instructions")))
            if drill.get("related_terms"):
                st.caption(f"Related terms: {drill.get('related_terms')}")
            st.divider()


def missed_terms_from_review_issues(issues: list[dict[str, Any]]) -> list[str]:
    vals = []
    for issue in issues:
        if str(issue.get("category", "")).lower() == "legal term check":
            exp = str(issue.get("expected", "") or "").strip()
            if exp:
                vals.append(exp)
    return sorted(set(vals), key=str.lower)


def get_exam_audio_path(cache_dir: Path, script_id: str, item_index: Any, role: str, lang: str, text: str, voice: str, rate: str) -> Path:
    key = short_hash(script_id, str(item_index), role, lang, text, voice, rate)
    filename = f"{sanitize_filename_part(script_id)}_exam_{sanitize_filename_part(str(item_index), 20)}_{role}_{lang}_{key}.mp3"
    return cache_dir / filename


def initialize_exam_session(script_id: str, exam_mode: str) -> str:
    current = st.session_state.get("exam_session_key")
    desired = f"{script_id}|{exam_mode}"
    if current != desired:
        st.session_state.exam_session_key = desired
        st.session_state.exam_session_id = f"exam-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{uuid.uuid4().hex[:8]}"
        st.session_state.exam_index = 0
        st.session_state.exam_rows = []
        st.session_state.exam_saved = {}
        st.session_state.exam_started = False
        st.session_state.exam_transcript_rows = []
        st.session_state.exam_last_transcript_txt = ""
        st.session_state.exam_last_transcript_csv = ""
        st.session_state.exam_sim_source_audio_path = ""
        st.session_state.exam_sim_source_group_count = None
    return st.session_state.exam_session_id

# -----------------------------
# Phase 2D simultaneous practice helpers
# -----------------------------
SIMULTANEOUS_LOG_FIELDS = [
    "timestamp", "session_id", "script_id", "title", "source_file", "practice_mode",
    "training_mode", "lag_seconds", "playback_speed", "repeat_loop",
    "chunk_index", "chunk_label", "source_language", "target_language", "speaker",
    "segment_type", "score", "error_tags", "notes", "source_text", "target_text",
]

SIMULTANEOUS_ERROR_TAGS = [
    "Omission",
    "Lag collapse",
    "False start",
    "Terminology",
    "Register",
    "Numbers/names",
    "Syntax interference",
    "Incomplete thought",
    "Too literal",
    "Lost thread",
]


def simultaneous_log_to_csv(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return ""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=SIMULTANEOUS_LOG_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        writer.writerow({k: row.get(k, "") for k in SIMULTANEOUS_LOG_FIELDS})
    return output.getvalue()


def normalize_simultaneous_log_row(row: dict[str, Any]) -> dict[str, Any]:
    normalized = {k: row.get(k, "") for k in SIMULTANEOUS_LOG_FIELDS}
    normalized["score"] = str(normalized.get("score", ""))
    return normalized


def load_simultaneous_log(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            return [normalize_simultaneous_log_row(row) for row in csv.DictReader(f)]
    except Exception as exc:
        st.warning(f"Could not read simultaneous practice log: {exc}")
        return []


def append_simultaneous_log(path: Path, row: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    file_exists = path.exists() and path.stat().st_size > 0
    with path.open("a", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=SIMULTANEOUS_LOG_FIELDS, extrasaction="ignore")
        if not file_exists:
            writer.writeheader()
        writer.writerow({k: row.get(k, "") for k in SIMULTANEOUS_LOG_FIELDS})


def simultaneous_rows_to_dataframe(rows: list[dict[str, Any]]) -> pd.DataFrame:
    if not rows:
        return pd.DataFrame(columns=SIMULTANEOUS_LOG_FIELDS)
    df = pd.DataFrame([normalize_simultaneous_log_row(r) for r in rows])
    for col in SIMULTANEOUS_LOG_FIELDS:
        if col not in df.columns:
            df[col] = ""
    df["score_numeric"] = pd.to_numeric(df["score"], errors="coerce")
    df["timestamp_dt"] = pd.to_datetime(df["timestamp"], errors="coerce")
    df["date"] = df["timestamp_dt"].dt.date.astype("string")
    df["direction"] = df["source_language"].str.title() + " → " + df["target_language"].str.title()
    return df


def build_simultaneous_chunks(data: dict[str, Any], mode: str, direction: str = "Source language first") -> list[dict[str, Any]]:
    paired = data.get("paired_segments", []) or []
    if not isinstance(paired, list):
        paired = []
    script_id = safe_script_id(data, "simultaneous")
    chunks: list[dict[str, Any]] = []
    valid_segments: list[dict[str, Any]] = []
    for idx, seg in enumerate(paired, start=1):
        if not isinstance(seg, dict):
            continue
        source, target = get_segment_source_and_target(seg, direction)
        source_text = str(seg.get(source, "")).strip()
        target_text = str(seg.get(target, "")).strip()
        if not source_text or not target_text:
            continue
        copy = dict(seg)
        copy.update({
            "_segment_index": idx,
            "_source_language": source,
            "_target_language": target,
            "_source_text": source_text,
            "_target_text": target_text,
        })
        valid_segments.append(copy)

    if mode == "Full passage":
        if not valid_segments:
            return []
        source_lang = valid_segments[0]["_source_language"]
        target_lang = valid_segments[0]["_target_language"]
        source_text = "\n\n".join(seg["_source_text"] for seg in valid_segments)
        target_text = "\n\n".join(seg["_target_text"] for seg in valid_segments)
        speakers = sorted({str(seg.get("speaker", "")).strip() for seg in valid_segments if str(seg.get("speaker", "")).strip()})
        segment_types = sorted({str(seg.get("segment_type", "")).strip() for seg in valid_segments if str(seg.get("segment_type", "")).strip()})
        return [{
            "chunk_index": 1,
            "chunk_label": "Full passage",
            "source_language": source_lang,
            "target_language": target_lang,
            "source_text": source_text,
            "target_text": target_text,
            "speaker": ", ".join(speakers[:6]),
            "segment_type": ", ".join(segment_types[:6]),
            "segments": valid_segments,
        }]

    for pos, seg in enumerate(valid_segments, start=1):
        chunks.append({
            "chunk_index": pos,
            "chunk_label": f"Chunk {pos} / JSON segment {seg.get('_segment_index')}",
            "source_language": seg["_source_language"],
            "target_language": seg["_target_language"],
            "source_text": seg["_source_text"],
            "target_text": seg["_target_text"],
            "speaker": str(seg.get("speaker", "")),
            "segment_type": str(seg.get("segment_type", "")),
            "segments": [seg],
        })
    return chunks


def get_simultaneous_audio_path(cache_dir: Path, script_id: str, chunk_index: Any, role: str, lang: str, text: str, voice: str, rate: str) -> Path:
    key = short_hash(script_id, str(chunk_index), role, lang, text, voice, rate)
    filename = f"{sanitize_filename_part(script_id)}_sim_{sanitize_filename_part(str(chunk_index), 20)}_{role}_{lang}_{key}.mp3"
    return cache_dir / filename


# -----------------------------
# Source-only full-passage audio helpers
# -----------------------------
def _normalize_source_only_speaker(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip()).upper()


def _clean_source_only_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def make_silence_mp3_for_app(out_path: Path, seconds: float) -> Path:
    """Create a small MP3 silence file for app-side concatenation."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    seconds = max(0.01, float(seconds or 0.01))
    if out_path.exists() and out_path.stat().st_size > 0:
        return out_path
    if not shutil.which("ffmpeg"):
        raise FileNotFoundError("ffmpeg was not found on PATH; cannot create timing gaps.")
    subprocess.run([
        "ffmpeg", "-y", "-hide_banner", "-loglevel", "error",
        "-f", "lavfi", "-i", "anullsrc=r=24000:cl=mono",
        "-t", f"{seconds:.3f}", "-q:a", "9", "-acodec", "libmp3lame", str(out_path)
    ], check=True)
    return out_path


def build_source_only_full_passage_groups(
    data: dict[str, Any],
    segments: list[dict[str, Any]],
    speed_mode_value: str,
    use_audio_profile: bool,
    max_merged_chars: int,
    direction: str = "Source language first",
) -> list[dict[str, Any]]:
    """Build source-only groups using the same idea as the generator's _source_only export.

    Adjacent source chunks are merged only when they have the same speaker, source
    language, resolved TTS voice, and rate. This avoids artificial pauses caused by
    JSON chunking while preserving natural pauses for speaker/language/voice changes.
    """
    entries: list[dict[str, Any]] = []
    for idx, seg in enumerate(segments or [], start=1):
        if not isinstance(seg, dict):
            continue
        source_lang, _target_lang = get_segment_source_and_target(seg, direction)
        source_text = _clean_source_only_text(seg.get(source_lang, ""))
        if not source_text:
            continue
        voice = get_practice_voice(data, seg, source_lang, use_audio_profile)
        rate = get_rate_for_language(speed_mode_value, source_lang)
        entries.append({
            "segment_index": seg.get("_segment_index") or seg.get("segment_number") or idx,
            "speaker": str(seg.get("speaker") or ""),
            "language": source_lang,
            "text": source_text,
            "voice": voice,
            "rate": rate,
            "segment_type": str(seg.get("segment_type") or ""),
        })

    groups: list[dict[str, Any]] = []
    max_chars = max(200, int(max_merged_chars or 2800))
    for entry in entries:
        if not groups:
            new_group = dict(entry)
            new_group["segments"] = [entry.get("segment_index")]
            new_group["break_kind"] = "start"
            groups.append(new_group)
            continue

        prev = groups[-1]
        same_speaker = _normalize_source_only_speaker(prev.get("speaker")) == _normalize_source_only_speaker(entry.get("speaker"))
        same_language = str(prev.get("language", "")).lower() == str(entry.get("language", "")).lower()
        same_voice = _clean_source_only_text(prev.get("voice")) == _clean_source_only_text(entry.get("voice"))
        same_rate = _clean_source_only_text(prev.get("rate")) == _clean_source_only_text(entry.get("rate"))
        merged_text = _clean_source_only_text(str(prev.get("text", "")) + " " + str(entry.get("text", "")))

        if same_speaker and same_language and same_voice and same_rate and len(merged_text) <= max_chars:
            prev["text"] = merged_text
            prev.setdefault("segments", []).append(entry.get("segment_index"))
            continue

        new_group = dict(entry)
        new_group["segments"] = [entry.get("segment_index")]
        if same_speaker and same_language and same_voice and same_rate:
            new_group["break_kind"] = "continuation"
        elif same_speaker:
            new_group["break_kind"] = "same_speaker"
        else:
            new_group["break_kind"] = "speaker_change"
        groups.append(new_group)
    return groups


async def generate_source_only_full_passage_audio(
    data: dict[str, Any],
    segments: list[dict[str, Any]],
    script_id: str,
    cache_dir: Path,
    role: str,
    speed_mode_value: str,
    use_audio_profile: bool,
    same_speaker_gap: float,
    speaker_change_gap: float,
    continuation_gap: float,
    max_merged_chars: int,
    generator_path: Path | None = None,
    direction: str = "Source language first",
) -> tuple[Path | None, list[dict[str, Any]]]:
    """Generate a natural full-passage source-only MP3 for practice/exam modes."""
    cache_dir.mkdir(parents=True, exist_ok=True)
    groups = build_source_only_full_passage_groups(
        data=data,
        segments=segments,
        speed_mode_value=speed_mode_value,
        use_audio_profile=use_audio_profile,
        max_merged_chars=max_merged_chars,
        direction=direction,
    )
    if not groups:
        return None, []

    signature = short_hash(
        script_id, role, speed_mode_value, str(use_audio_profile),
        str(same_speaker_gap), str(speaker_change_gap), str(continuation_gap), str(max_merged_chars),
        json.dumps([
            {
                "speaker": g.get("speaker", ""), "language": g.get("language", ""),
                "voice": g.get("voice", ""), "rate": g.get("rate", ""),
                "text": g.get("text", ""), "break_kind": g.get("break_kind", ""),
            }
            for g in groups
        ], ensure_ascii=False, sort_keys=True),
    )
    out_path = cache_dir / f"{sanitize_filename_part(script_id)}_{sanitize_filename_part(role, 40)}_source_only_full_{signature}.mp3"
    if out_path.exists() and out_path.stat().st_size > 0:
        return out_path, groups

    parts: list[Path] = []
    gap_dir = cache_dir / "gaps"
    same_gap_path = make_silence_mp3_for_app(gap_dir / f"same_{float(same_speaker_gap):.2f}.mp3", same_speaker_gap)
    speaker_gap_path = make_silence_mp3_for_app(gap_dir / f"speaker_{float(speaker_change_gap):.2f}.mp3", speaker_change_gap)
    continuation_gap_path = make_silence_mp3_for_app(gap_dir / f"continuation_{float(continuation_gap):.2f}.mp3", continuation_gap)

    for group_idx, group in enumerate(groups, start=1):
        if parts:
            break_kind = group.get("break_kind")
            if break_kind == "speaker_change":
                parts.append(speaker_gap_path)
            elif break_kind == "continuation":
                parts.append(continuation_gap_path)
            else:
                parts.append(same_gap_path)
        group_audio = cache_dir / f"{sanitize_filename_part(script_id)}_{sanitize_filename_part(role, 40)}_group_{group_idx:03d}_{short_hash(group.get('text',''), group.get('voice',''), group.get('rate',''))}.mp3"
        await generate_practice_audio_file(str(group.get("text", "")), str(group.get("voice", "")), str(group.get("rate", "")), group_audio, generator_path)
        parts.append(group_audio)

    combined = ffmpeg_concat_mp3s(parts, out_path)
    return combined, groups


def reset_simultaneous_state(reason: str = "") -> None:
    st.session_state.simultaneous_idx = 0
    st.session_state.simultaneous_revealed = False
    st.session_state.simultaneous_signature = None
    st.session_state.simultaneous_source_nonce = 0
    st.session_state.simultaneous_target_nonce = 0
    st.session_state.simultaneous_started_at = time.time()
    st.session_state.simultaneous_lag_started = False
    st.session_state.simultaneous_lag_start_time = None
    st.session_state.simultaneous_reason = reason
    st.session_state.simultaneous_session_id = f"sim-session-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{uuid.uuid4().hex[:8]}"

def reset_term_review_state(reason: str = "") -> None:
    st.session_state.term_review_idx = 0
    st.session_state.term_review_revealed = False
    st.session_state.term_review_started_at = time.time()
    st.session_state.term_review_reason = reason
    st.session_state.term_review_session_id = f"term-session-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{uuid.uuid4().hex[:8]}"
    st.session_state.term_review_shuffle_seed = uuid.uuid4().hex[:8]

def reset_practice_state(reason: str = "") -> None:
    st.session_state.practice_idx = 0
    st.session_state.practice_revealed = False
    st.session_state.practice_segment_signature = None
    st.session_state.source_replay_nonce = 0
    st.session_state.target_replay_nonce = 0
    st.session_state.target_reveal_nonce = 0
    st.session_state.practice_started_at = time.time()
    st.session_state.practice_reason = reason
    st.session_state.practice_session_id = f"session-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{uuid.uuid4().hex[:8]}"

st.set_page_config(page_title="Interpreter Audio Generator", page_icon="🎧", layout="wide")
inject_phase2e_css()
st.title("🎧 Interpreter Audio Generator")
st.caption("Cloud version: generated audio, recordings, transcripts, drill history, and ZIPs are temporary. Download anything you want to keep.")
st.caption("Local Streamlit dashboard for audio generation, practice, exam review, feedback, and interactive personalized drills.")

with st.sidebar:
    st.header("Generator")
    generator_path = Path(st.text_input("Generator script path", value=str(DEFAULT_GENERATOR_PATH))).expanduser()
    output_dir = Path(st.text_input("Output folder", value=str(RUNTIME_OUTPUT_DIR.resolve()))).expanduser()
    st.divider()
    st.header("Audio settings")
    speed_mode = st.selectbox("Speed mode", ["learning", "normal", "fast"], index=0)
    split_flashcards = st.checkbox("Split flashcards and script", value=True)
    include_flashcards = st.checkbox("Include flashcards", value=True)
    consecutive_mode = st.checkbox("Generate consecutive mode", value=True)
    source_only_mode = st.checkbox(
        "Generate source-only practice audio",
        value=True,
        help="Exports <script_id>_source_only.mp3 with only the original/source language side of each segment, preserving JSON speaker voices.",
    )
    full_speed = st.checkbox("Generate full speed-increase drill", value=True)
    targeted_terms = st.checkbox("Generate targeted hard-terms drill", value=False)
    inline_cues = st.checkbox("Generate inline term-cue drill", value=False)
    combined_all = st.checkbox("Generate combined audio from all script outputs", value=False)
    skip_existing = st.checkbox("Skip existing up-to-date MP3s", value=True)
    verify_voices = st.checkbox("Verify Edge TTS voices", value=True)

    st.subheader("Consecutive mode timing")
    consecutive_ratio = st.number_input("Pause ratio", min_value=0.1, max_value=5.0, value=1.35, step=0.05)
    consecutive_min = st.number_input("Minimum pause", min_value=0.0, max_value=60.0, value=3.0, step=0.5)
    consecutive_max = st.number_input("Maximum pause", min_value=0.5, max_value=120.0, value=14.0, step=0.5)
    consecutive_gap = st.number_input("Segment gap", min_value=0.0, max_value=20.0, value=1.8, step=0.1)

    st.subheader("Source-only timing")
    source_only_gap = st.number_input(
        "Source-only same-speaker gap",
        min_value=0.0,
        max_value=10.0,
        value=0.75,
        step=0.05,
        help="Pause between source-only chunks when the speaker does not change.",
    )
    source_only_speaker_gap = st.number_input(
        "Source-only speaker-change gap",
        min_value=0.0,
        max_value=10.0,
        value=1.35,
        step=0.05,
        help="Slightly longer natural pause when the JSON speaker changes, such as Judge to attorney.",
    )
    source_only_merge_same_speaker = st.checkbox(
        "Merge adjacent same-speaker source chunks",
        value=True,
        help="For source-only output, combines adjacent JSON segments when the same speaker continues in the same source language with the same voice. This removes unnatural pauses caused only by JSON chunking.",
    )
    source_only_max_merged_chars = st.number_input(
        "Max merged source-only block size",
        min_value=500,
        max_value=8000,
        value=2800,
        step=100,
        help="Safety limit for one merged TTS block. Larger values create longer same-speaker monologues; smaller values create more continuation breaks.",
    )
    source_only_continuation_gap = st.number_input(
        "Source-only continuation gap",
        min_value=0.0,
        max_value=3.0,
        value=0.18,
        step=0.01,
        help="Tiny pause if a same-speaker block must be split because it exceeds the max merged size.",
    )

    st.subheader("Interactive practice audio")
    practice_audio_enabled = st.checkbox("Enable practice audio", value=True)
    practice_audio_autoplay_source = st.checkbox("Auto-play source audio when segment loads", value=True)
    practice_audio_play_correction = st.checkbox("Play correction audio when revealing target", value=True)
    practice_audio_show_replay = st.checkbox("Show replay audio controls", value=True)
    practice_audio_use_profile = st.checkbox("Use JSON audio_profile speaker voices", value=True)
    audio_only_source_mode = st.checkbox("Consecutive audio-only source mode", value=False)
    consecutive_playback_speed = st.selectbox(
        "Consecutive playback speed",
        [0.85, 1.0, 1.10, 1.20, 1.30, 1.40, 1.50],
        index=1,
        format_func=lambda x: f"{x:.2f}x",
        key="consecutive_playback_speed",
        help="Applies to both the source audio and the revealed correction audio in Consecutive Practice.",
    )
    practice_audio_cache_dir = Path(st.text_input("Practice audio cache folder", value=str((RUNTIME_CACHE_DIR / "practice_audio_cache").resolve()))).expanduser()


    st.divider()
    st.header("Study progress")
    persistent_log_path = Path(st.text_input("Practice history CSV", value=str(DEFAULT_PRACTICE_LOG_PATH.resolve()))).expanduser()
    auto_save_practice_log = st.checkbox("Auto-save scored segments to study history", value=True)
    simultaneous_log_path = Path(st.text_input("Simultaneous practice history CSV", value=str(DEFAULT_SIMULTANEOUS_LOG_PATH.resolve()))).expanduser()
    auto_save_simultaneous_log = st.checkbox("Auto-save simultaneous practice attempts", value=True)
    term_review_log_path = Path(st.text_input("Term review history CSV", value=str(DEFAULT_TERM_REVIEW_LOG_PATH.resolve()))).expanduser()
    term_srs_state_path = Path(st.text_input("Term SRS state CSV", value=str(DEFAULT_TERM_SRS_STATE_PATH.resolve()))).expanduser()
    drill_studio_log_path = Path(st.text_input("Drill Studio history CSV", value=str(DEFAULT_DRILL_STUDIO_LOG_PATH.resolve()))).expanduser()
    auto_save_term_review_log = st.checkbox("Auto-save term review scores", value=True)

    st.subheader("Spoken term review audio")
    term_audio_enabled = st.checkbox("Enable term review audio", value=True)
    term_audio_autoplay_prompt = st.checkbox("Auto-play prompt term", value=True)
    term_audio_play_answer_on_reveal = st.checkbox("Play answer audio when revealing target", value=True)
    term_audio_show_replay = st.checkbox("Show term replay audio controls", value=True)
    term_audio_only_prompt = st.checkbox("Audio-only prompt mode", value=False)
    term_audio_cache_dir = Path(st.text_input("Term audio cache folder", value=str((RUNTIME_CACHE_DIR / "term_audio_cache").resolve()))).expanduser()

    show_history_preview_rows = st.number_input("History rows to preview", min_value=10, max_value=1000, value=100, step=10)

tab_generate, tab_practice, tab_simul, tab_exam, tab_drills, tab_history, tab_terms = st.tabs(["Generate outputs", "Consecutive practice", "Simultaneous practice", "Exam mode", "Drill Studio", "Study history", "Term review"])

with tab_generate:
    st.subheader("1. Choose source JSON files")
    source_mode = st.radio("Source", ["Upload JSON file(s)", "Use a local folder path"], horizontal=True)

    loaded_files: list[tuple[str, bytes]] = []
    if source_mode == "Upload JSON file(s)":
        uploads = st.file_uploader("Select one or more JSON scripts", type=["json"], accept_multiple_files=True)
        for uploaded in uploads or []:
            loaded_files.append((uploaded.name, uploaded.getvalue()))
    else:
        local_dir = Path(st.text_input("Local JSON folder", value=str((APP_DIR / "scripts").resolve()))).expanduser()
        pattern = st.text_input("Filename pattern", value="*.json")
        if local_dir.exists():
            paths = sorted(local_dir.glob(pattern))
            selected_names = st.multiselect("Files found", [p.name for p in paths], default=[p.name for p in paths[:10]])
            for p in paths:
                if p.name in selected_names:
                    loaded_files.append((p.name, p.read_bytes()))
        else:
            st.warning("That local folder does not exist yet.")

    payloads: list[tuple[str, dict[str, Any], bytes]] = []
    parse_errors: list[str] = []
    for name, raw in loaded_files:
        data, error = load_json_bytes(name, raw)
        if error:
            parse_errors.append(error)
        elif data is not None:
            payloads.append((name, data, raw))

    if parse_errors:
        for err in parse_errors:
            st.error(err)

    if payloads:
        st.subheader("2. Preview and validate")
        summaries = [summarize_payload(data, name) for name, data, _ in payloads]
        st.dataframe(summaries, use_container_width=True, hide_index=True)

        with st.expander("Validation details", expanded=True):
            for name, data, _ in payloads:
                warnings, errors = validate_json_payload(data)
                label = f"{name}: {len(errors)} error(s), {len(warnings)} warning(s)"
                with st.expander(label, expanded=bool(errors)):
                    if errors:
                        for err in errors:
                            st.error(err)
                    if warnings:
                        for warning in warnings[:50]:
                            st.warning(warning)
                        if len(warnings) > 50:
                            st.info(f"Showing first 50 warnings out of {len(warnings)}.")
                    if not errors and not warnings:
                        st.success("No validation issues found.")

        st.subheader("3. Run")
        if source_only_mode:
            if source_only_merge_same_speaker:
                st.info("Source-only export will create `<script_id>_source_only.mp3`: only the original/source-language side. Adjacent chunks from the same speaker/source language will be merged for more natural monologues; speaker changes keep a longer pause.")
            else:
                st.info("Source-only export will create `<script_id>_source_only.mp3`: only the original/source-language side of each paired segment, preserving speaker voices and adding a longer pause when the speaker changes.")
        settings = {
            "SPEED_MODE": speed_mode,
            "SPLIT_FLASHCARDS_AND_SCRIPT": split_flashcards,
            "INCLUDE_FLASHCARDS": include_flashcards,
            "GENERATE_CONSECUTIVE_MODE": consecutive_mode,
            "GENERATE_SOURCE_ONLY_AUDIO": source_only_mode,
            "SOURCE_ONLY_SEGMENT_GAP": float(source_only_gap),
            "SOURCE_ONLY_SPEAKER_CHANGE_GAP": float(source_only_speaker_gap),
            "SOURCE_ONLY_MERGE_ADJACENT_SAME_SPEAKER": bool(source_only_merge_same_speaker),
            "SOURCE_ONLY_MAX_MERGED_CHARS": int(source_only_max_merged_chars),
            "SOURCE_ONLY_MERGED_CONTINUATION_GAP": float(source_only_continuation_gap),
            "CONSECUTIVE_PAUSE_RATIO": float(consecutive_ratio),
            "CONSECUTIVE_MIN_PAUSE": float(consecutive_min),
            "CONSECUTIVE_MAX_PAUSE": float(consecutive_max),
            "CONSECUTIVE_SEGMENT_GAP": float(consecutive_gap),
            "GENERATE_FULL_SPEED_INCREASE": full_speed,
            "GENERATE_TARGETED_HARD_TERMS": targeted_terms,
            "GENERATE_INLINE_TERM_CUES": inline_cues,
            "GENERATE_COMBINED_AUDIO_ALL_SCRIPTS": combined_all,
            "SKIP_EXISTING_SCRIPT_MP3S": skip_existing,
            "VERIFY_TTS_VOICES": verify_voices,
        }

        disabled = not generator_path.exists() or not payloads
        if not generator_path.exists():
            st.error("Generator script path does not exist.")

        if st.button("Run audio generator", type="primary", disabled=disabled):
            output_dir.mkdir(parents=True, exist_ok=True)
            with tempfile.TemporaryDirectory(prefix="streamlit_audio_json_") as tmp:
                tmpdir = Path(tmp)
                json_paths: list[Path] = []
                for name, _data, raw in payloads:
                    path = tmpdir / Path(name).name
                    path.write_bytes(raw)
                    json_paths.append(path)

                log_buffer = io.StringIO()
                try:
                    with st.spinner("Generating audio..."):
                        with contextlib.redirect_stdout(log_buffer), contextlib.redirect_stderr(log_buffer):
                            outputs = asyncio.run(run_generator(generator_path, json_paths, output_dir, settings))
                    st.success(f"Generated {len(outputs)} output file(s).")
                    if outputs:
                        st.write("Output files:")
                        for p in outputs:
                            st.code(str(p))
                    with st.expander("Generator log", expanded=False):
                        st.text(log_buffer.getvalue() or "No log output.")
                except Exception as exc:
                    st.error(f"Generator failed: {exc}")
                    with st.expander("Generator log", expanded=True):
                        st.text(log_buffer.getvalue() or "No log output.")
    else:
        st.info("Choose one or more JSON files to preview settings and run the generator.")

with tab_practice:
    st.subheader("Interactive consecutive practice")
    st.caption("Practice segment by segment: read/listen to the source side, interpret aloud, reveal the correction, score yourself, and automatically save long-term progress.")
    st.caption(f"Current practice session: `{st.session_state.get('practice_session_id', 'not started')}`")

    with st.expander("Consecutive audio controls", expanded=True):
        ac1, ac2 = st.columns([1, 1])
        with ac1:
            consecutive_playback_speed = st.selectbox(
                "Playback speed",
                [0.85, 1.0, 1.10, 1.20, 1.30, 1.40, 1.50],
                index=1,
                format_func=lambda x: f"{x:.2f}x",
                key="consecutive_playback_speed_inline",
                help="Applies to both the source audio and the revealed correction audio in Consecutive Practice.",
            )
        with ac2:
            audio_only_source_mode = st.checkbox(
                "Audio-only source mode",
                value=False,
                key="consecutive_audio_only_source_mode_inline",
                help="Hide the source text so you practice from the audio first.",
            )
        st.caption("These controls affect only the Consecutive Practice tab.")

    if "practice_log" not in st.session_state:
        st.session_state.practice_log = []
    if "practice_session_id" not in st.session_state:
        st.session_state.practice_session_id = f"session-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{uuid.uuid4().hex[:8]}"
    if "practice_idx" not in st.session_state:
        reset_practice_state("initial")

    practice_files: list[tuple[str, bytes]] = []
    practice_source = st.radio("Practice source", ["Upload JSON file", "Use selected files from Generate tab", "Use local file path"], horizontal=True)
    if practice_source == "Upload JSON file":
        uploaded_practice = st.file_uploader("Choose one JSON for practice", type=["json"], accept_multiple_files=False, key="practice_upload")
        if uploaded_practice:
            practice_files.append((uploaded_practice.name, uploaded_practice.getvalue()))
    elif practice_source == "Use selected files from Generate tab":
        practice_files = loaded_files[:]
        if not practice_files:
            st.info("Select files in the Generate outputs tab first, or use Upload JSON file here.")
    else:
        practice_path = Path(st.text_input("Practice JSON path", value="")).expanduser()
        if practice_path.exists() and practice_path.suffix.lower() == ".json":
            practice_files.append((practice_path.name, practice_path.read_bytes()))
        elif str(practice_path):
            st.warning("That JSON path does not exist.")

    practice_payloads: list[tuple[str, dict[str, Any]]] = []
    for name, raw in practice_files:
        data, error = load_json_bytes(name, raw)
        if error:
            st.error(error)
        elif data:
            practice_payloads.append((name, data))

    if practice_payloads:
        file_options = [f"{safe_script_id(data, Path(name).stem)} — {data.get('title', name)}" for name, data in practice_payloads]
        selected_label = st.selectbox("Script", file_options)
        selected_idx = file_options.index(selected_label)
        selected_name, selected_data = practice_payloads[selected_idx]
        script_id = safe_script_id(selected_data, Path(selected_name).stem)
        title = str(selected_data.get("title", selected_name))

        paired = selected_data.get("paired_segments", []) or []
        if not isinstance(paired, list) or not paired:
            st.warning("This JSON does not contain paired_segments for interactive practice.")
        else:
            all_types = sorted({str(seg.get("segment_type", "")).strip() for seg in paired if isinstance(seg, dict) and str(seg.get("segment_type", "")).strip()})
            all_speakers = sorted({str(seg.get("speaker", "")).strip() for seg in paired if isinstance(seg, dict) and str(seg.get("speaker", "")).strip()})

            c1, c2, c3 = st.columns(3)
            with c1:
                direction = st.selectbox("Practice direction", ["Source language first", "English → Spanish only", "Spanish → English only"])
            with c2:
                selected_types = st.multiselect("Segment types", all_types, default=[])
            with c3:
                selected_speakers = st.multiselect("Speakers", all_speakers, default=[])

            randomize = st.checkbox("Randomize segment order", value=False)
            segments = filter_practice_segments(selected_data, direction, selected_types, selected_speakers)
            if randomize:
                # Stable deterministic randomization per script/load so Streamlit reruns do not jump unexpectedly.
                import random
                rng = random.Random(script_id + title + direction)
                segments = segments[:]
                rng.shuffle(segments)

            if not segments:
                st.warning("No segments match the current filters.")
            else:
                st.write(f"Practice set: **{len(segments)} segment(s)**")
                if st.button("Start / reset this practice set"):
                    reset_practice_state("manual reset")

                st.session_state.practice_idx = min(st.session_state.practice_idx, len(segments) - 1)
                idx = st.session_state.practice_idx
                seg = segments[idx]
                source_lang = seg["_source_language"]
                target_lang = seg["_target_language"]
                source_text = str(seg.get(source_lang, "")).strip()
                target_text = str(seg.get(target_lang, "")).strip()

                # Detect an actual segment change and reset audio/reveal state. This prevents
                # the source player from reusing the previous segment's browser audio element.
                segment_signature = short_hash(
                    script_id,
                    str(idx),
                    str(seg.get("_segment_index")),
                    source_lang,
                    target_lang,
                    source_text,
                    target_text,
                )
                if st.session_state.get("practice_segment_signature") != segment_signature:
                    st.session_state.practice_segment_signature = segment_signature
                    st.session_state.practice_revealed = False
                    st.session_state.source_replay_nonce = st.session_state.get("source_replay_nonce", 0) + 1
                    st.session_state.target_replay_nonce = 0
                    st.session_state.target_reveal_nonce = 0

                pause_estimate = estimate_dynamic_pause_seconds(source_text, float(consecutive_ratio), float(consecutive_min), float(consecutive_max))

                st.progress((idx + 1) / len(segments))
                st.markdown(f"### Segment {idx + 1} of {len(segments)}")
                render_compact_segment_meta(
                    primary_label=f"JSON segment {seg.get('_segment_index')}",
                    source_lang=source_lang,
                    target_lang=target_lang,
                    word_count=len(source_text.split()),
                    speaker=str(seg.get("speaker") or ""),
                    segment_type=str(seg.get("segment_type") or ""),
                    extra=f"pause ~{pause_estimate}s",
                )

                st.markdown("#### Practice workspace")
                render_side_by_side_text(
                    source_text=source_text,
                    target_text=target_text,
                    source_lang=source_lang,
                    target_lang=target_lang,
                    speaker=str(seg.get('speaker', '')),
                    segment_type=str(seg.get('segment_type', '')),
                    revealed=bool(st.session_state.practice_revealed),
                    source_visible=not bool(audio_only_source_mode),
                    context_label="Correction / target",
                )
                st.caption("Interpret this aloud before revealing the correction.")

                source_audio_path = None
                target_audio_path = None
                if practice_audio_enabled:
                    source_voice = get_practice_voice(selected_data, seg, source_lang, practice_audio_use_profile)
                    target_voice = get_practice_voice(selected_data, seg, target_lang, practice_audio_use_profile)
                    source_rate = get_rate_for_language(speed_mode, source_lang)
                    target_rate = get_rate_for_language(speed_mode, target_lang)
                    source_audio_path = get_practice_audio_path(practice_audio_cache_dir, script_id, seg.get("_segment_index", idx + 1), "source", source_lang, source_text, source_voice, source_rate)
                    target_audio_path = get_practice_audio_path(practice_audio_cache_dir, script_id, seg.get("_segment_index", idx + 1), "target", target_lang, target_text, target_voice, target_rate)

                    try:
                        with st.spinner("Preparing source audio..."):
                            asyncio.run(generate_practice_audio_file(source_text, source_voice, source_rate, source_audio_path, generator_path))
                        audio_cols = st.columns([2, 1, 1])
                        with audio_cols[0]:
                            render_audio_player(
                                source_audio_path,
                                label=f"Source audio — segment {seg.get('_segment_index')} — {source_lang.title()} — voice: {source_voice}",
                                autoplay=practice_audio_autoplay_source and not st.session_state.practice_revealed,
                                show_native_player=practice_audio_show_replay,
                                playback_speed=float(consecutive_playback_speed),
                                nonce=f"source-{segment_signature}-{st.session_state.get('source_replay_nonce', 0)}",
                            )
                            render_wpm_estimate("Source", source_text, source_audio_path, float(consecutive_playback_speed))
                        with audio_cols[1]:
                            if practice_audio_show_replay and st.button("Replay source audio", key=f"replay_source_{segment_signature}"):
                                st.session_state.source_replay_nonce = st.session_state.get("source_replay_nonce", 0) + 1
                                st.rerun()
                        with audio_cols[2]:
                            st.caption(f"Source voice: {source_voice}")
                    except Exception as exc:
                        st.warning(f"Could not prepare source audio: {exc}")

                reveal_clicked = st.button("Reveal correction / target text", type="primary")
                if reveal_clicked:
                    st.session_state.practice_revealed = True
                    st.session_state.target_reveal_nonce = st.session_state.get("target_reveal_nonce", 0) + 1
                    # Re-run immediately so the side-by-side reference panel above
                    # switches from hidden to visible on the first click, while the
                    # correction audio can still autoplay in the revealed state.
                    st.rerun()

                if st.session_state.practice_revealed:
                    st.caption("Correction is shown in the right-hand reference panel above.")

                    if practice_audio_enabled and target_audio_path is not None:
                        try:
                            target_voice = get_practice_voice(selected_data, seg, target_lang, practice_audio_use_profile)
                            target_rate = get_rate_for_language(speed_mode, target_lang)
                            with st.spinner("Preparing correction audio..."):
                                asyncio.run(generate_practice_audio_file(target_text, target_voice, target_rate, target_audio_path, generator_path))
                            correction_cols = st.columns([2, 1, 1])
                            with correction_cols[0]:
                                render_audio_player(
                                    target_audio_path,
                                    label=f"Correction audio — segment {seg.get('_segment_index')} — {target_lang.title()} — voice: {target_voice}",
                                    autoplay=practice_audio_play_correction,
                                    show_native_player=practice_audio_show_replay,
                                    playback_speed=float(consecutive_playback_speed),
                                    nonce=f"target-{segment_signature}-{st.session_state.get('target_reveal_nonce', 0)}-{st.session_state.get('target_replay_nonce', 0)}",
                                )
                                render_wpm_estimate("Correction", target_text, target_audio_path, float(consecutive_playback_speed))
                            with correction_cols[1]:
                                if practice_audio_show_replay and st.button("Replay correction audio", key=f"replay_target_{segment_signature}"):
                                    st.session_state.target_replay_nonce = st.session_state.get("target_replay_nonce", 0) + 1
                                    st.rerun()
                            with correction_cols[2]:
                                st.caption(f"Correction voice: {target_voice}")
                        except Exception as exc:
                            st.warning(f"Could not prepare correction audio: {exc}")

                    st.markdown("#### Self-score")
                    score = st.radio(
                        "Score",
                        options=[3, 2, 1, 0],
                        format_func=lambda x: {
                            3: "3 — Accurate",
                            2: "2 — Minor issue",
                            1: "1 — Major issue",
                            0: "0 — Missed / could not complete",
                        }[x],
                        horizontal=True,
                        key=f"score_{script_id}_{idx}_{seg.get('_segment_index')}",
                    )
                    error_tags = st.multiselect("Error tags", ERROR_TAGS, key=f"tags_{script_id}_{idx}_{seg.get('_segment_index')}")
                    notes = st.text_area("Notes", key=f"notes_{script_id}_{idx}_{seg.get('_segment_index')}", height=90)

                    col_prev, col_save, col_next = st.columns([1, 2, 1])
                    with col_prev:
                        if st.button("← Previous", disabled=idx == 0):
                            st.session_state.practice_idx -= 1
                            st.session_state.practice_revealed = False
                            st.rerun()
                    with col_save:
                        if st.button("Save score and next", type="primary"):
                            log_row = {
                                "timestamp": datetime.now().isoformat(timespec="seconds"),
                                "session_id": st.session_state.get("practice_session_id", ""),
                                "script_id": script_id,
                                "title": title,
                                "practice_direction": direction,
                                "practice_set_size": len(segments),
                                "segment_index": seg.get("_segment_index"),
                                "speaker": seg.get("speaker", ""),
                                "segment_type": seg.get("segment_type", ""),
                                "source_language": source_lang,
                                "target_language": target_lang,
                                "score": score,
                                "error_tags": "; ".join(error_tags),
                                "notes": notes,
                                "source_text": source_text,
                                "target_text": target_text,
                            }
                            st.session_state.practice_log.append(log_row)
                            if auto_save_practice_log:
                                try:
                                    append_persistent_practice_log(persistent_log_path, log_row)
                                except Exception as exc:
                                    st.warning(f"Could not auto-save to study history: {exc}")
                            if idx < len(segments) - 1:
                                st.session_state.practice_idx += 1
                                st.session_state.practice_revealed = False
                            st.rerun()
                    with col_next:
                        if st.button("Skip →", disabled=idx >= len(segments) - 1):
                            st.session_state.practice_idx += 1
                            st.session_state.practice_revealed = False
                            st.rerun()
                else:
                    nav1, nav2 = st.columns([1, 1])
                    with nav1:
                        if st.button("← Previous segment", disabled=idx == 0):
                            st.session_state.practice_idx -= 1
                            st.session_state.practice_revealed = False
                            st.rerun()
                    with nav2:
                        if st.button("Skip segment →", disabled=idx >= len(segments) - 1):
                            st.session_state.practice_idx += 1
                            st.session_state.practice_revealed = False
                            st.rerun()

                st.divider()
                st.markdown("### Session log")
                log_rows = st.session_state.practice_log
                if log_rows:
                    st.dataframe(log_rows, use_container_width=True, hide_index=True)
                    scores = [int(r.get("score", 0)) for r in log_rows if str(r.get("score", "")).isdigit()]
                    if scores:
                        st.write(f"Average score: **{sum(scores) / len(scores):.2f} / 3** across **{len(scores)} saved segment(s)**.")
                    if auto_save_practice_log:
                        st.caption(f"Scored segments are being saved to: `{persistent_log_path}`")
                    csv_text = practice_log_to_csv(log_rows)
                    st.download_button(
                        "Download session log CSV",
                        data=csv_text.encode("utf-8"),
                        file_name=f"{script_id}_practice_log.csv",
                        mime="text/csv",
                    )
                    if st.button("Clear session log"):
                        st.session_state.practice_log = []
                        st.rerun()
                else:
                    st.info("No scores saved yet in this app session.")
    else:
        st.info("Choose a JSON file to start interactive consecutive practice.")


with tab_simul:
    st.subheader("Simultaneous practice")
    st.caption("Listen to a continuous source passage, interpret aloud in real time, reveal the reference interpretation, then score and save your attempt.")
    st.caption(f"Current simultaneous session: `{st.session_state.get('simultaneous_session_id', 'not started')}`")

    if "simultaneous_log" not in st.session_state:
        st.session_state.simultaneous_log = []
    if "simultaneous_session_id" not in st.session_state:
        reset_simultaneous_state("initial")

    sim_files: list[tuple[str, bytes]] = []
    sim_source = st.radio("Simultaneous source", ["Upload JSON file", "Use selected files from Generate tab", "Use local file path"], horizontal=True, key="sim_source_radio")
    if sim_source == "Upload JSON file":
        uploaded_sim = st.file_uploader("Choose one simultaneous JSON", type=["json"], accept_multiple_files=False, key="sim_upload")
        if uploaded_sim:
            sim_files.append((uploaded_sim.name, uploaded_sim.getvalue()))
    elif sim_source == "Use selected files from Generate tab":
        sim_files = loaded_files[:]
        if not sim_files:
            st.info("Select files in the Generate outputs tab first, or upload a JSON here.")
    else:
        sim_path = Path(st.text_input("Simultaneous JSON path", value="", key="sim_path")).expanduser()
        if sim_path.exists() and sim_path.suffix.lower() == ".json":
            sim_files.append((sim_path.name, sim_path.read_bytes()))
        elif str(sim_path):
            st.warning("That JSON path does not exist.")

    sim_payloads: list[tuple[str, dict[str, Any]]] = []
    for name, raw in sim_files:
        data, error = load_json_bytes(name, raw)
        if error:
            st.error(error)
        elif data:
            sim_payloads.append((name, data))

    if sim_payloads:
        sim_options = [f"{safe_script_id(data, Path(name).stem)} — {data.get('title', name)}" for name, data in sim_payloads]
        sim_label = st.selectbox("Script", sim_options, key="sim_script_select")
        sim_selected_idx = sim_options.index(sim_label)
        sim_name, sim_data = sim_payloads[sim_selected_idx]
        sim_script_id = safe_script_id(sim_data, Path(sim_name).stem)
        sim_title = str(sim_data.get("title", sim_name))

        format_hint = str(sim_data.get("format", "")).lower()
        if "simultaneous" not in format_hint:
            st.info("This file does not say 'simultaneous' in its format field, but it can still be practiced here if it has paired_segments.")

        settings_cols = st.columns(4)
        with settings_cols[0]:
            sim_mode = st.selectbox("Practice chunk size", ["Paragraph / thought chunk", "Full passage"], key="sim_mode")
        with settings_cols[1]:
            sim_direction = st.selectbox("Direction", ["Source language first", "English → Spanish only", "Spanish → English only"], key="sim_direction")
        with settings_cols[2]:
            sim_audio_only = st.checkbox("Audio-only source mode", value=True, key="sim_audio_only")
        with settings_cols[3]:
            show_source_transcript = st.checkbox("Show source transcript", value=not sim_audio_only, key="sim_show_source")

        reveal_reference_audio = st.checkbox("Play reference audio when revealing target", value=True, key="sim_play_ref")
        show_reference_text_by_default = st.checkbox("Show reference text after reveal", value=True, key="sim_show_ref_text")

        with st.expander("Simultaneous training tools", expanded=True):
            tool_cols = st.columns(4)
            with tool_cols[0]:
                sim_training_mode = st.selectbox(
                    "Training mode",
                    ["Interpret", "Lag trainer", "Shadowing", "Shadow first, then interpret"],
                    key="sim_training_mode",
                )
            with tool_cols[1]:
                sim_lag_seconds = st.number_input("Lag target seconds", min_value=0, max_value=15, value=3, step=1, key="sim_lag_seconds")
            with tool_cols[2]:
                sim_playback_speed = st.selectbox("Playback speed", [0.85, 1.0, 1.10, 1.20, 1.30, 1.40, 1.50], index=1, format_func=lambda x: f"{x:.2f}x", key="sim_playback_speed")
            with tool_cols[3]:
                sim_repeat_loop = st.checkbox("Repeat-loop source audio", value=False, key="sim_repeat_loop")

            st.caption(
                "Lag trainer gives you a start cue after the selected delay. "
                "Shadowing means repeating the source language first; Shadow first, then interpret adds an interpretation pass after shadowing."
            )

        chunk_mode_internal = "Full passage" if sim_mode == "Full passage" else "Paragraph / thought chunk"
        chunks = build_simultaneous_chunks(sim_data, chunk_mode_internal, sim_direction)
        if not chunks:
            st.warning("No usable paired_segments were found for simultaneous practice.")
        else:
            st.write(f"Practice set: **{len(chunks)} chunk(s)**")
            if st.button("Start / reset simultaneous practice", key="sim_reset"):
                reset_simultaneous_state("manual reset")
                st.rerun()

            st.session_state.simultaneous_idx = min(st.session_state.simultaneous_idx, len(chunks) - 1)
            sim_idx = st.session_state.simultaneous_idx
            chunk = chunks[sim_idx]
            source_lang = str(chunk.get("source_language", "english"))
            target_lang = str(chunk.get("target_language", "spanish"))
            source_text = str(chunk.get("source_text", "")).strip()
            target_text = str(chunk.get("target_text", "")).strip()
            chunk_signature = short_hash(sim_script_id, str(sim_idx), chunk.get("chunk_label", ""), source_lang, target_lang, source_text, target_text)
            if st.session_state.get("simultaneous_signature") != chunk_signature:
                st.session_state.simultaneous_signature = chunk_signature
                st.session_state.simultaneous_revealed = False
                st.session_state.simultaneous_source_nonce = st.session_state.get("simultaneous_source_nonce", 0) + 1
                st.session_state.simultaneous_target_nonce = 0
                st.session_state.simultaneous_lag_started = False
                st.session_state.simultaneous_lag_start_time = None

            st.progress((sim_idx + 1) / len(chunks), text=f"Chunk {sim_idx + 1} of {len(chunks)}")
            chunk_label = chunk.get('chunk_label', f'Chunk {sim_idx + 1}')
            st.markdown(f"### {chunk_label}")
            render_compact_segment_meta(
                primary_label=str(chunk_label),
                source_lang=source_lang,
                target_lang=target_lang,
                word_count=len(source_text.split()),
                speaker=str(chunk.get("speaker") or ""),
                segment_type=str(chunk.get("segment_type") or ""),
            )

            # Use audio_profile speaker assignments for both paragraph chunks and full-passage mode.
            # For full-passage source audio, use the same natural source-only grouping logic
            # as the generator's _source_only export: adjacent same-speaker/source-language
            # chunks are merged, and speaker changes keep a natural pause. This avoids the
            # choppy pauses caused by JSON chunk boundaries during real-time practice.
            source_voice = get_practice_voice_for_chunk(sim_data, chunk, source_lang, practice_audio_use_profile)
            target_voice = get_practice_voice_for_chunk(sim_data, chunk, target_lang, practice_audio_use_profile)
            source_rate = get_rate_for_language(speed_mode, source_lang)
            target_rate = get_rate_for_language(speed_mode, target_lang)
            sim_cache_dir = practice_audio_cache_dir / "simultaneous"
            source_audio_path = get_simultaneous_audio_path(sim_cache_dir, sim_script_id, chunk.get("chunk_index", sim_idx + 1), "source", source_lang, source_text, source_voice, source_rate)
            target_audio_path = get_simultaneous_audio_path(sim_cache_dir, sim_script_id, chunk.get("chunk_index", sim_idx + 1), "reference", target_lang, target_text, target_voice, target_rate)
            source_group_count = None

            st.markdown("#### Source audio")
            if practice_audio_enabled:
                try:
                    with st.spinner("Preparing simultaneous source audio..."):
                        if sim_mode == "Full passage" and isinstance(chunk.get("segments"), list):
                            merged_audio, source_groups = asyncio.run(generate_source_only_full_passage_audio(
                                data=sim_data,
                                segments=chunk.get("segments") or [],
                                script_id=sim_script_id,
                                cache_dir=sim_cache_dir / "source_only_full_passage",
                                role="simultaneous_practice",
                                speed_mode_value=speed_mode,
                                use_audio_profile=practice_audio_use_profile,
                                same_speaker_gap=float(source_only_gap),
                                speaker_change_gap=float(source_only_speaker_gap),
                                continuation_gap=float(source_only_continuation_gap),
                                max_merged_chars=int(source_only_max_merged_chars),
                                generator_path=generator_path,
                                direction=sim_direction,
                            ))
                            if merged_audio:
                                source_audio_path = merged_audio
                                source_group_count = len(source_groups)
                            else:
                                asyncio.run(generate_practice_audio_file(source_text, source_voice, source_rate, source_audio_path, generator_path))
                        else:
                            asyncio.run(generate_practice_audio_file(source_text, source_voice, source_rate, source_audio_path, generator_path))
                    source_label = f"Simultaneous source — {source_lang.title()}"
                    if source_group_count:
                        source_label += f" — natural source-only full passage ({source_group_count} voice block(s))"
                    else:
                        source_label += f" — voice: {source_voice}"
                    render_simultaneous_audio_player(
                        source_audio_path,
                        label=source_label,
                        autoplay=practice_audio_autoplay_source and not st.session_state.simultaneous_revealed,
                        show_native_player=True,
                        nonce=f"sim-source-{chunk_signature}-{st.session_state.get('simultaneous_source_nonce', 0)}",
                        playback_speed=float(sim_playback_speed),
                        loop=bool(sim_repeat_loop),
                    )
                    render_wpm_estimate("Source", source_text, source_audio_path, float(sim_playback_speed))
                    if source_group_count:
                        st.caption("Full-passage source audio uses source-only grouping: adjacent same-speaker chunks are merged; speaker changes keep a natural pause.")
                    if st.button("↻ Replay source", key=f"sim_replay_source_{chunk_signature}"):
                        st.session_state.simultaneous_source_nonce = st.session_state.get("simultaneous_source_nonce", 0) + 1
                        st.rerun()
                except Exception as exc:
                    st.warning(f"Could not prepare source audio: {exc}")
            else:
                st.info("Practice audio is disabled in the sidebar. Enable practice audio to use this tab as intended.")

            st.markdown("#### Practice cue")
            if sim_training_mode == "Lag trainer":
                cue_cols = st.columns([1, 1, 2])
                with cue_cols[0]:
                    if st.button("Start lag timer", key=f"sim_start_lag_{chunk_signature}"):
                        st.session_state.simultaneous_lag_started = True
                        st.session_state.simultaneous_lag_start_time = time.time()
                        st.rerun()
                with cue_cols[1]:
                    if st.button("Reset lag timer", key=f"sim_reset_lag_{chunk_signature}"):
                        st.session_state.simultaneous_lag_started = False
                        st.session_state.simultaneous_lag_start_time = None
                        st.rerun()
                lag_start = st.session_state.get("simultaneous_lag_start_time")
                if st.session_state.get("simultaneous_lag_started") and lag_start:
                    elapsed = time.time() - float(lag_start)
                    remaining = max(0, int(sim_lag_seconds) - elapsed)
                    if remaining > 0:
                        st.warning(f"Hold your output. Begin interpreting in approximately {remaining:.1f} second(s).")
                    else:
                        st.success("Begin interpreting now. Maintain the lag while the source continues.")
                else:
                    st.info("Start the source audio, then press Start lag timer. Begin interpreting when the cue appears.")
            elif sim_training_mode == "Shadowing":
                st.info("Shadowing mode: repeat the source language aloud while listening. Reveal the reference only after the shadowing pass.")
            elif sim_training_mode == "Shadow first, then interpret":
                st.info("First pass: shadow the source language. Replay the source, then interpret into the target language on the second pass.")
            else:
                st.info("Interpret aloud in real time while the source audio plays.")

            st.markdown("#### Practice workspace")
            render_side_by_side_text(
                source_text=source_text,
                target_text=target_text,
                source_lang=source_lang,
                target_lang=target_lang,
                speaker=str(chunk.get('speaker') or ''),
                segment_type=str(chunk.get('segment_type') or ''),
                revealed=bool(st.session_state.simultaneous_revealed) and bool(show_reference_text_by_default),
                source_visible=bool(show_source_transcript),
                context_label="Reference interpretation",
            )

            st.markdown("#### Practice task")
            st.write("Use the selected training mode above. When finished, reveal the reference interpretation and score your performance.")

            if not st.session_state.simultaneous_revealed:
                if st.button("Reveal reference interpretation", type="primary", key=f"sim_reveal_{chunk_signature}"):
                    st.session_state.simultaneous_revealed = True
                    st.session_state.simultaneous_target_nonce = st.session_state.get("simultaneous_target_nonce", 0) + 1
                    st.rerun()
            else:
                if show_reference_text_by_default:
                    st.caption("Reference interpretation is shown in the right-hand panel above.")
                else:
                    with st.expander("Reference interpretation text", expanded=False):
                        st.success(target_text)

                if practice_audio_enabled:
                    try:
                        with st.spinner("Preparing reference audio..."):
                            asyncio.run(generate_practice_audio_file(target_text, target_voice, target_rate, target_audio_path, generator_path))
                        render_simultaneous_audio_player(
                            target_audio_path,
                            label=f"Reference interpretation — {target_lang.title()} — voice: {target_voice}",
                            autoplay=reveal_reference_audio,
                            show_native_player=True,
                            nonce=f"sim-target-{chunk_signature}-{st.session_state.get('simultaneous_target_nonce', 0)}",
                            playback_speed=float(sim_playback_speed),
                            loop=False,
                        )
                        render_wpm_estimate("Reference", target_text, target_audio_path, float(sim_playback_speed))
                        if st.button("↻ Replay reference", key=f"sim_replay_target_{chunk_signature}"):
                            st.session_state.simultaneous_target_nonce = st.session_state.get("simultaneous_target_nonce", 0) + 1
                            st.rerun()
                    except Exception as exc:
                        st.warning(f"Could not prepare reference audio: {exc}")

                st.markdown("#### Self-score")
                sim_score = st.radio(
                    "Score",
                    options=[3, 2, 1, 0],
                    format_func=lambda x: {3: "3 — Strong", 2: "2 — Usable with minor issues", 1: "1 — Major issues", 0: "0 — Lost / could not complete"}[x],
                    horizontal=True,
                    key=f"sim_score_{chunk_signature}",
                )
                sim_tags = st.multiselect("Simultaneous issue tags", SIMULTANEOUS_ERROR_TAGS, key=f"sim_tags_{chunk_signature}")
                sim_notes = st.text_area("Notes", key=f"sim_notes_{chunk_signature}", height=90)

                nav_cols = st.columns([1, 2, 1])
                with nav_cols[0]:
                    if st.button("← Previous chunk", disabled=sim_idx == 0, key=f"sim_prev_{chunk_signature}"):
                        st.session_state.simultaneous_idx -= 1
                        st.session_state.simultaneous_revealed = False
                        st.rerun()
                with nav_cols[1]:
                    if st.button("Save score and next chunk", type="primary", key=f"sim_save_{chunk_signature}"):
                        sim_row = {
                            "timestamp": datetime.now().isoformat(timespec="seconds"),
                            "session_id": st.session_state.get("simultaneous_session_id", ""),
                            "script_id": sim_script_id,
                            "title": sim_title,
                            "source_file": sim_name,
                            "practice_mode": sim_mode,
                            "training_mode": sim_training_mode,
                            "lag_seconds": sim_lag_seconds if sim_training_mode in {"Lag trainer", "Shadow first, then interpret"} else "",
                            "playback_speed": sim_playback_speed,
                            "repeat_loop": sim_repeat_loop,
                            "chunk_index": chunk.get("chunk_index", sim_idx + 1),
                            "chunk_label": chunk.get("chunk_label", ""),
                            "source_language": source_lang,
                            "target_language": target_lang,
                            "speaker": chunk.get("speaker", ""),
                            "segment_type": chunk.get("segment_type", ""),
                            "score": sim_score,
                            "error_tags": "; ".join(sim_tags),
                            "notes": sim_notes,
                            "source_text": source_text,
                            "target_text": target_text,
                        }
                        st.session_state.simultaneous_log.append(sim_row)
                        if auto_save_simultaneous_log:
                            try:
                                append_simultaneous_log(simultaneous_log_path, sim_row)
                            except Exception as exc:
                                st.warning(f"Could not auto-save simultaneous practice: {exc}")
                        if sim_idx < len(chunks) - 1:
                            st.session_state.simultaneous_idx += 1
                            st.session_state.simultaneous_revealed = False
                        st.rerun()
                with nav_cols[2]:
                    if st.button("Skip chunk →", disabled=sim_idx >= len(chunks) - 1, key=f"sim_skip_{chunk_signature}"):
                        st.session_state.simultaneous_idx += 1
                        st.session_state.simultaneous_revealed = False
                        st.rerun()

            if not st.session_state.simultaneous_revealed:
                nav1, nav2 = st.columns(2)
                with nav1:
                    if st.button("← Previous", disabled=sim_idx == 0, key=f"sim_prev_unrevealed_{chunk_signature}"):
                        st.session_state.simultaneous_idx -= 1
                        st.session_state.simultaneous_revealed = False
                        st.rerun()
                with nav2:
                    if st.button("Skip →", disabled=sim_idx >= len(chunks) - 1, key=f"sim_skip_unrevealed_{chunk_signature}"):
                        st.session_state.simultaneous_idx += 1
                        st.session_state.simultaneous_revealed = False
                        st.rerun()

            st.divider()
            st.markdown("### Simultaneous session log")
            sim_session_rows = st.session_state.get("simultaneous_log", [])
            if sim_session_rows:
                st.dataframe(sim_session_rows, use_container_width=True, hide_index=True)
                sim_scores = [int(r.get("score", 0)) for r in sim_session_rows if str(r.get("score", "")).isdigit()]
                if sim_scores:
                    st.write(f"Average simultaneous score: **{sum(sim_scores) / len(sim_scores):.2f} / 3** across **{len(sim_scores)} saved attempt(s)**.")
                if auto_save_simultaneous_log:
                    st.caption(f"Simultaneous attempts are being saved to: `{simultaneous_log_path}`")
                st.download_button(
                    "Download simultaneous session CSV",
                    data=simultaneous_log_to_csv(sim_session_rows).encode("utf-8"),
                    file_name=f"{sim_script_id}_simultaneous_session_log.csv",
                    mime="text/csv",
                )
                if st.button("Clear simultaneous session log", key="sim_clear_session"):
                    st.session_state.simultaneous_log = []
                    st.rerun()
            else:
                st.info("No simultaneous attempts saved in this app session yet.")

            with st.expander("Persistent simultaneous history", expanded=False):
                sim_history_rows = load_simultaneous_log(simultaneous_log_path)
                sim_history_df = simultaneous_rows_to_dataframe(sim_history_rows)
                st.write(f"Persistent simultaneous log file: `{simultaneous_log_path}`")
                if sim_history_df.empty:
                    st.info("No persistent simultaneous practice history yet.")
                else:
                    st.metric("Saved simultaneous attempts", len(sim_history_df))
                    valid_sim_scores = sim_history_df["score_numeric"].dropna()
                    if not valid_sim_scores.empty:
                        st.metric("Average saved simultaneous score", f"{valid_sim_scores.mean():.2f} / 3")
                    st.dataframe(sim_history_df.tail(int(show_history_preview_rows)), use_container_width=True, hide_index=True)
                    st.download_button(
                        "Download full simultaneous history CSV",
                        data=simultaneous_log_to_csv(sim_history_rows).encode("utf-8"),
                        file_name="simultaneous_practice_log.csv",
                        mime="text/csv",
                    )
    else:
        st.info("Choose a JSON file to start simultaneous practice.")



with tab_exam:
    st.subheader("Exam mode")
    st.caption("Source-only playback, microphone recording, transcription, transcript review, rubric score, and remedial practice JSON.")

    exam_mode = st.radio("Exam type", ["Consecutive", "Simultaneous"], horizontal=True, key="exam_mode_choice")
    exam_source = st.radio("Exam JSON source", ["Upload JSON file", "Use selected files from Generate tab", "Use local file path"], horizontal=True, key="exam_source_radio")

    exam_payload: dict[str, Any] | None = None
    exam_name = ""
    if exam_source == "Upload JSON file":
        exam_upload = st.file_uploader("Select one JSON script for exam mode", type=["json"], key="exam_file_upload")
        if exam_upload is not None:
            exam_name = exam_upload.name
            exam_payload, exam_error = load_json_bytes(exam_upload.name, exam_upload.getvalue())
            if exam_error:
                st.error(exam_error)
    elif exam_source == "Use selected files from Generate tab":
        if "payloads" in globals() and payloads:
            names = [name for name, _, _ in payloads]
            chosen = st.selectbox("Choose one loaded JSON", names, key="exam_loaded_json")
            for name, data, _raw in payloads:
                if name == chosen:
                    exam_name, exam_payload = name, data
                    break
        else:
            st.info("Load JSON files in the Generate outputs tab first, or use Upload JSON file here.")
    else:
        exam_path = Path(st.text_input("Exam JSON path", value="", key="exam_path")).expanduser()
        if exam_path.exists():
            exam_name = exam_path.name
            exam_payload, exam_error = load_json_bytes(exam_path.name, exam_path.read_bytes())
            if exam_error:
                st.error(exam_error)
        elif str(exam_path).strip():
            st.warning("That JSON path does not exist.")

    if exam_payload:
        script_id = safe_script_id(exam_payload, Path(exam_name).stem)
        title = str(exam_payload.get("title") or "")
        session_id = initialize_exam_session(script_id, exam_mode)

        exam_output_dir = Path(st.text_input("Exam output folder", value=str((DEFAULT_EXAM_OUTPUT_DIR / sanitize_filename_part(script_id)).resolve()), key="exam_output_dir")).expanduser()
        exam_cache_dir = exam_output_dir / "source_audio_cache"
        response_dir = exam_output_dir / "responses"
        combined_dir = exam_output_dir / "combined"
        exam_output_dir.mkdir(parents=True, exist_ok=True)

        left, right = st.columns([2, 1])
        with left:
            st.markdown(f"**{script_id}**")
            if title:
                st.caption(title)
        with right:
            st.caption(f"Session: `{session_id}`")

        use_profile_exam = st.checkbox("Use JSON audio_profile speaker voices", value=True, key="exam_use_profile")
        exam_speed_mode = st.selectbox("TTS voice speed", ["learning", "normal", "fast"], index=1, key="exam_tts_speed")
        exam_playback_speed = st.selectbox(
            "Playback speed",
            [0.85, 1.00, 1.10, 1.20, 1.30, 1.40, 1.50],
            index=1,
            format_func=lambda x: f"{x:.2f}x",
            key="exam_playback_speed",
        )
        show_exam_reference = st.checkbox("Allow reference text/audio after recording", value=True, key="exam_show_reference")

        with st.expander("Transcription settings", expanded=False):
            enable_exam_transcription = st.checkbox("Enable OpenAI transcription for saved exam recordings", value=False, key="exam_enable_transcription")
            transcript_model = st.selectbox("Basic transcription model", OPENAI_TRANSCRIBE_MODELS, index=0, key="exam_transcript_model")
            transcript_language_hint_label = st.selectbox("Language hint", ["Auto", "English", "Spanish"], index=0, key="exam_transcript_language_hint")
            manual_openai_key = st.text_input("OpenAI API key (optional; otherwise uses OPENAI_API_KEY)", type="password", key="exam_openai_api_key")
            st.caption("The key is used only for this Streamlit session. You can also set OPENAI_API_KEY as a Streamlit secret or environment variable.")
            st.markdown("**Timestamp alignment**")
            timestamp_lag_seconds = st.number_input("Estimated simultaneous lag in seconds", min_value=0.0, max_value=15.0, value=3.0, step=0.5, key="exam_timestamp_lag_seconds")
            timestamp_padding_seconds = st.number_input("Timestamp window padding in seconds", min_value=0.0, max_value=8.0, value=1.5, step=0.5, key="exam_timestamp_padding_seconds")
            timestamp_exact_timing = st.checkbox("Use exact source chunk timing map", value=True, key="exam_timestamp_exact_timing")
            st.caption("Timestamp-anchored alignment uses whisper-1 timestamps. The exact timing map generates/reuses per-chunk source audio, measures each chunk, scales to the full source MP3, adjusts for playback speed, then applies the lag setting.")

        if exam_mode == "Simultaneous":
            chunks = build_simultaneous_chunks(exam_payload, "Full passage", "Source language first")
            if not chunks:
                st.warning("No usable source/target paired segments found for simultaneous exam mode.")
            else:
                chunk = chunks[0]
                source_lang = chunk["source_language"]
                target_lang = chunk["target_language"]
                source_text = chunk["source_text"]
                target_text = chunk["target_text"]
                voice = get_practice_voice_for_chunk(exam_payload, chunk, source_lang, use_profile_exam)
                rate = get_rate_for_language(exam_speed_mode, source_lang)
                source_audio = get_exam_audio_path(exam_cache_dir, script_id, "full", "source", source_lang, source_text, voice, rate)
                exam_source_group_count = None
                if st.button("Prepare simultaneous source audio", key="exam_prepare_sim") or not source_audio.exists():
                    try:
                        merged_audio, source_groups = asyncio.run(generate_source_only_full_passage_audio(
                            data=exam_payload,
                            segments=chunk.get("segments") or [],
                            script_id=script_id,
                            cache_dir=exam_cache_dir / "source_only_full_passage",
                            role="exam_simultaneous",
                            speed_mode_value=exam_speed_mode,
                            use_audio_profile=use_profile_exam,
                            same_speaker_gap=float(source_only_gap),
                            speaker_change_gap=float(source_only_speaker_gap),
                            continuation_gap=float(source_only_continuation_gap),
                            max_merged_chars=int(source_only_max_merged_chars),
                            generator_path=DEFAULT_GENERATOR_PATH,
                            direction="Source language first",
                        ))
                        if merged_audio:
                            source_audio = merged_audio
                            exam_source_group_count = len(source_groups)
                            st.session_state.exam_sim_source_audio_path = str(source_audio)
                            st.session_state.exam_sim_source_group_count = exam_source_group_count
                        else:
                            asyncio.run(generate_practice_audio_file(source_text, voice, rate, source_audio, DEFAULT_GENERATOR_PATH))
                            st.session_state.exam_sim_source_audio_path = str(source_audio)
                            st.session_state.exam_sim_source_group_count = None
                    except Exception as exc:
                        st.error(f"Could not prepare source audio: {exc}")
                else:
                    cached_exam_source = st.session_state.get("exam_sim_source_audio_path")
                    if cached_exam_source and Path(cached_exam_source).exists():
                        source_audio = Path(cached_exam_source)
                        exam_source_group_count = st.session_state.get("exam_sim_source_group_count")

                st.markdown(compact_metadata_row([
                    "Simultaneous full passage", f"{source_lang[:2].upper()} → {target_lang[:2].upper()}",
                    f"{len(source_text.split())} source words", str(chunk.get("speaker") or ""), "Exam mode"
                ]), unsafe_allow_html=True)
                if source_audio.exists():
                    exam_source_label = f"Exam source — {source_lang.title()}"
                    if exam_source_group_count:
                        exam_source_label += f" — natural source-only full passage ({exam_source_group_count} voice block(s))"
                    else:
                        exam_source_label += f" — voice: {voice}"
                    render_audio_player(source_audio, exam_source_label, autoplay=False, playback_speed=float(exam_playback_speed))
                    render_wpm_estimate("Exam source", source_text, source_audio, float(exam_playback_speed))
                    if exam_source_group_count:
                        st.caption("Exam source audio uses source-only grouping: adjacent same-speaker chunks are merged; speaker changes keep a natural pause.")
                    st.download_button("Download source prompt MP3", data=source_audio.read_bytes(), file_name=source_audio.name, mime="audio/mpeg")
                else:
                    st.info("Prepare the source audio before starting the exam.")

                st.markdown("### Record your simultaneous interpretation")
                st.caption("Start the audio player, then use the recorder below to capture your interpretation. Stop recording when the source passage ends.")
                sim_audio_input = st.audio_input("Record simultaneous interpretation", key=f"exam_sim_audio_{session_id}")
                notes = st.text_area("Exam notes", key="exam_sim_notes")
                if st.button("Save simultaneous recording", key="save_sim_exam", disabled=sim_audio_input is None):
                    response_path = response_dir / f"{sanitize_filename_part(script_id)}_{session_id}_simultaneous_response.wav"
                    saved = save_audio_input(sim_audio_input, response_path)
                    if saved:
                        row = {
                            "timestamp": datetime.now().isoformat(timespec="seconds"), "session_id": session_id,
                            "script_id": script_id, "title": title, "source_file": exam_name,
                            "exam_mode": "Simultaneous", "item_index": "full", "item_label": "Full passage",
                            "source_language": source_lang, "target_language": target_lang,
                            "speaker": chunk.get("speaker", ""), "segment_type": chunk.get("segment_type", ""),
                            "source_text": source_text, "source_audio": str(source_audio), "response_audio": str(saved),
                            "reference_text": target_text, "notes": notes,
                            "playback_speed": float(exam_playback_speed), "tts_voice": voice, "tts_rate": rate,
                            **wpm_fields_for_text_audio(source_text, source_audio, float(exam_playback_speed)),
                            "alignment_segments_json": json.dumps(chunk.get("segments", []), ensure_ascii=False),
                        }
                        st.session_state.exam_rows = [row]
                        st.success(f"Saved recording: {saved.name}")

                if show_exam_reference:
                    with st.expander("Reveal reference interpretation", expanded=False):
                        render_source_reference_cards(
                            source_title=f"Source — {source_lang.title()}", source_text=source_text,
                            target_title=f"Reference — {target_lang.title()}", target_text=target_text,
                            source_meta=f"Voice: {voice}", target_meta="Reference text",
                            source_visible=True, target_visible=True,
                        )
                        target_voice = get_practice_voice_for_chunk(exam_payload, chunk, target_lang, use_profile_exam)
                        target_rate = get_rate_for_language(exam_speed_mode, target_lang)
                        ref_audio = get_exam_audio_path(exam_cache_dir, script_id, "full", "reference", target_lang, target_text, target_voice, target_rate)
                        try:
                            asyncio.run(generate_practice_audio_file(target_text, target_voice, target_rate, ref_audio, DEFAULT_GENERATOR_PATH))
                            render_audio_player(ref_audio, f"Reference — {target_lang.title()} — voice: {target_voice}", autoplay=False, playback_speed=float(exam_playback_speed))
                            render_wpm_estimate("Reference", target_text, ref_audio, float(exam_playback_speed))
                        except Exception as exc:
                            st.warning(f"Could not prepare reference audio: {exc}")

        else:
            paired = exam_payload.get("paired_segments", []) or []
            segments = []
            for idx, seg in enumerate(paired, start=1):
                if not isinstance(seg, dict):
                    continue
                source, target = get_segment_source_and_target(seg, "Source language first")
                source_text = str(seg.get(source, "")).strip()
                target_text = str(seg.get(target, "")).strip()
                if source_text:
                    segments.append((idx, seg, source, target, source_text, target_text))
            if not segments:
                st.warning("No usable source-language paired segments found for consecutive exam mode.")
            else:
                max_idx = len(segments) - 1
                if "exam_index" not in st.session_state:
                    st.session_state.exam_index = 0
                st.session_state.exam_index = max(0, min(int(st.session_state.exam_index), max_idx))
                pos = st.session_state.exam_index
                json_idx, seg, source_lang, target_lang, source_text, target_text = segments[pos]
                speaker = str(seg.get("speaker") or "")
                segment_type = str(seg.get("segment_type") or "")
                voice = get_practice_voice(exam_payload, seg, source_lang, use_profile_exam)
                rate = get_rate_for_language(exam_speed_mode, source_lang)
                source_audio = get_exam_audio_path(exam_cache_dir, script_id, json_idx, "source", source_lang, source_text, voice, rate)
                try:
                    asyncio.run(generate_practice_audio_file(source_text, voice, rate, source_audio, DEFAULT_GENERATOR_PATH))
                except Exception as exc:
                    st.error(f"Could not prepare source audio: {exc}")

                st.progress((pos + 1) / len(segments), text=f"Segment {pos + 1} of {len(segments)}")
                st.markdown(compact_metadata_row([
                    f"JSON segment {json_idx}", f"{source_lang[:2].upper()} → {target_lang[:2].upper()}",
                    f"{len(source_text.split())} words", speaker, segment_type
                ]), unsafe_allow_html=True)
                render_source_reference_cards(
                    source_title=f"Source — {source_lang.title()}", source_text=source_text,
                    target_title=f"Expected target — {target_lang.title()}", target_text=target_text,
                    source_meta=f"{speaker} · voice: {voice}", target_meta="Hidden during exam",
                    source_visible=True, target_visible=False,
                )
                if source_audio.exists():
                    render_audio_player(source_audio, f"Exam source — {source_lang.title()} — voice: {voice}", autoplay=False, playback_speed=float(exam_playback_speed))
                    render_wpm_estimate("Exam source", source_text, source_audio, float(exam_playback_speed))

                st.markdown("### Record your consecutive interpretation")
                st.caption("After listening to the source segment, record your interpretation. Save it before moving to the next segment.")
                response_audio = st.audio_input("Record interpretation for this segment", key=f"exam_consec_audio_{session_id}_{json_idx}")
                notes = st.text_area("Segment notes", key=f"exam_consec_notes_{session_id}_{json_idx}")

                col_a, col_b, col_c, col_d = st.columns([1.2, 1, 1, 1])
                with col_a:
                    if st.button("Save response", key=f"save_consec_{json_idx}", disabled=response_audio is None):
                        response_path = response_dir / f"{sanitize_filename_part(script_id)}_{session_id}_seg_{json_idx:04d}_response.wav"
                        saved = save_audio_input(response_audio, response_path)
                        if saved:
                            row = {
                                "timestamp": datetime.now().isoformat(timespec="seconds"), "session_id": session_id,
                                "script_id": script_id, "title": title, "source_file": exam_name,
                                "exam_mode": "Consecutive", "item_index": json_idx,
                                "item_label": f"Segment {pos + 1} / JSON segment {json_idx}",
                                "source_language": source_lang, "target_language": target_lang,
                                "speaker": speaker, "segment_type": segment_type,
                                "source_text": source_text, "source_audio": str(source_audio), "response_audio": str(saved),
                                "reference_text": target_text, "notes": notes,
                                "playback_speed": float(exam_playback_speed), "tts_voice": voice, "tts_rate": rate,
                                **wpm_fields_for_text_audio(source_text, source_audio, float(exam_playback_speed)),
                            }
                            st.session_state.exam_saved[str(json_idx)] = row
                            rows_by_index = {str(r.get("item_index")): r for r in st.session_state.get("exam_rows", [])}
                            rows_by_index[str(json_idx)] = row
                            st.session_state.exam_rows = [rows_by_index[k] for k in sorted(rows_by_index, key=lambda x: int(x) if str(x).isdigit() else 999999)]
                            st.success(f"Saved response for segment {json_idx}.")
                with col_b:
                    if st.button("Previous", disabled=pos <= 0, key="exam_prev"):
                        st.session_state.exam_index -= 1
                        st.rerun()
                with col_c:
                    if st.button("Next", disabled=pos >= max_idx, key="exam_next"):
                        st.session_state.exam_index += 1
                        st.rerun()
                with col_d:
                    if show_exam_reference and st.button("Reveal reference", key=f"exam_reveal_ref_{json_idx}"):
                        st.session_state[f"exam_ref_revealed_{json_idx}"] = True

                if show_exam_reference and st.session_state.get(f"exam_ref_revealed_{json_idx}", False):
                    render_source_reference_cards(
                        source_title=f"Source — {source_lang.title()}", source_text=source_text,
                        target_title=f"Reference — {target_lang.title()}", target_text=target_text,
                        source_meta=f"{speaker}", target_meta="Reference text",
                        source_visible=True, target_visible=True,
                    )

                saved_count = len(st.session_state.get("exam_saved", {}))
                st.caption(f"Saved responses: {saved_count} / {len(segments)}")

        st.divider()
        st.markdown("### Exam outputs")
        exam_rows = st.session_state.get("exam_rows", [])
        if exam_rows:
            st.dataframe(exam_rows, use_container_width=True, hide_index=True)
            csv_text = exam_log_to_csv(exam_rows)
            csv_path = exam_output_dir / f"{sanitize_filename_part(script_id)}_{session_id}_exam_log.csv"
            csv_path.write_text(csv_text, encoding="utf-8")
            st.download_button("Download exam log CSV", data=csv_text.encode("utf-8"), file_name=csv_path.name, mime="text/csv")

            if enable_exam_transcription:
                st.markdown("#### Transcription")
                st.caption("Transcribes saved response recordings only. With whisper-1, the app now saves plain text and timestamp JSON in one pass, then reuses that for structured review and timestamp alignment.")
                if st.button("Transcribe saved exam recordings", key="transcribe_exam_recordings"):
                    try:
                        api_key = get_openai_api_key(manual_openai_key)
                        language_hint = language_hint_to_api_value(transcript_language_hint_label)
                        transcript_dir = exam_output_dir / "transcripts"
                        transcript_rows, txt_path, transcript_csv_path = build_exam_transcript_outputs(
                            exam_rows=exam_rows, script_id=script_id, session_id=session_id,
                            transcript_dir=transcript_dir, model=transcript_model, api_key=api_key,
                            language_hint=language_hint,
                        )
                        st.session_state.exam_transcript_rows = transcript_rows
                        st.session_state.exam_last_transcript_txt = str(txt_path)
                        st.session_state.exam_last_transcript_csv = str(transcript_csv_path)
                        rows_by_index = {str(r.get("item_index")): r for r in exam_rows}
                        for tr in transcript_rows:
                            rows_by_index[str(tr.get("item_index"))] = tr
                        st.session_state.exam_rows = [rows_by_index[k] for k in sorted(rows_by_index, key=lambda x: int(x) if str(x).isdigit() else 999999)]
                        if transcript_model == "whisper-1":
                            # Build timestamp-aligned rows immediately from the cached Whisper JSON.
                            # This avoids a second transcription pass and makes the transcript usable
                            # by both Structured Review and Timestamp Alignment.
                            align_dir = exam_output_dir / "timestamp_alignment"
                            align_rows, align_txt_path, align_csv_path, align_docx_path = build_timestamp_anchored_alignment_outputs(
                                exam_rows=st.session_state.exam_rows,
                                script_id=script_id,
                                session_id=session_id,
                                out_dir=align_dir,
                                api_key=api_key,
                                language_hint=language_hint,
                                lag_seconds=float(timestamp_lag_seconds),
                                padding_seconds=float(timestamp_padding_seconds),
                                use_exact_timing_map=bool(timestamp_exact_timing),
                            )
                            st.session_state.exam_timestamp_alignment_rows = align_rows
                            st.session_state.exam_last_alignment_txt = str(align_txt_path)
                            st.session_state.exam_last_alignment_csv = str(align_csv_path)
                            st.session_state.exam_last_alignment_docx = str(align_docx_path)
                            st.success(f"Transcribed {len(transcript_rows)} recording(s) and built timestamp alignment for {len(align_rows)} chunk(s) from the same Whisper-1 pass.")
                        else:
                            st.success(f"Transcribed {len(transcript_rows)} recording(s). For unified timestamp review, use whisper-1.")
                    except Exception as exc:
                        st.error(f"Could not transcribe exam recordings: {exc}")

                transcript_rows = st.session_state.get("exam_transcript_rows", [])
                if transcript_rows:
                    st.dataframe([{k: r.get(k, "") for k in ["item_label", "source_language", "target_language", "transcript_text"]} for r in transcript_rows], use_container_width=True, hide_index=True)
                    txt_value = str(st.session_state.get("exam_last_transcript_txt", "") or "").strip()
                    csv_value = str(st.session_state.get("exam_last_transcript_csv", "") or "").strip()
                    txt_path = Path(txt_value) if txt_value else None
                    transcript_csv_path = Path(csv_value) if csv_value else None
                    if txt_path and txt_path.exists() and txt_path.is_file():
                        st.download_button("Download transcript TXT", data=txt_path.read_bytes(), file_name=txt_path.name, mime="text/plain")
                    if transcript_csv_path and transcript_csv_path.exists() and transcript_csv_path.is_file():
                        st.download_button("Download transcript CSV", data=transcript_csv_path.read_bytes(), file_name=transcript_csv_path.name, mime="text/csv")

            st.markdown("#### Structured exam report")
            st.caption("Builds Source | Reference interpretation | Your transcript. Full transcript view shows the complete transcript; timestamp/chunk rows remain available as optional detail.")
            current_report_rows = choose_structured_review_rows(exam_rows, st.session_state.get("exam_transcript_rows", []), st.session_state.get("exam_timestamp_alignment_rows", []))
            if current_report_rows:
                render_exam_report_review(current_report_rows)
            if st.button("Build structured exam report", key="build_structured_exam_report"):
                try:
                    report_dir = exam_output_dir / "reports"
                    report_rows, report_txt_path, report_csv_path, report_docx_path = build_structured_exam_report_outputs(
                        exam_rows=exam_rows,
                        transcript_rows=st.session_state.get("exam_transcript_rows", []),
                        script_id=script_id,
                        session_id=session_id,
                        report_dir=report_dir,
                        alignment_rows=st.session_state.get("exam_timestamp_alignment_rows", []),
                    )
                    st.session_state.exam_report_rows = report_rows
                    st.session_state.exam_last_report_txt = str(report_txt_path)
                    st.session_state.exam_last_report_csv = str(report_csv_path)
                    st.session_state.exam_last_report_docx = str(report_docx_path)
                    st.success("Structured exam report prepared.")
                except Exception as exc:
                    st.error(f"Could not build structured exam report: {exc}")

            report_txt_value = str(st.session_state.get("exam_last_report_txt", "") or "").strip()
            report_csv_value = str(st.session_state.get("exam_last_report_csv", "") or "").strip()
            report_docx_value = str(st.session_state.get("exam_last_report_docx", "") or "").strip()
            report_txt_path = Path(report_txt_value) if report_txt_value else None
            report_csv_path = Path(report_csv_value) if report_csv_value else None
            report_docx_path = Path(report_docx_value) if report_docx_value else None
            if report_docx_path and report_docx_path.exists() and report_docx_path.is_file():
                st.download_button("Download structured report DOCX", data=report_docx_path.read_bytes(), file_name=report_docx_path.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            if report_csv_path and report_csv_path.exists() and report_csv_path.is_file():
                st.download_button("Download structured report CSV", data=report_csv_path.read_bytes(), file_name=report_csv_path.name, mime="text/csv")
            if report_txt_path and report_txt_path.exists() and report_txt_path.is_file():
                st.download_button("Download structured report TXT", data=report_txt_path.read_bytes(), file_name=report_txt_path.name, mime="text/plain")

            st.markdown("#### Timestamp-anchored alignment")
            st.caption("Uses Whisper-1 word timestamps plus a source timing map. If you already transcribed with whisper-1, this reuses the cached timestamp JSON instead of transcribing again.")
            if st.button("Build timestamp-anchored alignment", key="build_timestamp_alignment", disabled=not enable_exam_transcription):
                try:
                    api_key = get_openai_api_key(manual_openai_key)
                    language_hint = language_hint_to_api_value(transcript_language_hint_label)
                    align_dir = exam_output_dir / "timestamp_alignment"
                    align_rows, align_txt_path, align_csv_path, align_docx_path = build_timestamp_anchored_alignment_outputs(
                        exam_rows=st.session_state.get("exam_rows", exam_rows),
                        script_id=script_id,
                        session_id=session_id,
                        out_dir=align_dir,
                        api_key=api_key,
                        language_hint=language_hint,
                        lag_seconds=float(timestamp_lag_seconds),
                        padding_seconds=float(timestamp_padding_seconds),
                        use_exact_timing_map=bool(timestamp_exact_timing),
                    )
                    st.session_state.exam_timestamp_alignment_rows = align_rows
                    st.session_state.exam_last_alignment_txt = str(align_txt_path)
                    st.session_state.exam_last_alignment_csv = str(align_csv_path)
                    st.session_state.exam_last_alignment_docx = str(align_docx_path)
                    st.success(f"Timestamp-anchored alignment prepared for {len(align_rows)} chunk(s).")
                except Exception as exc:
                    st.error(f"Could not build timestamp-anchored alignment: {exc}")
            if not enable_exam_transcription:
                st.caption("Enable OpenAI transcription above to use timestamp-anchored alignment.")
            alignment_rows = st.session_state.get("exam_timestamp_alignment_rows", [])
            if alignment_rows:
                render_timestamp_alignment_review(alignment_rows)
                align_txt_value = str(st.session_state.get("exam_last_alignment_txt", "") or "").strip()
                align_csv_value = str(st.session_state.get("exam_last_alignment_csv", "") or "").strip()
                align_docx_value = str(st.session_state.get("exam_last_alignment_docx", "") or "").strip()
                align_txt_path = Path(align_txt_value) if align_txt_value else None
                align_csv_path = Path(align_csv_value) if align_csv_value else None
                align_docx_path = Path(align_docx_value) if align_docx_value else None
                if align_docx_path and align_docx_path.exists() and align_docx_path.is_file():
                    st.download_button("Download timestamp alignment DOCX", data=align_docx_path.read_bytes(), file_name=align_docx_path.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                if align_csv_path and align_csv_path.exists() and align_csv_path.is_file():
                    st.download_button("Download timestamp alignment CSV", data=align_csv_path.read_bytes(), file_name=align_csv_path.name, mime="text/csv")
                if align_txt_path and align_txt_path.exists() and align_txt_path.is_file():
                    st.download_button("Download timestamp alignment TXT", data=align_txt_path.read_bytes(), file_name=align_txt_path.name, mime="text/plain")

            st.markdown("#### Rubric score + remedial practice JSON")
            st.caption(
                "Use the source transcript, reference translation, and your attempt transcript "
                "to create an estimated practice score plus a usable remedial script JSON for audio/drill work."
            )
            drill_source = st.radio(
                "Review source",
                ["Full structured transcript", "Timestamp-aligned transcript"],
                index=0,
                horizontal=True,
                key="exam_simplified_rubric_source",
            )
            if drill_source == "Timestamp-aligned transcript" and st.session_state.get("exam_timestamp_alignment_rows", []):
                drill_rows_for_builder = st.session_state.get("exam_timestamp_alignment_rows", [])
            else:
                drill_rows_for_builder = choose_structured_review_rows(
                    exam_rows,
                    st.session_state.get("exam_transcript_rows", []),
                    st.session_state.get("exam_timestamp_alignment_rows", []),
                )
            simple_cols = st.columns([1, 1, 1, 2])
            with simple_cols[0]:
                drill_target_wpm = st.number_input("Target WPM", min_value=80, max_value=220, value=140, step=5, key="drill_target_wpm")
            with simple_cols[1]:
                max_drill_count = st.number_input("Max remedial sections", min_value=5, max_value=40, value=18, step=1, key="max_personalized_drills")
            with simple_cols[2]:
                rubric_ai_model = st.selectbox("Rubric AI model", OPENAI_AI_FEEDBACK_MODELS, index=0 if OPENAI_AI_FEEDBACK_MODELS else 0, key="exam_rubric_ai_model")
            with simple_cols[3]:
                st.caption("This uses one stronger AI review call over source + reference + your attempt + WPM data + project terminology preferences. It does not use the original JSON terms_used list. The remedial JSON should include weak/high-value transcript sections in paired_segments.")

            if st.button("Generate AI rubric score + remedial JSON", key="build_personalized_drill_plan", disabled=not drill_rows_for_builder):
                try:
                    drill_dir = exam_output_dir / "rubric_remedial"
                    api_key = get_openai_api_key(manual_openai_key)
                    (
                        score_summary,
                        rubric_units,
                        drills,
                        score_txt_path,
                        score_csv_path,
                        score_docx_path,
                        rubric_txt_path,
                        rubric_csv_path,
                        remedial_json_path,
                    ) = save_ai_rubric_remedial_outputs(
                        payload=exam_payload,
                        rows=drill_rows_for_builder,
                        script_id=script_id,
                        session_id=session_id,
                        out_dir=drill_dir,
                        model=str(rubric_ai_model),
                        api_key=api_key,
                        target_wpm=int(drill_target_wpm),
                        passing_target=70,
                        max_segments=int(max_drill_count),
                    )
                    srs_terms = remedial_terms_from_drills(exam_payload, drills, limit=40)
                    drill_txt_path = drill_dir / f"{sanitize_filename_part(script_id)}_{session_id}_remedial_drills_preview.txt"
                    drill_csv_path = drill_dir / f"{sanitize_filename_part(script_id)}_{session_id}_remedial_drills_preview.csv"
                    drill_docx_path = drill_dir / f"{sanitize_filename_part(script_id)}_{session_id}_remedial_drills_preview.docx"
                    drill_json_path = drill_dir / f"{sanitize_filename_part(script_id)}_{session_id}_interactive_drills.json"
                    drill_txt_path.write_text(personalized_drills_to_txt(drills), encoding="utf-8")
                    drill_csv_path.write_text(personalized_drills_to_csv(drills), encoding="utf-8")
                    build_personalized_drills_docx(drills, drill_docx_path)
                    interactive_json = build_interactive_drill_json(exam_payload, drills, srs_terms, script_id, session_id, int(drill_target_wpm))
                    drill_json_path.write_text(json.dumps(interactive_json, ensure_ascii=False, indent=2), encoding="utf-8")
                    st.session_state.exam_personalized_drills = drills
                    st.session_state.exam_personalized_drill_srs_terms = srs_terms
                    st.session_state.exam_last_personalized_drills_txt = str(drill_txt_path)
                    st.session_state.exam_last_personalized_drills_csv = str(drill_csv_path)
                    st.session_state.exam_last_personalized_drills_docx = str(drill_docx_path)
                    st.session_state.exam_last_personalized_drills_json = str(drill_json_path)
                    st.session_state.latest_interactive_drill_json = interactive_json
                    st.session_state.exam_rubric_key_units = rubric_units
                    st.session_state.exam_last_rubric_key_units_txt = str(rubric_txt_path)
                    st.session_state.exam_last_rubric_key_units_csv = str(rubric_csv_path)
                    st.session_state.exam_last_remedial_drill_script_json = str(remedial_json_path)
                    st.session_state.exam_rubric_score_summary = score_summary
                    st.session_state.exam_last_rubric_score_txt = str(score_txt_path)
                    st.session_state.exam_last_rubric_score_csv = str(score_csv_path)
                    st.session_state.exam_last_rubric_score_docx = str(score_docx_path)
                    st.success(f"AI rubric score + remedial practice JSON prepared with {len(rubric_units)} key-unit finding(s) and {len(drills)} drill item(s).")
                except Exception as exc:
                    st.error(f"Could not generate AI rubric score + remedial JSON: {exc}")

            rubric_score_state = st.session_state.get("exam_rubric_score_summary", {})
            rubric_units_state = st.session_state.get("exam_rubric_key_units", [])
            personalized_drills_state = st.session_state.get("exam_personalized_drills", [])
            personalized_srs_terms_state = st.session_state.get("exam_personalized_drill_srs_terms", [])
            if rubric_score_state:
                render_rubric_score_summary(rubric_score_state)
            if rubric_units_state:
                with st.expander("Key units used to build the score and remedial JSON", expanded=False):
                    st.dataframe(rubric_units_state, use_container_width=True, hide_index=True)
            if personalized_drills_state:
                with st.expander("Remedial drill items preview", expanded=False):
                    render_personalized_drills(personalized_drills_state, personalized_srs_terms_state)
                if personalized_srs_terms_state:
                    if st.button("Add remedial terms to Term Review SRS", key="add_drill_terms_to_srs"):
                        try:
                            latest_srs_rows = merge_terms_into_srs_state(load_term_srs_state(term_srs_state_path), personalized_srs_terms_state)
                            save_term_srs_state(term_srs_state_path, latest_srs_rows)
                            st.success(f"Added/updated {len(personalized_srs_terms_state)} term candidate(s) in the Term Review SRS.")
                        except Exception as exc:
                            st.error(f"Could not add remedial terms to SRS: {exc}")

            drill_txt_value = str(st.session_state.get("exam_last_personalized_drills_txt", "") or "").strip()
            drill_csv_value = str(st.session_state.get("exam_last_personalized_drills_csv", "") or "").strip()
            drill_docx_value = str(st.session_state.get("exam_last_personalized_drills_docx", "") or "").strip()
            drill_json_value = str(st.session_state.get("exam_last_personalized_drills_json", "") or "").strip()
            rubric_txt_value = str(st.session_state.get("exam_last_rubric_key_units_txt", "") or "").strip()
            rubric_csv_value = str(st.session_state.get("exam_last_rubric_key_units_csv", "") or "").strip()
            rubric_score_txt_value = str(st.session_state.get("exam_last_rubric_score_txt", "") or "").strip()
            rubric_score_csv_value = str(st.session_state.get("exam_last_rubric_score_csv", "") or "").strip()
            rubric_score_docx_value = str(st.session_state.get("exam_last_rubric_score_docx", "") or "").strip()
            remedial_json_value = str(st.session_state.get("exam_last_remedial_drill_script_json", "") or "").strip()
            drill_txt_path = Path(drill_txt_value) if drill_txt_value else None
            drill_csv_path = Path(drill_csv_value) if drill_csv_value else None
            drill_docx_path = Path(drill_docx_value) if drill_docx_value else None
            drill_json_path = Path(drill_json_value) if drill_json_value else None
            rubric_txt_path = Path(rubric_txt_value) if rubric_txt_value else None
            rubric_csv_path = Path(rubric_csv_value) if rubric_csv_value else None
            rubric_score_txt_path = Path(rubric_score_txt_value) if rubric_score_txt_value else None
            rubric_score_csv_path = Path(rubric_score_csv_value) if rubric_score_csv_value else None
            rubric_score_docx_path = Path(rubric_score_docx_value) if rubric_score_docx_value else None
            remedial_json_path = Path(remedial_json_value) if remedial_json_value else None
            if remedial_json_path and remedial_json_path.exists() and remedial_json_path.is_file():
                st.download_button("Download remedial practice script JSON", data=remedial_json_path.read_bytes(), file_name=remedial_json_path.name, mime="application/json")
            if rubric_score_docx_path and rubric_score_docx_path.exists() and rubric_score_docx_path.is_file():
                st.download_button("Download rubric score summary DOCX", data=rubric_score_docx_path.read_bytes(), file_name=rubric_score_docx_path.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            if rubric_score_csv_path and rubric_score_csv_path.exists() and rubric_score_csv_path.is_file():
                st.download_button("Download rubric score summary CSV", data=rubric_score_csv_path.read_bytes(), file_name=rubric_score_csv_path.name, mime="text/csv")
            if rubric_score_txt_path and rubric_score_txt_path.exists() and rubric_score_txt_path.is_file():
                st.download_button("Download rubric score summary TXT", data=rubric_score_txt_path.read_bytes(), file_name=rubric_score_txt_path.name, mime="text/plain")
            if rubric_csv_path and rubric_csv_path.exists() and rubric_csv_path.is_file():
                st.download_button("Download key-unit CSV", data=rubric_csv_path.read_bytes(), file_name=rubric_csv_path.name, mime="text/csv")
            if rubric_txt_path and rubric_txt_path.exists() and rubric_txt_path.is_file():
                st.download_button("Download key-unit TXT", data=rubric_txt_path.read_bytes(), file_name=rubric_txt_path.name, mime="text/plain")
            if drill_json_path and drill_json_path.exists() and drill_json_path.is_file():
                st.download_button("Download Drill Studio JSON", data=drill_json_path.read_bytes(), file_name=drill_json_path.name, mime="application/json")
            if drill_docx_path and drill_docx_path.exists() and drill_docx_path.is_file():
                st.download_button("Download short remedial drill plan DOCX", data=drill_docx_path.read_bytes(), file_name=drill_docx_path.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            if drill_csv_path and drill_csv_path.exists() and drill_csv_path.is_file():
                st.download_button("Download short remedial drill plan CSV", data=drill_csv_path.read_bytes(), file_name=drill_csv_path.name, mime="text/csv")
            if drill_txt_path and drill_txt_path.exists() and drill_txt_path.is_file():
                st.download_button("Download short remedial drill plan TXT", data=drill_txt_path.read_bytes(), file_name=drill_txt_path.name, mime="text/plain")

            if st.button("Build combined exam audio files", key="build_exam_combined"):
                try:
                    response_mp3s = []
                    interleaved_parts = []
                    for row in exam_rows:
                        src_mp3 = Path(str(row.get("source_audio", "")))
                        response_path = Path(str(row.get("response_audio", "")))
                        if response_path.exists():
                            response_mp3 = combined_dir / (response_path.stem + ".mp3")
                            ffmpeg_convert_to_mp3(response_path, response_mp3)
                            response_mp3s.append(response_mp3)
                            if src_mp3.exists():
                                interleaved_parts.append(src_mp3)
                            interleaved_parts.append(response_mp3)
                    made = []
                    resp_combined = ffmpeg_concat_mp3s(response_mp3s, combined_dir / f"{sanitize_filename_part(script_id)}_{session_id}_responses_only.mp3")
                    if resp_combined:
                        made.append(resp_combined)
                    interleaved = ffmpeg_concat_mp3s(interleaved_parts, combined_dir / f"{sanitize_filename_part(script_id)}_{session_id}_interleaved_exam.mp3")
                    if interleaved:
                        made.append(interleaved)
                    transcript_paths = []
                    for key in ["exam_last_transcript_txt", "exam_last_transcript_csv", "exam_last_report_txt", "exam_last_report_csv", "exam_last_report_docx", "exam_last_alignment_txt", "exam_last_alignment_csv", "exam_last_alignment_docx", "exam_last_personalized_drills_txt", "exam_last_personalized_drills_csv", "exam_last_personalized_drills_docx", "exam_last_personalized_drills_json", "exam_last_rubric_key_units_txt", "exam_last_rubric_key_units_csv", "exam_last_rubric_score_txt", "exam_last_rubric_score_csv", "exam_last_rubric_score_docx", "exam_last_remedial_drill_script_json"]:
                        val = str(st.session_state.get(key, "") or "").strip()
                        if val:
                            transcript_paths.append(Path(val))
                    for r in st.session_state.get("exam_transcript_rows", []):
                        val = str(r.get("transcript_file", "") or "").strip()
                        if val:
                            transcript_paths.append(Path(val))
                    for r in st.session_state.get("exam_timestamp_alignment_rows", []):
                        for field in ["timestamp_json_file", "timestamp_full_transcript_file", "timing_map_file"]:
                            val = str(r.get(field, "") or "").strip()
                            if val:
                                transcript_paths.append(Path(val))
                    all_paths = made + [csv_path] + transcript_paths + [Path(str(r.get("response_audio", ""))) for r in exam_rows]
                    zip_path = make_zip_from_paths(all_paths, exam_output_dir / f"{sanitize_filename_part(script_id)}_{session_id}_exam_outputs.zip")
                    st.session_state.exam_last_zip = str(zip_path) if zip_path else ""
                    st.success("Combined exam outputs prepared.")
                except Exception as exc:
                    st.error(f"Could not build combined exam files: {exc}")

            zip_value = str(st.session_state.get("exam_last_zip", "") or "").strip()
            zip_path = Path(zip_value) if zip_value else None
            if zip_path and zip_path.exists() and zip_path.is_file():
                st.download_button("Download exam outputs ZIP", data=zip_path.read_bytes(), file_name=zip_path.name, mime="application/zip")
        else:
            st.info("No exam recordings saved yet.")


with tab_drills:
    st.subheader("Drill Studio")
    st.caption("Practice remedial JSON files with term flashcards from `terms_used` and segment drills from `paired_segments`.")

    if "drill_studio_session_id" not in st.session_state:
        st.session_state.drill_studio_session_id = f"drill-session-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{uuid.uuid4().hex[:8]}"
    if "drill_studio_session_rows" not in st.session_state:
        st.session_state.drill_studio_session_rows = []

    source_choice = st.radio("Drill set source", ["Use latest generated remedial JSON", "Upload drill JSON"], horizontal=True, key="drill_studio_source_choice")
    drill_json_data: dict[str, Any] | None = None
    if source_choice == "Use latest generated remedial JSON":
        latest_candidates = [
            st.session_state.get("exam_last_remedial_drill_script_json", ""),
            st.session_state.get("exam_last_personalized_drills_json", ""),
        ]
        if isinstance(st.session_state.get("latest_interactive_drill_json"), dict):
            drill_json_data = st.session_state.latest_interactive_drill_json
            st.success("Loaded the latest generated drill set from Exam Mode.")
        else:
            for latest_path_value in latest_candidates:
                latest_path_value = str(latest_path_value or "").strip()
                if latest_path_value and Path(latest_path_value).exists():
                    try:
                        drill_json_data = load_drill_json_from_bytes(Path(latest_path_value).read_bytes())
                        st.session_state.latest_interactive_drill_json = drill_json_data
                        st.success("Loaded the latest generated remedial JSON from Exam Mode.")
                        break
                    except Exception as exc:
                        st.warning(f"Could not load latest drill JSON: {exc}")
            if drill_json_data is None:
                st.info("No remedial JSON has been generated yet. Build one in Exam Mode or upload one here.")
    else:
        drill_upload = st.file_uploader("Upload remedial or interactive drill JSON", type=["json"], key="drill_json_upload")
        if drill_upload is not None:
            try:
                drill_json_data = load_drill_json_from_bytes(drill_upload.getvalue())
                st.session_state.latest_uploaded_drill_json = drill_json_data
            except Exception as exc:
                st.error(f"Could not load drill JSON: {exc}")
        elif isinstance(st.session_state.get("latest_uploaded_drill_json"), dict):
            drill_json_data = st.session_state.latest_uploaded_drill_json

    if not drill_json_data:
        st.stop()

    drill_set_id = str(drill_json_data.get("drill_set_id", "") or drill_json_data.get("script_id", "drill_set") or "drill_set")
    term_drills, segment_drills, legacy_drills = get_drill_studio_collections(drill_json_data)
    total_items = len(term_drills) + len(segment_drills) + len(legacy_drills)
    if total_items == 0:
        st.info("This JSON does not contain terms, paired segments, or drill cards that Drill Studio can practice.")
        st.stop()

    top_cols = st.columns(4)
    top_cols[0].metric("Term flashcards", len(term_drills))
    top_cols[1].metric("Segment drills", len(segment_drills))
    top_cols[2].metric("Other drills", len(legacy_drills))
    top_cols[3].metric("Session rows", len(st.session_state.drill_studio_session_rows))

    if drill_json_data.get("title"):
        st.markdown(f"**{drill_json_data.get('title')}**")
    if drill_json_data.get("performance_targets"):
        with st.expander("Performance targets", expanded=False):
            for target in ensure_list(drill_json_data.get("performance_targets")):
                if str(target).strip():
                    st.write(f"- {target}")

    with st.expander("Drill Studio settings", expanded=False):
        a1, a2, a3, a4 = st.columns(4)
        with a1:
            drill_audio_enabled = st.checkbox("Enable drill audio", value=True, key="drill_audio_enabled")
        with a2:
            drill_autoplay_prompt = st.checkbox("Auto-play prompt", value=True, key="drill_autoplay_prompt")
        with a3:
            drill_play_answer = st.checkbox("Play answer on reveal", value=True, key="drill_play_answer")
        with a4:
            drill_audio_only = st.checkbox("Audio-only prompt", value=False, key="drill_audio_only")
        drill_audio_speed = st.selectbox("Playback speed", [0.85, 1.0, 1.10, 1.20, 1.30, 1.40, 1.50], index=1, format_func=lambda x: f"{x:.2f}x", key="drill_audio_speed")
        drill_audio_cache_dir = Path(st.text_input("Drill audio cache folder", value=str((RUNTIME_CACHE_DIR / "drill_audio_cache").resolve()), key="drill_audio_cache_dir")).expanduser()

    practice_tabs = []
    labels = []
    if term_drills:
        labels.append("Term Flashcards")
    if segment_drills:
        labels.append("Segment Drills")
    if legacy_drills:
        labels.append("Other Drill Cards")
    if not labels:
        st.stop()
    practice_tabs = st.tabs(labels)
    tab_i = 0
    if term_drills:
        with practice_tabs[tab_i]:
            st.markdown("### Term Flashcards")
            st.caption("Built from the remedial JSON `terms_used` section. Use these to make weak/high-value legal terms automatic.")
            render_drill_practice_panel(term_drills, drill_json_data, drill_set_id, "term_flashcards", "Term flashcard", drill_audio_enabled, drill_autoplay_prompt, drill_play_answer, drill_audio_only, float(drill_audio_speed), drill_audio_cache_dir, speed_mode, generator_path, drill_studio_log_path)
        tab_i += 1
    if segment_drills:
        with practice_tabs[tab_i]:
            st.markdown("### Segment Drills")
            st.caption("Built from the remedial JSON `paired_segments` section. Interpret each weak/high-value source segment aloud, then compare with the reference.")
            render_drill_practice_panel(segment_drills, drill_json_data, drill_set_id, "segment_drills", "Segment drill", drill_audio_enabled, drill_autoplay_prompt, drill_play_answer, drill_audio_only, float(drill_audio_speed), drill_audio_cache_dir, speed_mode, generator_path, drill_studio_log_path)
        tab_i += 1
    if legacy_drills:
        with practice_tabs[tab_i]:
            st.markdown("### Other Drill Cards")
            st.caption("Compatibility view for older interactive drill JSON files that contain a top-level `drills` list.")
            render_drill_practice_panel(legacy_drills, drill_json_data, drill_set_id, "other_drills", "Drill", drill_audio_enabled, drill_autoplay_prompt, drill_play_answer, drill_audio_only, float(drill_audio_speed), drill_audio_cache_dir, speed_mode, generator_path, drill_studio_log_path)

    history_rows = load_drill_studio_log(drill_studio_log_path)
    with st.expander("Drill Studio history", expanded=False):
        st.write(f"History file: `{drill_studio_log_path}`")
        if history_rows:
            st.dataframe(pd.DataFrame(history_rows).tail(100), use_container_width=True, hide_index=True)
            st.download_button("Download full Drill Studio history CSV", data=drill_studio_log_to_csv(history_rows).encode("utf-8"), file_name="drill_studio_log.csv", mime="text/csv")
        else:
            st.info("No Drill Studio history yet.")


with tab_history:
    st.subheader("Study history")
    st.caption("Review saved practice scores and long-term practice patterns.")

    history_rows = load_persistent_practice_log(persistent_log_path)
    history_df = rows_to_dataframe(history_rows)

    top_cols = st.columns([2, 1, 1])
    with top_cols[0]:
        st.write(f"Persistent log file: `{persistent_log_path}`")
    with top_cols[1]:
        if st.button("Reload history"):
            st.rerun()
    with top_cols[2]:
        if history_rows:
            st.download_button(
                "Download full history CSV",
                data=practice_log_to_csv(history_rows).encode("utf-8"),
                file_name="practice_log.csv",
                mime="text/csv",
            )

    if history_df.empty:
        st.info("No persistent study history yet. Score a few segments in the Consecutive practice tab with auto-save enabled.")
    else:
        scripts = sorted([s for s in history_df["script_id"].dropna().unique().tolist() if str(s).strip()])
        directions = sorted([d for d in history_df["direction"].dropna().unique().tolist() if "nan" not in str(d).lower()])
        speakers = sorted([s for s in history_df["speaker"].dropna().unique().tolist() if str(s).strip()])

        f1, f2, f3 = st.columns(3)
        with f1:
            script_filter = st.multiselect("Filter by script", scripts, default=[])
        with f2:
            direction_filter = st.multiselect("Filter by direction", directions, default=[])
        with f3:
            speaker_filter = st.multiselect("Filter by speaker", speakers, default=[])

        filtered = history_df.copy()
        if script_filter:
            filtered = filtered[filtered["script_id"].isin(script_filter)]
        if direction_filter:
            filtered = filtered[filtered["direction"].isin(direction_filter)]
        if speaker_filter:
            filtered = filtered[filtered["speaker"].isin(speaker_filter)]

        valid_scores = filtered["score_numeric"].dropna()
        metric_cols = st.columns(5)
        metric_cols[0].metric("Segments practiced", int(len(filtered)))
        metric_cols[1].metric("Average score", f"{valid_scores.mean():.2f} / 3" if len(valid_scores) else "—")
        metric_cols[2].metric("Weak segments", int((filtered["score_numeric"] <= 2).sum()))
        metric_cols[3].metric("Scripts", int(filtered["script_id"].nunique()))
        metric_cols[4].metric("Sessions", int(filtered["session_id"].nunique()))

        st.markdown("#### Progress by script")
        by_script = (
            filtered.dropna(subset=["score_numeric"])
            .groupby("script_id", as_index=False)
            .agg(segments=("score_numeric", "count"), average_score=("score_numeric", "mean"))
            .sort_values(["average_score", "segments"], ascending=[True, False])
        )
        if not by_script.empty:
            st.dataframe(by_script, use_container_width=True, hide_index=True)

        st.markdown("#### Average score by direction")
        by_direction = (
            filtered.dropna(subset=["score_numeric"])
            .groupby("direction", as_index=False)
            .agg(segments=("score_numeric", "count"), average_score=("score_numeric", "mean"))
            .sort_values("average_score")
        )
        if not by_direction.empty:
            st.dataframe(by_direction, use_container_width=True, hide_index=True)

        st.markdown("#### Common error tags")
        tag_counts = split_error_tags(filtered.to_dict("records"))
        if tag_counts:
            tag_df = pd.DataFrame([{"error_tag": k, "count": v} for k, v in tag_counts.items()])
            st.dataframe(tag_df, use_container_width=True, hide_index=True)
        else:
            st.info("No error tags recorded in the current filter.")

        st.markdown("#### Weak / missed segments for review")
        weak = filtered[filtered["score_numeric"] <= 2].sort_values("timestamp_dt", ascending=False)
        if weak.empty:
            st.success("No weak segments in the current filter.")
        else:
            show_cols = [
                "timestamp", "script_id", "segment_index", "speaker", "segment_type",
                "direction", "score", "error_tags", "notes", "source_text", "target_text",
            ]
            st.dataframe(weak[show_cols].head(int(show_history_preview_rows)), use_container_width=True, hide_index=True)
            weak_csv = weak[PRACTICE_LOG_FIELDS].to_csv(index=False)
            st.download_button(
                "Download weak segments CSV",
                data=weak_csv.encode("utf-8"),
                file_name="weak_segments_for_review.csv",
                mime="text/csv",
            )

        with st.expander("Raw history table"):
            st.dataframe(filtered[PRACTICE_LOG_FIELDS].sort_values("timestamp", ascending=False).head(int(show_history_preview_rows)), use_container_width=True, hide_index=True)

        with st.expander("History maintenance"):
            st.warning("This clears the runtime CSV for the current cloud session. Download your history first if you want to keep it.")
            if st.button("Back up history CSV now"):
                backup = persistent_log_path.with_name(f"practice_log_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv")
                try:
                    persistent_log_path.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copy2(persistent_log_path, backup)
                    st.success(f"Backup created: {backup}")
                except Exception as exc:
                    st.error(f"Could not create backup: {exc}")
            confirm_delete = st.checkbox("I understand this will delete all persistent history")
            if st.button("Delete persistent history CSV", disabled=not confirm_delete):
                try:
                    if persistent_log_path.exists():
                        persistent_log_path.unlink()
                    st.success("Persistent history deleted.")
                    st.rerun()
                except Exception as exc:
                    st.error(f"Could not delete history: {exc}")



with tab_terms:
    st.subheader("Term review")
    st.caption("Practice legal terms with spaced review, spoken prompts, answer audio, and persistent scheduling.")

    st.markdown("### 1. Choose source JSON files")
    term_source_mode = st.radio("Term source", ["Upload JSON file(s)", "Use a local folder path"], horizontal=True, key="term_source_mode")

    term_loaded_files: list[tuple[str, bytes]] = []
    if term_source_mode == "Upload JSON file(s)":
        term_uploads = st.file_uploader("Select JSON scripts for term review", type=["json"], accept_multiple_files=True, key="term_json_uploads")
        for uploaded in term_uploads or []:
            term_loaded_files.append((uploaded.name, uploaded.getvalue()))
    else:
        term_local_dir = Path(st.text_input("Local JSON folder for term review", value=str((APP_DIR / "scripts").resolve()), key="term_local_dir")).expanduser()
        term_pattern = st.text_input("Term file pattern", value="*.json", key="term_file_pattern")
        if term_local_dir.exists():
            term_paths = sorted(term_local_dir.glob(term_pattern))
            selected_term_names = st.multiselect("Term files found", [p.name for p in term_paths], default=[p.name for p in term_paths[:10]], key="selected_term_names")
            for pth in term_paths:
                if pth.name in selected_term_names:
                    term_loaded_files.append((pth.name, pth.read_bytes()))
        else:
            st.warning("That local folder does not exist yet.")

    term_payloads, term_parse_errors = load_term_payloads_from_files(term_loaded_files)
    for err in term_parse_errors:
        st.error(err)

    all_terms: list[dict[str, Any]] = []
    for filename, data, _raw in term_payloads:
        all_terms.extend(extract_terms_from_payload(data, filename))

    if not all_terms:
        st.info("Choose JSON files with a terms_used list to start term review.")
    else:
        term_history_rows = load_term_review_log(term_review_log_path)
        term_history_df = term_rows_to_dataframe(term_history_rows)
        term_srs_rows = merge_terms_into_srs_state(load_term_srs_state(term_srs_state_path), all_terms)
        try:
            save_term_srs_state(term_srs_state_path, term_srs_rows)
        except Exception as exc:
            st.warning(f"Could not update term SRS state file: {exc}")
        term_srs_df = term_srs_to_dataframe(term_srs_rows)

        setup_cols = st.columns([1.4, 1, 1, 1])
        with setup_cols[0]:
            script_options = sorted(set(t["script_id"] for t in all_terms))
            selected_scripts_for_terms = st.multiselect("Scripts", script_options, default=script_options, key="term_selected_scripts")
        with setup_cols[1]:
            drill_direction = st.selectbox("Drill direction", ["English → Spanish", "Spanish → English"], key="term_drill_direction")
        with setup_cols[2]:
            review_mode = st.selectbox("Review order", ["Due today", "Review terms", "New terms", "All SRS priority", "Hard/missed first", "New terms first", "Script order", "Random"], key="term_review_order")
        with setup_cols[3]:
            limit_terms = st.number_input("Session size", min_value=5, max_value=500, value=min(30, max(5, len(all_terms))), step=5, key="term_session_size")

        terms = [t for t in all_terms if t["script_id"] in selected_scripts_for_terms]
        if review_mode in {"Due today", "Review terms", "New terms", "All SRS priority"}:
            terms = rank_terms_for_srs(terms, term_srs_rows, review_mode)[: int(limit_terms)]
        else:
            terms = rank_terms_for_review(terms, term_history_rows, review_mode)[: int(limit_terms)]

        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Terms in deck", len(terms))
        if not term_srs_df.empty:
            due_count = int((term_srs_df["next_review_date"].isna() | (term_srs_df["next_review_date"] <= today_date())).sum())
            new_count = int((term_srs_df["attempts"] == 0).sum())
        else:
            due_count = 0
            new_count = 0
        m2.metric("Due today", due_count)
        m3.metric("New terms", new_count)
        if not term_srs_df.empty:
            mastered_count = int((term_srs_df["mastery_level"] >= 4).sum())
            m4.metric("Mastered", mastered_count)
        elif not term_history_df.empty and term_history_df["score_numeric"].notna().any():
            m4.metric("Average term score", f"{term_history_df['score_numeric'].mean():.2f} / 3")
        else:
            m4.metric("Mastered", "—")

        ctrl1, ctrl2, ctrl3 = st.columns([1, 1, 2])
        with ctrl1:
            if st.button("Start new term session", type="primary"):
                reset_term_review_state("manual_restart")
                st.rerun()
        with ctrl2:
            if st.button("Clear term session log"):
                st.session_state.term_review_log = []
                st.rerun()
        with ctrl3:
            st.caption(f"Persistent term log: `{term_review_log_path}`  ·  SRS state: `{term_srs_state_path}`")

        if "term_review_log" not in st.session_state:
            st.session_state.term_review_log = []
        if "term_review_session_id" not in st.session_state:
            reset_term_review_state("initial")
        if "term_review_idx" not in st.session_state:
            st.session_state.term_review_idx = 0
        if "term_review_revealed" not in st.session_state:
            st.session_state.term_review_revealed = False

        if not terms:
            st.warning("No terms match the current filters.")
        else:
            st.session_state.term_review_idx = max(0, min(int(st.session_state.term_review_idx), len(terms) - 1))
            term_idx = int(st.session_state.term_review_idx)
            current_term = terms[term_idx]
            progress = (term_idx + 1) / max(len(terms), 1)
            st.progress(progress, text=f"Term {term_idx + 1} of {len(terms)}")

            prompt_lang = "english" if drill_direction == "English → Spanish" else "spanish"
            answer_lang = "spanish" if prompt_lang == "english" else "english"
            prompt_text = current_term[prompt_lang]
            answer_text = current_term[answer_lang]
            term_audio_rate = get_rate_for_language(speed_mode, prompt_lang)
            prompt_voice = DEFAULT_PRACTICE_VOICES.get(prompt_lang, "en-US-JennyNeural")
            answer_voice = DEFAULT_PRACTICE_VOICES.get(answer_lang, "en-US-JennyNeural")
            term_audio_key_base = short_hash(
                current_term.get("term_id", ""),
                current_term.get("script_id", ""),
                drill_direction,
                str(term_idx),
                prompt_text,
                answer_text,
            )
            show_prompt_text_key = f"term_show_prompt_text_{term_audio_key_base}"

            srs_lookup = {str(r.get("term_key", "")): normalize_srs_row(r) for r in term_srs_rows}
            current_srs_key = term_srs_key(current_term)
            current_srs = srs_lookup.get(current_srs_key, {})

            st.markdown("### Current term")

            if term_audio_enabled:
                prompt_audio_path = get_practice_audio_path(
                    term_audio_cache_dir,
                    current_term.get("script_id", "terms"),
                    term_audio_key_base,
                    "prompt",
                    prompt_lang,
                    prompt_text,
                    prompt_voice,
                    term_audio_rate,
                )
                try:
                    asyncio.run(generate_practice_audio_file(prompt_text, prompt_voice, term_audio_rate, prompt_audio_path, generator_path))
                    render_audio_player(
                        prompt_audio_path,
                        label=f"Prompt term audio — {prompt_lang.title()} — voice: {prompt_voice}",
                        autoplay=term_audio_autoplay_prompt,
                        show_native_player=True,
                        nonce=f"prompt-{term_audio_key_base}-{st.session_state.get('term_review_replay_nonce', '')}",
                    )
                except Exception as exc:
                    st.warning(f"Could not generate/play prompt term audio: {exc}")

                if term_audio_show_replay:
                    if st.button("↻ Replay prompt term", key=f"replay_prompt_term_{term_audio_key_base}"):
                        st.session_state.term_review_replay_nonce = uuid.uuid4().hex[:8]
                        st.rerun()

            prompt_display_text = prompt_text
            if term_audio_only_prompt and not st.session_state.get(show_prompt_text_key, False):
                prompt_display_text = "🎧 Audio-only prompt mode: listen first, retrieve aloud, then reveal or show text."
                if st.button("Show prompt text", key=f"show_prompt_text_{term_audio_key_base}"):
                    st.session_state[show_prompt_text_key] = True
                    st.rerun()

            st.markdown(
                f"""
                <div style="border:1px solid #ddd; border-radius:18px; padding:22px; background:#fafafa; margin-bottom:12px;">
                  <div style="font-size:0.85rem; color:#666; margin-bottom:8px;">Prompt · {html.escape(prompt_lang.title())}</div>
                  <div style="font-size:1.65rem; font-weight:700; line-height:1.35;">{html.escape(str(prompt_display_text))}</div>
                  <div style="font-size:0.85rem; color:#666; margin-top:14px;">Script: {html.escape(str(current_term.get('script_id', '')))}</div>
                  <div style="font-size:0.85rem; color:#666; margin-top:6px;">SRS: attempts {html.escape(str(current_srs.get('attempts', '0')))} · next review {html.escape(str(current_srs.get('next_review_at', 'today')))} · interval {html.escape(str(current_srs.get('interval_days', '0')))} day(s) · mastery {html.escape(str(current_srs.get('mastery_level', '0')))}/5</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

            answer_box = st.empty()
            if st.session_state.term_review_revealed:
                answer_box.markdown(
                    f"""
                    <div style="border:1px solid #cfcfcf; border-radius:18px; padding:20px; background:#ffffff; margin-bottom:12px;">
                      <div style="font-size:0.85rem; color:#666; margin-bottom:8px;">Answer · {html.escape(answer_lang.title())}</div>
                      <div style="font-size:1.35rem; font-weight:650; line-height:1.35;">{html.escape(str(answer_text))}</div>
                      <div style="font-size:0.9rem; color:#555; margin-top:12px;">Alternatives: {html.escape(str(current_term.get('alternatives') or '—'))}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                if term_audio_enabled:
                    answer_audio_path = get_practice_audio_path(
                        term_audio_cache_dir,
                        current_term.get("script_id", "terms"),
                        term_audio_key_base,
                        "answer",
                        answer_lang,
                        answer_text,
                        answer_voice,
                        get_rate_for_language(speed_mode, answer_lang),
                    )
                    try:
                        asyncio.run(generate_practice_audio_file(answer_text, answer_voice, get_rate_for_language(speed_mode, answer_lang), answer_audio_path, generator_path))
                        render_audio_player(
                            answer_audio_path,
                            label=f"Answer term audio — {answer_lang.title()} — voice: {answer_voice}",
                            autoplay=term_audio_play_answer_on_reveal,
                            show_native_player=True,
                            nonce=f"answer-{term_audio_key_base}-{st.session_state.get('term_review_answer_replay_nonce', '')}",
                        )
                    except Exception as exc:
                        st.warning(f"Could not generate/play answer term audio: {exc}")

                    if term_audio_show_replay:
                        if st.button("↻ Replay answer term", key=f"replay_answer_term_{term_audio_key_base}"):
                            st.session_state.term_review_answer_replay_nonce = uuid.uuid4().hex[:8]
                            st.rerun()
            else:
                answer_box.info("Say the interpretation aloud, then reveal the preferred answer.")

            if not st.session_state.term_review_revealed:
                if st.button("👁 Reveal answer", type="primary"):
                    st.session_state.term_review_revealed = True
                    st.rerun()
            else:
                st.markdown("#### Score this term")
                score_label = st.radio(
                    "How did you do?",
                    list(TERM_SCORE_OPTIONS.keys()),
                    horizontal=True,
                    key=f"term_score_{current_term['term_id']}_{term_idx}",
                )
                term_notes = st.text_area("Notes", key=f"term_notes_{current_term['term_id']}_{term_idx}", height=80)

                nav_cols = st.columns([1, 2, 1])
                with nav_cols[0]:
                    if st.button("← Previous term", disabled=term_idx == 0):
                        st.session_state.term_review_idx -= 1
                        st.session_state.term_review_revealed = False
                        st.rerun()
                with nav_cols[1]:
                    if st.button("Save score and next term", type="primary"):
                        term_row = {
                            "timestamp": datetime.now().isoformat(timespec="seconds"),
                            "session_id": st.session_state.get("term_review_session_id", ""),
                            "script_id": current_term.get("script_id", ""),
                            "title": current_term.get("title", ""),
                            "source_file": current_term.get("source_file", ""),
                            "prompt_language": prompt_lang,
                            "answer_language": answer_lang,
                            "english": current_term.get("english", ""),
                            "spanish": current_term.get("spanish", ""),
                            "alternatives": current_term.get("alternatives", ""),
                            "score_label": score_label,
                            "score_numeric": TERM_SCORE_OPTIONS.get(score_label, ""),
                            "notes": term_notes,
                        }
                        st.session_state.term_review_log.append(term_row)
                        if auto_save_term_review_log:
                            try:
                                append_term_review_log(term_review_log_path, term_row)
                            except Exception as exc:
                                st.warning(f"Could not auto-save term review: {exc}")
                        try:
                            latest_srs_rows = merge_terms_into_srs_state(load_term_srs_state(term_srs_state_path), all_terms)
                            latest_by_key = {str(r.get("term_key", "")): normalize_srs_row(r) for r in latest_srs_rows}
                            key = term_srs_key(current_term)
                            base_srs = latest_by_key.get(key, {
                                "term_key": key,
                                "english": current_term.get("english", ""),
                                "spanish": current_term.get("spanish", ""),
                                "script_ids": current_term.get("script_id", ""),
                                "titles": current_term.get("title", ""),
                                "source_files": current_term.get("source_file", ""),
                                "created_at": datetime.now().isoformat(timespec="seconds"),
                            })
                            latest_by_key[key] = calculate_next_srs_state(base_srs, score_label)
                            save_term_srs_state(term_srs_state_path, list(latest_by_key.values()))
                        except Exception as exc:
                            st.warning(f"Could not update term SRS state: {exc}")
                        if term_idx < len(terms) - 1:
                            st.session_state.term_review_idx += 1
                            st.session_state.term_review_revealed = False
                        st.rerun()
                with nav_cols[2]:
                    if st.button("Skip term →", disabled=term_idx >= len(terms) - 1):
                        st.session_state.term_review_idx += 1
                        st.session_state.term_review_revealed = False
                        st.rerun()

            st.divider()
            st.markdown("### Term session log")
            session_terms = st.session_state.get("term_review_log", [])
            if session_terms:
                st.dataframe(session_terms, use_container_width=True, hide_index=True)
                scores = [int(r.get("score_numeric", 0)) for r in session_terms if str(r.get("score_numeric", "")).isdigit()]
                if scores:
                    st.write(f"Average term score: **{sum(scores) / len(scores):.2f} / 3** across **{len(scores)} saved term(s)**.")
                st.download_button(
                    "Download term session CSV",
                    data=term_review_to_csv(session_terms).encode("utf-8"),
                    file_name="term_review_session_log.csv",
                    mime="text/csv",
                )
            else:
                st.info("No terms scored in this app session yet.")

            st.markdown("### Weakest terms from saved history")
            updated_term_history_rows = load_term_review_log(term_review_log_path)
            stats = get_term_stats(updated_term_history_rows)
            if stats.empty:
                st.info("No persistent term-review history yet.")
            else:
                weak_terms = stats.sort_values(["average_score", "misses", "attempts"], ascending=[True, False, False]).head(30)
                st.dataframe(weak_terms, use_container_width=True, hide_index=True)
                st.download_button(
                    "Download full term review history CSV",
                    data=term_review_to_csv(updated_term_history_rows).encode("utf-8"),
                    file_name="term_review_log.csv",
                    mime="text/csv",
                )

            with st.expander("Term review maintenance"):
                st.warning("This clears the runtime term review CSV for the current cloud session. Download your history first if you want to keep it.")
                if st.button("Back up term review CSV now"):
                    backup = term_review_log_path.with_name(f"term_review_log_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv")
                    try:
                        term_review_log_path.parent.mkdir(parents=True, exist_ok=True)
                        shutil.copy2(term_review_log_path, backup)
                        st.success(f"Backup created: {backup}")
                    except Exception as exc:
                        st.error(f"Could not create backup: {exc}")
                confirm_term_delete = st.checkbox("I understand this will delete all persistent term review history")
                if st.button("Delete persistent term review CSV", disabled=not confirm_term_delete):
                    try:
                        if term_review_log_path.exists():
                            term_review_log_path.unlink()
                        st.success("Persistent term review history deleted.")
                        st.rerun()
                    except Exception as exc:
                        st.error(f"Could not delete term review history: {exc}")

st.divider()
st.caption("Interpreter training app with audio generation, exam practice, transcript review, rubric feedback, and personalized drills.")
