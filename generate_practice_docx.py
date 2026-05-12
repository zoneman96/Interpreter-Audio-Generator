import json
import os
import shutil
import subprocess
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from xml.sax.saxutils import escape

SCRIPTS_DIR = Path("scripts")
OUTPUT_DIR = Path("docx_output")
OUTPUT_DIR.mkdir(exist_ok=True)

TITLE_FONT_SIZE = 16
HEADING_FONT_SIZE = 13
BODY_FONT_SIZE = 11
SMALL_FONT_SIZE = 10

HIGHLIGHT_TERMS = True  # set to False to disable glossary term highlighting
GENERATE_MASTER_DOCX = True
GENERATE_MASTER_PDF = True
GENERATE_INDIVIDUAL_PDFS = True
MASTER_DOCX_NAME = "combined_practice_packet.docx"


def get_target_names() -> set[str] | None:
    raw = os.environ.get("PIPELINE_TARGETS")
    if not raw:
        return None
    try:
        items = json.loads(raw)
        if isinstance(items, list):
            return {str(x) for x in items}
    except json.JSONDecodeError:
        pass
    return None


def is_valid_script_json(data: dict[str, Any]) -> bool:
    if "script_id" not in data or "terms_used" not in data:
        return False
    has_paired = isinstance(data.get("paired_segments"), list) and len(data.get("paired_segments", [])) > 0
    has_full = isinstance(data.get("english_script"), str) and isinstance(data.get("spanish_script"), str)
    return has_paired or has_full


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(TITLE_FONT_SIZE)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(HEADING_FONT_SIZE)


def add_body_paragraph(doc: Document, text: str = ""):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.size = Pt(BODY_FONT_SIZE)
    return p


def add_terms_table(doc: Document, terms: list[dict[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "English Term"
    hdr[1].text = "Spanish Contextual Choice"
    hdr[2].text = "All Approved Spanish Equivalents"

    for term in terms:
        row = table.add_row().cells
        english = str(term.get("english", "")).strip()
        contextual = str(term.get("spanish_contextual_choice", "")).strip()
        alternatives = term.get("spanish_alternatives_from_sheet", []) or []
        alternatives_text = "; ".join(str(x).strip() for x in alternatives if str(x).strip())

        row[0].text = english
        row[1].text = contextual
        row[2].text = alternatives_text


def build_highlight_patterns(terms: list[dict[str, Any]], language: str) -> list[str]:
    patterns = []
    seen = set()

    for term in terms:
        if language == "english":
            text = str(term.get("english", "")).strip()
            if text:
                lowered = text.casefold()
                if lowered not in seen:
                    seen.add(lowered)
                    patterns.append(text)
        elif language == "spanish":
            values = []
            contextual = str(term.get("spanish_contextual_choice", "")).strip()
            if contextual:
                values.append(contextual)
            for alt in term.get("spanish_alternatives_from_sheet", []) or []:
                alt = str(alt).strip()
                if alt:
                    values.append(alt)
            for text in values:
                lowered = text.casefold()
                if lowered not in seen:
                    seen.add(lowered)
                    patterns.append(text)

    patterns.sort(key=len, reverse=True)
    return patterns


def add_highlighted_paragraph(doc: Document, text: str, patterns: list[str]) -> None:
    p = doc.add_paragraph()
    if not text:
        return

    if not HIGHLIGHT_TERMS or not patterns:
        run = p.add_run(text)
        run.font.size = Pt(BODY_FONT_SIZE)
        return

    lower_text = text.casefold()
    i = 0

    while i < len(text):
        matched = None
        match_len = 0

        for pat in patterns:
            pat_lower = pat.casefold()
            if lower_text.startswith(pat_lower, i):
                before_ok = i == 0 or not text[i - 1].isalnum()
                end = i + len(pat)
                after_ok = end >= len(text) or not text[end:end + 1].isalnum()
                if before_ok and after_ok:
                    matched = text[i:end]
                    match_len = len(pat)
                    break

        if matched:
            run = p.add_run(matched)
            run.font.size = Pt(BODY_FONT_SIZE)
            run.bold = True
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            i += match_len
        else:
            start = i
            i += 1
            while i < len(text):
                any_match = False
                for pat in patterns:
                    pat_lower = pat.casefold()
                    if lower_text.startswith(pat_lower, i):
                        before_ok = i == 0 or not text[i - 1].isalnum()
                        end = i + len(pat)
                        after_ok = end >= len(text) or not text[end:end + 1].isalnum()
                        if before_ok and after_ok:
                            any_match = True
                            break
                if any_match:
                    break
                i += 1
            run = p.add_run(text[start:i])
            run.font.size = Pt(BODY_FONT_SIZE)


def add_paired_segments_table(
    doc: Document,
    paired_segments: list[dict[str, Any]],
    english_patterns: list[str],
    spanish_patterns: list[str],
) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "English"
    hdr[1].text = "Español"

    for pair in paired_segments:
        en = str(pair.get("english", "")).strip()
        es = str(pair.get("spanish", "")).strip()
        row = table.add_row().cells

        p_en = row[0].paragraphs[0]
        p_es = row[1].paragraphs[0]

        write_highlighted_into_paragraph(p_en, en, english_patterns)
        write_highlighted_into_paragraph(p_es, es, spanish_patterns)


def write_highlighted_into_paragraph(paragraph, text: str, patterns: list[str]) -> None:
    if not text:
        return

    if not HIGHLIGHT_TERMS or not patterns:
        run = paragraph.add_run(text)
        run.font.size = Pt(BODY_FONT_SIZE)
        return

    lower_text = text.casefold()
    i = 0

    while i < len(text):
        matched = None
        match_len = 0

        for pat in patterns:
            pat_lower = pat.casefold()
            if lower_text.startswith(pat_lower, i):
                before_ok = i == 0 or not text[i - 1].isalnum()
                end = i + len(pat)
                after_ok = end >= len(text) or not text[end:end + 1].isalnum()
                if before_ok and after_ok:
                    matched = text[i:end]
                    match_len = len(pat)
                    break

        if matched:
            run = paragraph.add_run(matched)
            run.font.size = Pt(BODY_FONT_SIZE)
            run.bold = True
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            i += match_len
        else:
            start = i
            i += 1
            while i < len(text):
                any_match = False
                for pat in patterns:
                    pat_lower = pat.casefold()
                    if lower_text.startswith(pat_lower, i):
                        before_ok = i == 0 or not text[i - 1].isalnum()
                        end = i + len(pat)
                        after_ok = end >= len(text) or not text[end:end + 1].isalnum()
                        if before_ok and after_ok:
                            any_match = True
                            break
                if any_match:
                    break
                i += 1
            run = paragraph.add_run(text[start:i])
            run.font.size = Pt(BODY_FONT_SIZE)


def set_page_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def iter_text_blocks(text: str) -> list[str]:
    normalized = str(text)
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\\n", "\n")
    return [block.strip() for block in normalized.split("\n\n") if block.strip()]


def append_script_to_doc(doc: Document, data: dict[str, Any], add_page_break_after: bool = False) -> None:
    script_id = data["script_id"]
    title = data.get("title", script_id)
    category = data.get("category", "")
    subcategory = data.get("subcategory", "")
    terms = data.get("terms_used", [])
    paired_segments = data.get("paired_segments", [])
    english_script = data.get("english_script", "")
    spanish_script = data.get("spanish_script", "")

    english_patterns = build_highlight_patterns(terms, "english")
    spanish_patterns = build_highlight_patterns(terms, "spanish")

    add_title(doc, f"Interpreter Practice – {script_id}")
    meta = add_body_paragraph(doc)
    meta_run = meta.add_run(f"Title: {title}\nCategory: {category}\nSubcategory: {subcategory}")
    meta_run.font.size = Pt(BODY_FONT_SIZE)

    add_heading(doc, "Terms Covered")
    add_terms_table(doc, terms)

    if paired_segments:
        add_heading(doc, "Paired Practice Script")
        add_paired_segments_table(doc, paired_segments, english_patterns, spanish_patterns)
    else:
        add_heading(doc, "English Script")
        for block in iter_text_blocks(english_script):
            add_highlighted_paragraph(doc, block, english_patterns)

        add_heading(doc, "Spanish Script")
        for block in iter_text_blocks(spanish_script):
            add_highlighted_paragraph(doc, block, spanish_patterns)

    if english_script and spanish_script:
        doc.add_page_break()
        add_heading(doc, "Full English Script")
        for block in iter_text_blocks(english_script):
            add_highlighted_paragraph(doc, block, english_patterns)

        add_heading(doc, "Full Spanish Script")
        for block in iter_text_blocks(spanish_script):
            add_highlighted_paragraph(doc, block, spanish_patterns)

    if add_page_break_after:
        doc.add_page_break()


def add_master_title_page(doc: Document, payloads: list[dict[str, Any]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Combined Interpreter Practice Packet")
    r.bold = True
    r.font.size = Pt(18)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run("Generated from script JSON files")
    r2.italic = True
    r2.font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(18)
    r3 = p3.add_run(f"Included packets: {len(payloads)}")
    r3.font.size = Pt(SMALL_FONT_SIZE)

    add_heading(doc, "Contents")
    for data in payloads:
        add_body_paragraph(doc, f"{data.get('script_id', '')} — {data.get('title', '')}")

    doc.add_page_break()


def generate_docx(script_path: Path) -> Path | None:
    with script_path.open("r", encoding="utf-8") as f:
        data = json.load(f)

    if not is_valid_script_json(data):
        print(f"Skipping invalid or non-script JSON: {script_path.name}")
        return None

    doc = Document()
    set_page_margins(doc)
    append_script_to_doc(doc, data, add_page_break_after=False)

    out_path = OUTPUT_DIR / f"{data['script_id']}.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


def generate_master_docx(script_paths: list[Path]) -> Path | None:
    valid_payloads: list[dict[str, Any]] = []
    for script_path in script_paths:
        with script_path.open("r", encoding="utf-8") as f:
            data = json.load(f)
        if is_valid_script_json(data):
            valid_payloads.append(data)

    if not valid_payloads:
        print("No valid scripts found for combined packet.")
        return None

    doc = Document()
    set_page_margins(doc)
    add_master_title_page(doc, valid_payloads)

    for idx, data in enumerate(valid_payloads):
        append_script_to_doc(doc, data, add_page_break_after=(idx < len(valid_payloads) - 1))

    out_path = OUTPUT_DIR / MASTER_DOCX_NAME
    doc.save(out_path)
    print(f"Saved combined DOCX: {out_path}")
    return out_path




def iter_text_lines(text: str) -> list[str]:
    normalized = str(text)
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\\n", "\n")
    return [line.strip() for line in normalized.split("\n") if line.strip()]


def generate_single_script_pdf(script_path: Path, pdf_path: Path) -> Path | None:
    try:
        with script_path.open("r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        print(f"Pure-Python single-script PDF fallback failed to read {script_path.name}: {e}")
        return None

    if not is_valid_script_json(data):
        print(f"Pure-Python single-script PDF fallback skipped invalid script JSON: {script_path.name}")
        return None

    styles = getSampleStyleSheet()
    section_title = ParagraphStyle(
        "SectionTitleSingle",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        spaceBefore=10,
        spaceAfter=6,
    )
    script_title = ParagraphStyle(
        "ScriptTitleSingle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=19,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    meta_style = ParagraphStyle(
        "MetaSingle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=13.5,
        spaceAfter=10,
    )
    body_style = ParagraphStyle(
        "BodySingle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=13.5,
        spaceAfter=6,
    )
    table_cell = ParagraphStyle(
        "TableCellSingle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        spaceAfter=0,
    )
    table_head = ParagraphStyle(
        "TableHeadSingle",
        parent=table_cell,
        fontName="Helvetica-Bold",
        textColor=colors.white,
        alignment=TA_CENTER,
    )

    script_id = escape(str(data.get("script_id", "")))
    title = escape(str(data.get("title", "")))
    category = escape(str(data.get("category", "")))
    subcategory = escape(str(data.get("subcategory", "")))
    terms = data.get("terms_used", []) or []
    paired_segments = data.get("paired_segments", []) or []
    english_script = data.get("english_script", "")
    spanish_script = data.get("spanish_script", "")

    story = []
    story.append(Paragraph(f"Interpreter Practice - {script_id}", script_title))
    story.append(Paragraph(
        f"<b>Title:</b> {title}<br/><b>Category:</b> {category}<br/><b>Subcategory:</b> {subcategory}",
        meta_style,
    ))

    story.append(Paragraph("Terms Covered", section_title))
    term_rows = [[
        Paragraph("English Term", table_head),
        Paragraph("Spanish Contextual Choice", table_head),
        Paragraph("All Approved Spanish Equivalents", table_head),
    ]]
    for term in terms:
        english = escape(str(term.get("english", "")).strip())
        contextual = escape(str(term.get("spanish_contextual_choice", "")).strip())
        alternatives = "; ".join(
            str(x).strip() for x in (term.get("spanish_alternatives_from_sheet", []) or []) if str(x).strip()
        )
        term_rows.append([
            Paragraph(english, table_cell),
            Paragraph(escape(contextual), table_cell),
            Paragraph(escape(alternatives), table_cell),
        ])
    terms_table = Table(term_rows, colWidths=[1.7 * inch, 2.05 * inch, 2.65 * inch], repeatRows=1)
    terms_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5597")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C6D9")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.HexColor("#EEF3F9")]),
    ]))
    story.append(terms_table)
    story.append(Spacer(1, 0.12 * inch))

    if paired_segments:
        story.append(Paragraph("Paired Practice Script", section_title))
        pair_rows = [[Paragraph("English", table_head), Paragraph("Español", table_head)]]
        for pair in paired_segments:
            en = escape(str(pair.get("english", "")).strip())
            es = escape(str(pair.get("spanish", "")).strip())
            pair_rows.append([Paragraph(en, table_cell), Paragraph(es, table_cell)])
        pair_table = Table(pair_rows, colWidths=[3.15 * inch, 3.15 * inch], repeatRows=1)
        pair_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5597")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C6D9")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8FB")]),
        ]))
        story.append(pair_table)
        story.append(PageBreak())

    story.append(Paragraph("Full English Script", section_title))
    for block in iter_text_blocks(english_script):
        for line in iter_text_lines(block):
            story.append(Paragraph(escape(line), body_style))
        story.append(Spacer(1, 0.04 * inch))

    story.append(Paragraph("Full Spanish Script", section_title))
    for block in iter_text_blocks(spanish_script):
        for line in iter_text_lines(block):
            story.append(Paragraph(escape(line), body_style))
        story.append(Spacer(1, 0.04 * inch))

    try:
        doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=letter,
            leftMargin=0.7 * inch,
            rightMargin=0.7 * inch,
            topMargin=0.7 * inch,
            bottomMargin=0.7 * inch,
            title=f"Interpreter Practice - {data.get('script_id', '')}",
        )
        doc.build(story)
        if pdf_path.exists():
            print(f"Saved PDF for {script_path.stem} via pure-Python fallback: {pdf_path}")
            return pdf_path
    except Exception as e:
        print(f"Pure-Python single-script PDF fallback failed for {script_path.name}: {e}")

    return None


def generate_master_pdf_from_scripts(script_paths: list[Path], pdf_path: Path) -> Path | None:
    valid_payloads: list[dict[str, Any]] = []
    for script_path in script_paths:
        try:
            with script_path.open("r", encoding="utf-8") as f:
                data = json.load(f)
            if is_valid_script_json(data):
                valid_payloads.append(data)
        except Exception:
            continue

    if not valid_payloads:
        print("Pure-Python PDF fallback skipped: no valid script JSON payloads were available.")
        return None

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "PacketTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        spaceAfter=14,
    )
    subtitle_style = ParagraphStyle(
        "PacketSubtitle",
        parent=styles["BodyText"],
        fontName="Helvetica-Oblique",
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#555555"),
        spaceAfter=8,
    )
    center_small = ParagraphStyle(
        "CenterSmall",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    section_title = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        spaceBefore=10,
        spaceAfter=6,
    )
    script_title = ParagraphStyle(
        "ScriptTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=19,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    meta_style = ParagraphStyle(
        "Meta",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=13.5,
        spaceAfter=10,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=13.5,
        spaceAfter=6,
    )
    contents_style = ParagraphStyle(
        "Contents",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=13.5,
        leftIndent=12,
        bulletIndent=0,
        spaceAfter=3,
    )
    table_cell = ParagraphStyle(
        "TableCell",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        spaceAfter=0,
    )
    table_head = ParagraphStyle(
        "TableHead",
        parent=table_cell,
        fontName="Helvetica-Bold",
        textColor=colors.white,
        alignment=TA_CENTER,
    )

    story = []
    story.append(Spacer(1, 1.2 * inch))
    story.append(Paragraph("Combined Interpreter Practice Packet", title_style))
    story.append(Paragraph("Generated from script JSON files", subtitle_style))
    story.append(Paragraph(f"Included packets: {len(valid_payloads)}", center_small))
    story.append(Spacer(1, 0.22 * inch))
    story.append(Paragraph("Contents", section_title))
    for data in valid_payloads:
        line = f"{escape(str(data.get('script_id', '')))} - {escape(str(data.get('title', '')))}"
        story.append(Paragraph(line, contents_style, bulletText="-"))
    story.append(PageBreak())

    for idx, data in enumerate(valid_payloads):
        script_id = escape(str(data.get("script_id", "")))
        title = escape(str(data.get("title", "")))
        category = escape(str(data.get("category", "")))
        subcategory = escape(str(data.get("subcategory", "")))
        terms = data.get("terms_used", []) or []
        paired_segments = data.get("paired_segments", []) or []
        english_script = data.get("english_script", "")
        spanish_script = data.get("spanish_script", "")

        story.append(Paragraph(f"Interpreter Practice - {script_id}", script_title))
        story.append(Paragraph(
            f"<b>Title:</b> {title}<br/><b>Category:</b> {category}<br/><b>Subcategory:</b> {subcategory}",
            meta_style,
        ))

        story.append(Paragraph("Terms Covered", section_title))
        term_rows = [[
            Paragraph("English Term", table_head),
            Paragraph("Spanish Contextual Choice", table_head),
            Paragraph("All Approved Spanish Equivalents", table_head),
        ]]
        for term in terms:
            english = escape(str(term.get("english", "")).strip())
            contextual = escape(str(term.get("spanish_contextual_choice", "")).strip())
            alternatives = "; ".join(
                str(x).strip() for x in (term.get("spanish_alternatives_from_sheet", []) or []) if str(x).strip()
            )
            term_rows.append([
                Paragraph(english, table_cell),
                Paragraph(escape(contextual), table_cell),
                Paragraph(escape(alternatives), table_cell),
            ])
        terms_table = Table(term_rows, colWidths=[1.7 * inch, 2.05 * inch, 2.65 * inch], repeatRows=1)
        terms_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5597")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C6D9")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.HexColor("#EEF3F9")]),
        ]))
        story.append(terms_table)
        story.append(Spacer(1, 0.12 * inch))

        if paired_segments:
            story.append(Paragraph("Paired Practice Script", section_title))
            pair_rows = [[Paragraph("English", table_head), Paragraph("Español", table_head)]]
            for pair in paired_segments:
                en = escape(str(pair.get("english", "")).strip())
                es = escape(str(pair.get("spanish", "")).strip())
                pair_rows.append([Paragraph(en, table_cell), Paragraph(es, table_cell)])
            pair_table = Table(pair_rows, colWidths=[3.15 * inch, 3.15 * inch], repeatRows=1)
            pair_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5597")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C6D9")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8FB")]),
            ]))
            story.append(pair_table)
            story.append(PageBreak())

        story.append(Paragraph("Full English Script", section_title))
        for block in iter_text_blocks(english_script):
            for line in iter_text_lines(block):
                story.append(Paragraph(escape(line), body_style))
            story.append(Spacer(1, 0.04 * inch))

        story.append(Paragraph("Full Spanish Script", section_title))
        for block in iter_text_blocks(spanish_script):
            for line in iter_text_lines(block):
                story.append(Paragraph(escape(line), body_style))
            story.append(Spacer(1, 0.04 * inch))

        if idx < len(valid_payloads) - 1:
            story.append(PageBreak())

    try:
        doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=letter,
            leftMargin=0.7 * inch,
            rightMargin=0.7 * inch,
            topMargin=0.7 * inch,
            bottomMargin=0.7 * inch,
            title="Combined Interpreter Practice Packet",
        )
        doc.build(story)
        if pdf_path.exists():
            print(f"Saved combined PDF via pure-Python fallback: {pdf_path}")
            return pdf_path
    except Exception as e:
        print(f"Pure-Python PDF fallback failed: {e}")

    return None


def convert_docx_to_pdf(
    docx_path: Path,
    source_script_paths: list[Path] | None = None,
    label: str | None = None,
    source_script_path: Path | None = None,
) -> Path | None:
    pdf_path = docx_path.with_suffix(".pdf")
    doc_label = label or docx_path.name

    office_bin = shutil.which("soffice") or shutil.which("libreoffice")
    if office_bin:
        try:
            subprocess.run(
                [
                    office_bin,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(docx_path.parent),
                    str(docx_path),
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            if pdf_path.exists():
                print(f"Saved PDF for {doc_label}: {pdf_path}")
                return pdf_path
        except Exception as e:
            print(f"Direct PDF conversion failed with {Path(office_bin).name}: {e}")

    render_script = Path("/home/oai/skills/docx/render_docx.py")
    if render_script.exists():
        try:
            subprocess.run(
                [
                    "python",
                    str(render_script),
                    str(docx_path),
                    "--output_dir",
                    str(docx_path.parent),
                    "--emit_pdf",
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            if pdf_path.exists():
                print(f"Saved PDF for {doc_label} via render_docx.py: {pdf_path}")
                return pdf_path
        except Exception as e:
            print(f"Fallback PDF conversion via render_docx.py failed: {e}")

    if source_script_paths:
        pure_python_pdf = generate_master_pdf_from_scripts(source_script_paths, pdf_path)
        if pure_python_pdf:
            return pure_python_pdf

    if source_script_path:
        pure_python_pdf = generate_single_script_pdf(source_script_path, pdf_path)
        if pure_python_pdf:
            return pure_python_pdf

    print(
        "PDF conversion skipped: no LibreOffice executable was found in PATH, "
        "the render_docx.py fallback did not produce a PDF, and the pure-Python "
        "fallback was unavailable or failed."
    )
    return None



def generate_practice_packets(
    script_paths: list[Path],
    output_dir: Path,
    *,
    generate_individual_docx: bool = True,
    generate_individual_pdfs: bool = False,
    generate_master_docx: bool = False,
    generate_master_pdf: bool = False,
    highlight_terms: bool = True,
    master_docx_name: str = "combined_practice_packet.docx",
) -> list[Path]:
    """Generate DOCX/PDF practice packets from explicit JSON file paths.

    This is the Streamlit-friendly entry point. It avoids relying on the
    module-level scripts/ and docx_output/ folders, while preserving the
    existing standalone CLI behavior in main().
    """
    global OUTPUT_DIR, HIGHLIGHT_TERMS, MASTER_DOCX_NAME

    previous_output_dir = OUTPUT_DIR
    previous_highlight_terms = HIGHLIGHT_TERMS
    previous_master_docx_name = MASTER_DOCX_NAME

    OUTPUT_DIR = Path(output_dir)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    HIGHLIGHT_TERMS = bool(highlight_terms)
    MASTER_DOCX_NAME = master_docx_name or "combined_practice_packet.docx"

    generated: list[Path] = []
    valid_script_paths: list[Path] = []

    try:
        for path in script_paths:
            if Path(path).name == "index.json":
                continue
            try:
                with Path(path).open("r", encoding="utf-8") as f:
                    data = json.load(f)
            except Exception as exc:
                print(f"Skipping unreadable JSON {Path(path).name}: {exc}")
                continue

            if not is_valid_script_json(data):
                print(f"Skipping invalid or non-script JSON: {Path(path).name}")
                continue

            valid_script_paths.append(Path(path))

            if generate_individual_docx:
                docx_path = generate_docx(Path(path))
                if docx_path:
                    generated.append(docx_path)
                    if generate_individual_pdfs:
                        pdf_path = convert_docx_to_pdf(docx_path, label=docx_path.stem, source_script_path=Path(path))
                        if pdf_path:
                            generated.append(pdf_path)

        if generate_master_docx and valid_script_paths:
            master_docx = generate_master_docx_from_paths(valid_script_paths)
            if master_docx:
                generated.append(master_docx)
                if generate_master_pdf:
                    master_pdf = convert_docx_to_pdf(master_docx, valid_script_paths, label=master_docx.stem)
                    if master_pdf:
                        generated.append(master_pdf)

        return generated
    finally:
        OUTPUT_DIR = previous_output_dir
        HIGHLIGHT_TERMS = previous_highlight_terms
        MASTER_DOCX_NAME = previous_master_docx_name


def generate_master_docx_from_paths(script_paths: list[Path]) -> Path | None:
    """Alias kept separate from the Streamlit option name to avoid shadowing."""
    return generate_master_docx(script_paths)

def main() -> None:
    files = sorted(SCRIPTS_DIR.glob("*.json"))
    if not files:
        print("No JSON script files found in scripts/")
        return

    target_names = get_target_names()
    if target_names is not None:
        files = [p for p in files if p.name in target_names]

    count = 0
    generated_paths: list[Path] = []
    generated_docx_paths: list[Path] = []
    for path in files:
        if path.name == "index.json":
            continue
        result = generate_docx(path)
        if result:
            count += 1
            generated_paths.append(path)
            generated_docx_paths.append(result)
            if GENERATE_INDIVIDUAL_PDFS:
                convert_docx_to_pdf(result, label=result.stem, source_script_path=path)

    if GENERATE_MASTER_DOCX and generated_paths:
        master_docx = generate_master_docx(generated_paths)
        if master_docx and GENERATE_MASTER_PDF:
            convert_docx_to_pdf(master_docx, generated_paths, label="combined practice packet")

    print(f"Done. Generated {count} document(s) in {OUTPUT_DIR}/")


if __name__ == "__main__":
    main()
